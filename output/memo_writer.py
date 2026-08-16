"""
Generates the SS Investment Memo as a .docx file using python-docx.

Follows the exact section structure of the SS Investment Memo Template.
"""

import os

import config as cfg
from output import safe_filename
from output.page_budget import (PageBudget, MARGIN_X_IN, MARGIN_Y_IN,
                                PAGE_HEIGHT_IN, PAGE_MIN_PT, PAGE_WIDTH_IN,
                                paragraph_pt, table_pt)
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def generate_memo(property_name: str, cim_data, gate_results: list,
                  market_analysis: dict, physical_analysis: dict,
                  financial_analysis: dict, rent_analysis: dict,
                  scenario_results: dict, value_add: dict,
                  risk_analysis: dict, max_offer: dict,
                  va_results: dict = None, va_max_offer: dict = None,
                  checks: list = None, sources_uses: dict = None,
                  levered: dict = None, debt: dict = None,
                  levered_max_offer: dict = None,
                  assumption_fill_log: list = None,
                  assumption_register: list = None,
                  output_dir: str = ".") -> str:
    """
    Generate the SS Investment Memo .docx.

    Returns: path to generated file.
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    safe_name = safe_filename(property_name or "Unknown_Property")
    filename = f"SS_Investment_Memo_{safe_name}.docx"
    filepath = os.path.join(output_dir, filename)

    # ── Title Page ──────────────────────────────────────────────
    _add_title_page(doc, cim_data)

    # A portfolio-suspect CIM gets its caveat BEFORE section 1: every
    # number after this line may mix portfolio- and property-level values,
    # so the reader must meet the warning before the first number.
    _add_portfolio_warning(doc, cim_data)

    # ── Section 1: Investment Summary ───────────────────────────
    _add_section_1(doc, cim_data, gate_results, scenario_results, max_offer,
                   checks, assumption_fill_log)

    # ── Section 2: Market Overview ──────────────────────────────
    _add_section_2(doc, market_analysis)

    # ── Section 3: Property Description ─────────────────────────
    _add_section_3(doc, physical_analysis)

    # ── Section 4: Financial Analysis ───────────────────────────
    _add_section_4(doc, financial_analysis, cim_data)

    # ── Section 5: Unit Mix & Rent Analysis ─────────────────────
    _add_section_5(doc, rent_analysis)

    # ── Section 6: Valuation & Returns ──────────────────────────
    _add_section_6(doc, scenario_results, max_offer, sources_uses,
                   levered, debt, levered_max_offer)

    # ── Section 7: Value-Add Opportunities ──────────────────────
    _add_section_7(doc, value_add, va_results, va_max_offer)

    # ── Section 8: Risk Analysis ────────────────────────────────
    _add_section_8(doc, risk_analysis)

    # ── Section 9: Due Diligence Items ──────────────────────────
    _add_section_9(doc)

    # ── Section 10: Recommendation ──────────────────────────────
    _add_section_10(doc, gate_results, scenario_results, max_offer, risk_analysis, cim_data)

    # ── Appendix A: Assumptions Filled From Defaults ────────────
    # Last, because sections 1-10 carry their number in the heading and an
    # eleventh numbered section would renumber a document IC readers
    # navigate by number. Section 1 carries the count so a reader knows
    # to turn here (item T Category 4).
    _add_assumptions_appendix(doc, assumption_fill_log)

    # ── Appendix B: Assumption Register ─────────────────────────
    # Item T Category 6, and the item's acceptance criterion: every number
    # that moved an output, with the provenance that produced it, in one
    # auditable place. It CONTAINS Appendix A's rows — an invented value
    # appears here as one `fallback` row among the rest — so an auditor
    # who reads only this has still seen everything. A stays because
    # "what did the model invent?" is a sharper question than "what did
    # the model use?", and nine invented numbers inside a hundred and
    # forty do not read as an answer to it.
    _add_assumption_register(doc, assumption_register)

    doc.save(filepath)
    return filepath


# ── Section Builders ────────────────────────────────────────────────

def _hold_years(scenarios: dict, noi_key: str = "noi_projection") -> int:
    """Hold length these scenarios were computed on. Falls back to the
    config default rather than a literal 5, so the two cannot diverge."""
    for scen in (scenarios or {}).values():
        if isinstance(scen, dict):
            hold = scen.get("hold_years") or len(scen.get(noi_key) or [])
            if hold:
                return hold
    return cfg.DEFAULT_HOLD_YEARS


def _add_title_page(doc, cim_data):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\nSELF-STORAGE INVESTMENT MEMO\n\n")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name = cim_data.property_name or "Property Name TBD"
    run2 = p2.add_run(name)
    run2.bold = True
    run2.font.size = Pt(18)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    addr = cim_data.address or "Address TBD"
    city_state = f"{cim_data.city or 'City'}, {cim_data.state or 'ST'}"
    p3.add_run(f"{addr}\n{city_state}").font.size = Pt(14)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run("\n\nPrepared by CIM Analyst\nConfidential").font.size = Pt(11)

    doc.add_page_break()


def _add_portfolio_warning(doc, cim_data):
    """The portfolio caveat, before the first number (page 1 of the body).

    Renders THE shared sentence (`extract.portfolio.warning_text`) so the
    memo cannot drift from the run warnings the results page shows, then
    the evidence lines — this is the analyst-facing document, so the
    evidence belongs here. No-op for a single-asset CIM.
    """
    signal = getattr(cim_data, "portfolio_signal", None)
    if not signal:
        return

    from extract.portfolio import warning_text

    p = doc.add_paragraph()
    run = p.add_run("PORTFOLIO CIM — READ FIRST: " + warning_text())
    run.bold = True
    run.font.color.rgb = RGBColor(153, 0, 0)

    for ev in signal.get("evidence", []):
        doc.add_paragraph(str(ev), style="List Bullet")


def _add_section_1(doc, cim_data, gate_results, scenario_results, max_offer,
                   checks=None, assumption_fill_log=None):
    doc.add_heading("1. Investment Summary", level=1)

    # Key metrics table
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"

    metrics = [
        ("Property", cim_data.property_name or "TBD"),
        ("Location", f"{cim_data.city or 'TBD'}, {cim_data.state or 'TBD'}"),
        ("Asking Price", _fmt_currency(cim_data.asking_price)),
        ("NRSF", _fmt_number(cim_data.nrsf, suffix=" SF")),
        ("Total Units", str(cim_data.total_units or "TBD")),
        ("Physical Occupancy", _fmt_pct(cim_data.physical_occupancy)),
        ("Economic Occupancy", _fmt_pct(cim_data.economic_occupancy)),
        ("Price / SF", _fmt_currency(cim_data.price_per_sf)),
        ("Year Built", str(cim_data.year_built or "TBD")),
    ]
    for label, val in metrics:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = val

    _add_occupancy_spread_note(doc, cim_data)

    _add_model_checks(doc, checks)

    _add_fill_count(doc, assumption_fill_log)

    doc.add_paragraph()

    # Gate summary
    doc.add_heading("Screening Gates", level=2)
    gate_table = doc.add_table(rows=1, cols=4)
    gate_table.style = "Light Grid Accent 1"
    gh = gate_table.rows[0].cells
    gh[0].text = "Gate"
    gh[1].text = "Threshold"
    gh[2].text = "Actual"
    gh[3].text = "Result"

    for g in gate_results:
        row = gate_table.add_row().cells
        row[0].text = g["name"]
        row[1].text = str(g["threshold"])
        row[2].text = str(g["actual"])
        row[3].text = g["result"]

    doc.add_paragraph()

    # Returns snapshot
    if scenario_results:
        doc.add_heading("Returns Snapshot (Unlevered)", level=2)
        ret_table = doc.add_table(rows=1, cols=4)
        ret_table.style = "Light Grid Accent 1"
        rh = ret_table.rows[0].cells
        rh[0].text = "Metric"
        rh[1].text = "Bear"
        rh[2].text = "Base"
        rh[3].text = "Bull"

        ret_hold = _hold_years(scenario_results)
        for label, key in [("Yr1 Yield on Cost", "yield_on_cost"),
                           (f"{ret_hold}-Yr IRR", "irr"),
                           (f"{ret_hold}-Yr MOIC", "moic")]:
            row = ret_table.add_row().cells
            row[0].text = label
            for i, scen in enumerate(["bear", "base", "bull"]):
                val = scenario_results.get(scen, {}).get(key)
                if key == "moic":
                    row[i + 1].text = f"{val:.2f}x" if val else "N/A"
                else:
                    row[i + 1].text = _fmt_pct(val)

    # Max offer
    if max_offer:
        doc.add_paragraph()
        mp = max_offer.get("max_price")
        doc.add_paragraph(
            f"Maximum Offer Price (for "
            f"{max_offer.get('target_irr', cfg.SOLVER_TARGET_IRR):.0%} "
            f"Base Case IRR): {_fmt_currency(mp)}"
        ).bold = True


def _add_model_checks(doc, checks):
    """Model error-check register, section 1. Findings only — the full
    register (passes and not-testable rows included) is the Excel Checks
    sheet. A memo that lists ten green checks buries the one red one."""
    if not checks:
        return

    failed = [c for c in checks if c.get("status") == "fail"]
    doc.add_paragraph()
    if not failed:
        tested = sum(1 for c in checks if c.get("status") == "pass")
        skipped = len(checks) - tested
        doc.add_paragraph(
            f"Model checks: {tested} of {len(checks)} integrity checks passed "
            f"and none were flagged"
            + (f"; {skipped} were not testable from the data supplied."
               if skipped else ".")
        )
        return

    doc.add_heading("Model Checks", level=2)
    doc.add_paragraph(
        f"{len(failed)} of {len(checks)} integrity checks flagged. Blocking "
        f"findings were accepted by the analyst with the discrepancy recorded; "
        f"advisory findings do not stop the model but change how its outputs "
        f"should be read."
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Check"
    hdr[1].text = "Severity"
    hdr[2].text = "Finding"
    for c in failed:
        row = table.add_row().cells
        row[0].text = c.get("label") or c.get("id") or ""
        row[1].text = (c.get("severity") or "").title()
        row[2].text = c.get("message") or ""


def _add_fill_count(doc, assumption_fill_log):
    """One sentence in section 1 naming how many inputs this run invented.

    The appendix carries the detail, but an appendix nobody is told about
    is a stamp nobody reads — the exact failure item T exists to close.
    It sits beside the check register because both answer "how far should
    I trust the numbers below", and a reader who sees neither line has
    been told the CIM supplied everything.
    """
    if not assumption_fill_log:
        return
    n = len(assumption_fill_log)
    doc.add_paragraph(
        f"{n} model input{'' if n == 1 else 's'} {'was' if n == 1 else 'were'} "
        f"not stated in the CIM and {'was' if n == 1 else 'were'} filled from "
        f"a default or a benchmark. Every one is listed with its source in "
        f"Appendix A; the returns below are computed on them."
    )


def _add_assumptions_appendix(doc, assumption_fill_log):
    """Appendix A — every value this run invented, with its provenance.

    Item T Category 4. Rendered only when the log is non-empty, matching
    `_add_model_checks`: an appendix reading "none" on a complete CIM is
    a page IC readers learn to skip, and the section-1 sentence above is
    already silent in that case.
    """
    if not assumption_fill_log:
        return

    from analysis.fills import format_value, from_dicts

    doc.add_page_break()
    doc.add_heading("Appendix A. Assumptions Filled From Defaults", level=1)
    doc.add_paragraph(
        "The CIM did not state the inputs below. The model used the value in "
        "the second column, from the source in the third — not a figure from "
        "this deal. Each one moves the returns in section 6, so an assumption "
        "an IC reader disagrees with is an assumption to change on the "
        "assumptions page and re-run, not a number to discount by eye."
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Input"
    hdr[1].text = "Value Used"
    hdr[2].text = "Source"
    hdr[3].text = "What This Means"
    for fill in from_dicts(assumption_fill_log):
        row = table.add_row().cells
        row[0].text = fill.field
        row[1].text = format_value(fill)
        row[2].text = fill.source_label
        row[3].text = fill.label


def _add_assumption_register(doc, assumption_register):
    """Appendix B — every number that moved an output, and who chose it.

    Two tables, because they answer to two readers. The first lists only
    the rows something other than the shipped model produced — a deal
    entry, a dated settings row, a fallback — which is what is unusual
    about THIS run and is typically ten to twenty lines. The second is the
    whole register, grouped by subject, for the auditor.

    Neither table omits anything. A "defaults suppressed for brevity"
    register asks the reader to trust that absence means default, which is
    the same act of faith item T exists to end — so the bulk goes to the
    back of the document rather than out of it.
    """
    if not assumption_register:
        return

    from analysis.assumptions import (CHOSEN, PROVENANCE_LABELS, format_value,
                                      from_dicts, summarize)

    rows = from_dicts(assumption_register)
    counts = summarize(rows)

    doc.add_page_break()
    doc.add_heading("Appendix B. Assumption Register", level=1)
    doc.add_paragraph(
        f"Every number this analysis used, with its source. Of "
        f"{counts['total']} assumptions, {counts['chosen']} came from "
        f"something other than the model's shipped defaults: "
        f"{counts['deal']} entered for this deal, {counts['settings']} from "
        f"a dated settings override, and {counts['fallback']} filled in "
        f"because the CIM did not state a value. Those are listed first. "
        f"The full register follows — nothing is omitted from it, so an "
        f"assumption absent below is an assumption this run did not use."
    )

    chosen = [r for r in rows if r.provenance in CHOSEN]
    doc.add_heading("B.1 Assumptions not taken from the model defaults",
                    level=2)
    if not chosen:
        doc.add_paragraph(
            "None. Every number below is the model's shipped default, and "
            "every input came from the CIM as stated.")
    else:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Assumption"
        hdr[1].text = "Value Used"
        hdr[2].text = "Source"
        hdr[3].text = "Replaced"
        for row in chosen:
            cells = table.add_row().cells
            cells[0].text = row.label
            cells[1].text = format_value(row)
            cells[2].text = PROVENANCE_LABELS.get(row.provenance,
                                                  row.provenance)
            cells[3].text = (_register_prior(row) if row.was is not None
                             else (row.detail or "—"))

    doc.add_heading("B.2 Full register", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Group"
    hdr[1].text = "Assumption"
    hdr[2].text = "Value Used"
    hdr[3].text = "Source"
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row.group
        cells[1].text = row.label
        cells[2].text = format_value(row)
        cells[3].text = PROVENANCE_LABELS.get(row.provenance, row.provenance)


def _register_prior(row):
    """What the winning value displaced, rendered in its own units.

    Built by swapping the value on a copy rather than by a second
    formatter: `was` is the same quantity in the same unit as `value`, and
    a separate rendering path is how one of them ends up printing 0.8
    where the other prints 80%.
    """
    import dataclasses

    from analysis.assumptions import format_value

    return format_value(dataclasses.replace(row, value=row.was))


def _add_occupancy_spread_note(doc, cim_data):
    """Economic-vs-physical occupancy read — the fastest value-add screen
    and the most common place a CIM hides weakness."""
    from config import GATES

    phys = cim_data.physical_occupancy
    econ = cim_data.economic_occupancy
    if phys is None and econ is None:
        return

    doc.add_paragraph()
    if econ is None:
        doc.add_paragraph(
            "Occupancy check: the CIM quotes physical occupancy only. Request "
            "economic occupancy (collected rent vs gross potential at street "
            "rates) before proceeding — a broker quoting a single occupancy "
            "figure is almost always quoting physical, and the spread between "
            "the two is where CIMs hide weakness."
        )
        return
    if phys is None:
        doc.add_paragraph(
            f"Occupancy check: economic occupancy of {econ:.1%} stated without "
            f"physical occupancy — request the physical figure to compute the spread."
        )
        return

    spread = phys - econ
    if spread >= GATES["econ_phys_spread_flag"]:
        doc.add_paragraph(
            f"Occupancy check: economic occupancy trails physical by "
            f"{spread * 100:.0f} pts ({phys:.1%} physical vs {econ:.1%} economic). "
            f"A spread this wide signals revenue leakage — concessions, "
            f"delinquency, or in-place rents below street — and is the primary "
            f"mismanagement value-add screen. Decompose the spread from the "
            f"rent roll during diligence."
        )
    else:
        doc.add_paragraph(
            f"Occupancy check: economic occupancy of {econ:.1%} vs physical of "
            f"{phys:.1%} ({spread * 100:.0f}-pt spread) — within normal operating "
            f"range; upside must come from rate growth or expense control rather "
            f"than collections recovery."
        )


def _add_section_2(doc, market):
    doc.add_heading("2. Market Overview", level=1)

    demos = market.get("demographics", {})
    doc.add_heading("Demographics", level=2)
    doc.add_paragraph(demos.get("pop_narrative", "TBD"))
    doc.add_paragraph(demos.get("hhi_narrative", "TBD"))

    msa = market.get("msa_info", {})
    doc.add_heading("MSA Classification", level=2)
    doc.add_paragraph(msa.get("narrative", "TBD"))

    supply = market.get("supply_assessment", {})
    doc.add_heading("Supply Assessment", level=2)
    doc.add_paragraph(supply.get("narrative", "TBD"))

    demand = market.get("demand_drivers", {})
    positives = demand.get("positives", [])
    negatives = demand.get("negatives", [])
    if positives:
        doc.add_heading("Demand Positives", level=2)
        for p in positives:
            doc.add_paragraph(p, style="List Bullet")
    if negatives:
        doc.add_heading("Demand Concerns", level=2)
        for n in negatives:
            doc.add_paragraph(n, style="List Bullet")

    doc.add_paragraph(f"\nOverall Market Rating: {market.get('overall_rating', 'TBD')}")


def _add_section_3(doc, physical):
    doc.add_heading("3. Property Description", level=1)

    profile = physical.get("property_profile", {})
    for key, label in [("property_name", "Property"), ("address", "Address"),
                       ("city_state", "City/State"), ("year_built", "Year Built"),
                       ("acreage", "Acreage"), ("nrsf", "NRSF"),
                       ("total_units", "Total Units"), ("cc_pct", "Climate-Controlled %"),
                       ("physical_occupancy", "Physical Occupancy"),
                       ("economic_occupancy", "Economic Occupancy")]:
        val = profile.get(key)
        if val is not None:
            if isinstance(val, float) and val < 1:
                val = f"{val:.1%}"
            elif isinstance(val, float):
                val = f"{val:,.0f}"
            doc.add_paragraph(f"{label}: {val}")

    doc.add_paragraph(profile.get("age_narrative", ""))
    doc.add_paragraph(profile.get("condition_note", ""))

    # Replacement cost
    repl = physical.get("replacement_cost", {})
    if repl.get("estimable"):
        doc.add_heading("Replacement Cost Estimate", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        table.rows[0].cells[0].text = "Component"
        table.rows[0].cells[1].text = "Cost"

        # Use facility-type detail rows if available, else legacy
        type_details = repl.get("facility_type_details", [])
        items = []
        if type_details:
            for td in type_details:
                items.append((f"{td['type']} Hard Cost ({td['sf']:,.0f} SF)", td["hard_cost"]))
                if td["site_cost"] > 0:
                    items.append((f"{td['type']} Site Work", td["site_cost"]))
        else:
            items.append(("Non-CC Hard Cost", repl.get("non_cc_cost")))
            items.append(("CC Hard Cost", repl.get("cc_cost")))
            items.append(("Site Work", repl.get("site_work")))
        items.extend([
            ("Soft Costs", repl.get("soft_costs")),
            ("Developer Profit", repl.get("dev_profit")),
            ("Total Replacement Cost", repl.get("total_replacement")),
        ])
        for label, val in items:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = _fmt_currency(val)

        # Assumptions disclosure
        doc.add_heading("Replacement Cost Assumptions", level=3)
        type_details = repl.get("facility_type_details", [])
        if type_details:
            assumptions_table = doc.add_table(rows=1, cols=3)
            assumptions_table.style = "Light Grid Accent 1"
            assumptions_table.rows[0].cells[0].text = "Facility Type"
            assumptions_table.rows[0].cells[1].text = "Hard Cost $/SF"
            assumptions_table.rows[0].cells[2].text = "Site Work $/SF"
            for td in type_details:
                arow = assumptions_table.add_row().cells
                arow[0].text = td["type"]
                arow[1].text = f"${td['hard_rate']:,.0f}"
                arow[2].text = f"${td['site_rate']:,.0f}" if td["site_rate"] > 0 else "Incl."
        soft_pct = repl.get("soft_cost_pct", 0)
        dev_pct = repl.get("dev_profit_pct", 0)
        doc.add_paragraph(
            f"Soft costs assumed at {soft_pct:.0%} of hard + site costs. "
            f"Developer profit assumed at {dev_pct:.0%} of total development cost. "
            f"Hard cost rates represent midpoints of benchmark ranges based on "
            f"2025/2026 construction cost data for each facility type."
        )

    comp = physical.get("price_vs_replacement", {})
    if comp.get("narrative"):
        doc.add_paragraph()
        doc.add_paragraph(comp["narrative"])


def _add_section_4(doc, fin, cim_data):
    doc.add_heading("4. Financial Analysis", level=1)

    # Income summary
    income = fin.get("income_summary", {})
    doc.add_heading("Income Summary", level=2)
    for label, key in [("Gross Potential Rent", "gpr"), ("Vacancy", "vacancy_loss"),
                       ("Effective Gross Revenue", "egr"),
                       ("Other Income", "other_income"),
                       ("Total Revenue", "total_revenue")]:
        val = income.get(key)
        doc.add_paragraph(f"{label}: {_fmt_currency(val)}")

    # Expense analysis
    doc.add_heading("Expense Benchmarking", level=2)
    exp = fin.get("expense_analysis", {})
    lines = exp.get("lines", [])
    if lines:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdrs = table.rows[0].cells
        hdrs[0].text = "Category"
        hdrs[1].text = "CIM Value"
        hdrs[2].text = "$/NRSF"
        hdrs[3].text = "Benchmark"
        hdrs[4].text = "Flag"

        for line in lines:
            row = table.add_row().cells
            category = line["category"]
            if line.get("source") == "analyst":
                category += " (analyst)"
            row[0].text = category
            row[1].text = _fmt_currency(line.get("cim_value"))
            pn = line.get("per_nrsf")
            row[2].text = f"${pn:.2f}" if pn else "N/A"
            br = line.get("benchmark_range")
            row[3].text = f"${br[0]:.2f}-${br[1]:.2f}" if br else "N/A"
            row[4].text = line.get("flag") or ""

    # Adjustments
    adjustments = fin.get("adjustments", [])
    if adjustments:
        doc.add_heading("Analyst Adjustments", level=2)
        for adj in adjustments:
            doc.add_paragraph(adj, style="List Bullet")

    # Adjusted NOI
    adj_noi = fin.get("adjusted_ttm_noi", {})
    doc.add_heading("Adjusted TTM NOI", level=2)
    doc.add_paragraph(adj_noi.get("narrative", "TBD"))


def _add_section_5(doc, rent):
    doc.add_heading("5. Unit Mix & Rent Analysis", level=1)
    doc.add_paragraph(rent.get("narrative", "Unit mix data not available."))

    summary = rent.get("unit_mix_summary", [])
    if summary:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdrs = table.rows[0].cells
        hdrs[0].text = "Size"
        hdrs[1].text = "Count"
        hdrs[2].text = "SF"
        hdrs[3].text = "Rate/Mo"
        hdrs[4].text = "$/SF/Mo"

        for s in summary:
            row = table.add_row().cells
            row[0].text = s.get("size_label") or ""
            row[1].text = str(s.get("count") or "")
            row[2].text = f"{s.get('unit_sf') or 0:,.0f}"
            row[3].text = _fmt_currency(s.get("monthly_rate"))
            r = s.get("rate_per_sf")
            row[4].text = f"${r:.2f}" if r else "N/A"

    gap = rent.get("rent_gap_analysis", {})
    if gap.get("narrative"):
        doc.add_heading("Rent Gap to Market", level=2)
        doc.add_paragraph(gap["narrative"])


def _add_section_6(doc, scenario_results, max_offer, sources_uses=None,
                   levered=None, debt=None, levered_max_offer=None):
    doc.add_heading("6. Valuation & Returns", level=1)

    if not scenario_results:
        doc.add_paragraph("Scenario analysis not available — insufficient data.")
        return

    for scen_name in ("bear", "base", "bull"):
        s = scenario_results.get(scen_name, {})
        doc.add_heading(f"{scen_name.title()} Case", level=2)

        noi_proj = s.get("noi_projection", [])
        if noi_proj:
            table = doc.add_table(rows=2, cols=len(noi_proj) + 1)
            table.style = "Light Grid Accent 1"
            table.rows[0].cells[0].text = "Year"
            table.rows[1].cells[0].text = "NOI"
            for i, noi in enumerate(noi_proj):
                table.rows[0].cells[i + 1].text = f"Yr {i + 1}"
                table.rows[1].cells[i + 1].text = _fmt_currency(noi)

        hold = s.get("hold_years") or len(noi_proj) or cfg.DEFAULT_HOLD_YEARS
        doc.add_paragraph(f"Entry Cap: {_fmt_pct(s.get('entry_cap'))}")
        doc.add_paragraph(f"Exit Cap: {_fmt_cap(s.get('exit_cap'))}")
        # The exit cap is derived, not typed. Printing only the result
        # would make the memo less auditable than the constant it
        # replaced, so the terms go in beside it.
        derivation = _exit_cap_derivation(s)
        if derivation:
            doc.add_paragraph(f"  {derivation}")
        # Under the trailing default the capitalized NOI is the last
        # column of the table above, so nothing is added and the memo is
        # unchanged. Under "forward" it is year N+1's — absent from that
        # table — and an exit value printed beside a table it does not
        # tie to is the kind of unauditable figure this memo refuses
        # everywhere else (`config.EXIT_NOI_CONVENTION`, decision 5).
        if s.get("exit_noi_convention") == "forward":
            doc.add_paragraph(
                f"Exit NOI (forward, Yr {hold + 1}): "
                f"{_fmt_currency(s.get('exit_noi'))}")
        doc.add_paragraph(f"Exit Value (gross): {_fmt_currency(s.get('exit_value'))}")
        # Costs are stated, not folded silently into the return: an IRR
        # quoted net of costs the reader can't see is not auditable.
        doc.add_paragraph(
            f"Acquisition Closing Costs: {_fmt_currency(s.get('acquisition_cost'))}")
        doc.add_paragraph(
            f"Disposition Costs: {_fmt_currency(s.get('disposition_cost'))}")
        doc.add_paragraph(
            f"Net Sale Proceeds: {_fmt_currency(s.get('net_exit_proceeds'))}")
        doc.add_paragraph(f"{hold}-Year IRR (net of costs): {_fmt_pct(s.get('irr'))}")
        doc.add_paragraph(f"{hold}-Year MOIC: {s.get('moic', 0):.2f}x" if s.get("moic") else "MOIC: N/A")
        doc.add_paragraph(f"Yield on Cost: {_fmt_pct(s.get('yield_on_cost'))}")

    _add_sources_uses(doc, sources_uses)
    _add_levered_returns(doc, levered, debt, scenario_results)

    # Max offer
    if max_offer or levered_max_offer:
        doc.add_heading("Maximum Offer Price", level=2)
    if max_offer:
        doc.add_paragraph(
            f"At a target "
            f"{max_offer.get('target_irr', cfg.SOLVER_TARGET_IRR):.0%} "
            f"base case unlevered IRR, "
            f"the maximum offer price is {_fmt_currency(max_offer.get('max_price'))} "
            f"(implied entry cap: {_fmt_pct(max_offer.get('implied_entry_cap'))})."
        )
    _add_levered_max_offer(doc, levered_max_offer)


def _add_levered_max_offer(doc, levered_max_offer):
    """The levered max offer (item E4), beside the unlevered one.

    Beside and not instead of: the unlevered figure is the price the
    primary 10% gate is measured against, and the two answer different
    questions — "what can the property carry" versus "what can the fund
    pay and still clear its LP net target". A memo showing only the
    second would leave the primary screen with no price of its own.

    Silent when the deal priced no loan, for the same reason
    `_add_levered_returns` returns early: a block of N/A reads as a
    failed calculation rather than an absent one.
    """
    offer = levered_max_offer or {}
    if not offer.get("max_price"):
        return

    doc.add_paragraph(
        f"At a target {offer.get('target_irr', 0.15):.0%} LP NET IRR — after "
        f"debt service, the asset-management fee and the GP promote — the "
        f"maximum offer price is {_fmt_currency(offer.get('max_price'))} "
        f"(implied entry cap: {_fmt_pct(offer.get('implied_entry_cap'))}). "
        f"At that price the deal supports "
        f"{_fmt_currency(offer.get('senior_debt'))} of senior debt "
        f"({_fmt_pct(offer.get('ltv'))} LTV, bound by "
        f"{_binding_label(offer)}) and requires "
        f"{_fmt_currency(offer.get('total_equity'))} of equity.")

    # The bisection assumes LP net IRR falls as price rises. It does over
    # every range measured (see model/solver.py), but an observed
    # inversion means this price may be the wrong root, and a memo must
    # not print a suspect number as a clean one.
    if offer.get("monotonicity_warning"):
        doc.add_paragraph(f"WARNING: {offer['monotonicity_warning']}")
    if not offer.get("converged"):
        doc.add_paragraph(
            "WARNING: the levered solver did not converge to its target "
            "within the iteration budget — treat this price as indicative.")

    # Same rule as every other LP net figure in this memo: no LP net IRR
    # without its assumption stamp (CLAUDE.md key design decision 7). A
    # price derived FROM an LP net IRR inherits the requirement.
    stamp = offer.get("assumption_stamp") or []
    if stamp:
        doc.add_paragraph(
            "This price is computed on the levered assumption set stated "
            "under Levered Returns above: "
            + "; ".join(row.get("label", "") for row in stamp) + ".")


def _binding_label(offer) -> str:
    from model.debt import CONSTRAINT_LABELS
    key = (offer or {}).get("binding_constraint")
    return CONSTRAINT_LABELS.get(key, key or "n/a")


def _add_sources_uses(doc, sources_uses):
    """Capital stack under section 6 — what the deal costs and where the
    money comes from. Every figure is read off build_sources_uses; nothing
    is recomputed here, so the memo cannot state a basis the model did not
    use. Senior debt and financing costs read $0 until item E1 sizes a
    loan, and they are PRINTED as zero rather than omitted, so a reader
    can see this is an all-equity underwrite rather than guess."""
    if not sources_uses:
        return
    doc.add_heading("Sources & Uses", level=2)

    for title, lines, total_key in (
            ("Uses", sources_uses.get("uses"), "total_uses"),
            ("Sources", sources_uses.get("sources"), "total_sources")):
        rows = list(lines or [])
        if not rows:
            continue
        table = doc.add_table(rows=len(rows) + 2, cols=3)
        table.style = "Light Grid Accent 1"
        header = table.rows[0].cells
        header[0].text = title
        header[1].text = "Amount"
        header[2].text = "% of Total"
        total = sources_uses.get(total_key) or 0
        for i, line in enumerate(rows, start=1):
            amount = line.get("amount") or 0
            cells = table.rows[i].cells
            cells[0].text = line.get("label") or ""
            cells[1].text = _fmt_currency(amount)
            cells[2].text = _fmt_pct(amount / total) if total else "N/A"
        last = table.rows[len(rows) + 1].cells
        last[0].text = f"Total {title}"
        last[1].text = _fmt_currency(total)
        last[2].text = _fmt_pct(1.0) if total else "N/A"
        doc.add_paragraph()

    doc.add_paragraph(
        f"Total equity required: "
        f"{_fmt_currency(sources_uses.get('total_equity'))} "
        f"(GP co-invest {_fmt_pct(sources_uses.get('gp_coinvest_pct'))}: "
        f"{_fmt_currency(sources_uses.get('gp_equity'))}; LP "
        f"{_fmt_currency(sources_uses.get('lp_equity'))}).")
    if not sources_uses.get("balanced"):
        doc.add_paragraph(
            f"WARNING: Sources and Uses are out of balance by "
            f"{_fmt_currency(abs(sources_uses.get('delta') or 0))}. The "
            f"returns above are computed on the Uses figure.")


def _add_levered_returns(doc, levered, debt, scenario_results):
    """The levered second lens under section 6 (item E3b).

    A level-2 subsection, not a new numbered section: ten numbered
    sections are referenced from the recommendation text and the CLI
    summary, and renumbering them to insert a presentation block is churn
    with a real chance of an off-by-one.

    Every figure is read off the persisted `levered` / `debt` payload —
    nothing is recomputed, so the memo cannot state a return the model did
    not produce. Absent entirely on a deal that priced no loan, which is
    why this returns early instead of printing a table of N/A.
    """
    levered = levered or {}
    base = levered.get("base") or {}
    if not base:
        return
    debt = debt or {}
    terms = debt.get("terms") or {}
    hold = _hold_years(scenario_results)

    doc.add_heading("Levered Returns (LP Net)", level=2)

    from model.debt import binding_constraint_label, displayed_rate

    rate = displayed_rate(terms)
    doc.add_paragraph(
        f"Senior loan {_fmt_currency(debt.get('loan'))} at "
        f"{_fmt_pct(rate)}, {_fmt_number(terms.get('amort_years'))}-year "
        f"amortization, {_fmt_number(terms.get('io_months'))} months IO, "
        f"{_fmt_number(terms.get('term_years'))}-year term. Sized off the "
        f"base case at the lesser of LTV, DSCR and debt yield; bound by "
        f"{binding_constraint_label(debt)} at "
        f"{_fmt_pct(debt.get('ltv'))} LTV, "
        f"{_fmt_x(debt.get('dscr_year_1'))} Year-1 DSCR and "
        f"{_fmt_pct(debt.get('debt_yield'))} debt yield. The SAME loan is "
        f"carried through all three scenarios — sizing per scenario would "
        f"hand the bear case a smaller loan and flatten the downside it "
        f"exists to show.")

    rows = [("LP Net IRR", "lp_net_irr", _fmt_pct),
            ("LP MOIC", "lp_moic", _fmt_x),
            ("GP Promote", "gp_promote", _fmt_currency),
            ("AM Fee (total)", "am_fee_total", _fmt_currency),
            ("Equity Required", "total_equity", _fmt_currency)]
    table = doc.add_table(rows=len(rows) + 1, cols=4)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].text = f"{hold}-Year Levered"
    for i, scen in enumerate(("bear", "base", "bull"), start=1):
        header[i].text = scen.title()
    for r, (label, key, fmt) in enumerate(rows, start=1):
        cells = table.rows[r].cells
        cells[0].text = label
        for i, scen in enumerate(("bear", "base", "bull"), start=1):
            cells[i].text = fmt((levered.get(scen) or {}).get(key))
    doc.add_paragraph()

    # Leverage is allowed to be dilutive, and on config defaults it often
    # is. Saying so in the memo stops a reader treating an LP net IRR
    # below the unlevered screen as an arithmetic error.
    dilutive = [scen.title() for scen in ("bear", "base", "bull")
                if (levered.get(scen) or {}).get("lp_net_irr") is not None
                and ((scenario_results or {}).get(scen) or {}).get("irr")
                is not None
                and (levered[scen]["lp_net_irr"]
                     < scenario_results[scen]["irr"])]
    if dilutive:
        doc.add_paragraph(
            f"Leverage is DILUTIVE in the {', '.join(dilutive)} case(s): LP "
            f"net IRR is below the unlevered IRR because the loan constant "
            f"sits above the deal's yield on cost. This is an outcome, not "
            f"an error.")
    if debt.get("matures_before_exit"):
        doc.add_paragraph(
            "WARNING: the loan matures before the hold ends. These figures "
            "amortize straight past maturity — no refinancing, rate reset "
            "or prepayment cost is modeled.")
    if base.get("called_capital_after_close"):
        doc.add_paragraph(
            f"WARNING: this deal calls "
            f"{_fmt_currency(base.get('capital_calls_total'))} of capital "
            f"after close "
            f"({_fmt_currency(base.get('reserve_drawn_total'))} covered by "
            f"the operating reserve). Called capital accrues preferred "
            f"return from the following period.")
    unrecovered = (base.get("waterfall") or {}).get("unrecovered_promote")
    if unrecovered:
        doc.add_paragraph(
            f"WARNING: {_fmt_currency(unrecovered)} of promote was paid "
            f"before a later capital call and the deal ends short of "
            f"capital plus preferred return. These fund terms carry no "
            f"clawback, so the GP keeps it.")

    # Not optional and not a footnote: every row is an LPA convention
    # that moves the LP net IRR printed above, and the stamp says which
    # ones the operator has actually read.
    doc.add_heading("Levered Assumptions", level=3)
    for row in base.get("assumption_stamp") or []:
        doc.add_paragraph(f"{row.get('label')} — {row.get('question')}",
                          style="List Bullet")


def _add_section_7(doc, value_add, va_results=None, va_max_offer=None):
    doc.add_heading("7. Value-Add Opportunities", level=1)
    doc.add_paragraph(value_add.get("narrative", "No opportunities identified."))

    # Qualitative opportunities table
    opps = value_add.get("opportunities", [])
    if opps:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdrs = table.rows[0].cells
        hdrs[0].text = "#"
        hdrs[1].text = "Opportunity"
        hdrs[2].text = "Est. Annual Impact"
        hdrs[3].text = "Timeline"

        for o in opps:
            row = table.add_row().cells
            row[0].text = str(o.get("priority") or "")
            row[1].text = o.get("opportunity") or ""
            impact = o.get("estimated_annual_impact") or 0
            row[2].text = _fmt_currency(impact) if impact else "TBD"
            row[3].text = o.get("timeline") or ""

    # Value-Add Financial Model (if available)
    if va_results:
        doc.add_heading("Value-Add Financial Model", level=2)

        base = va_results.get("base", {})
        in_place = base.get("in_place_rent_psf")
        market = base.get("market_rent_psf")
        current_occ = base.get("current_occupancy")
        target_occ = base.get("target_occupancy")

        # A rent gap is only a measurement when a market rent was
        # measured. With none, this engine sets market equal to in-place
        # and the sentence below would print "0% rent gap" — which reads
        # as "rents are already at market", the opposite of "we do not
        # know what market is" (item T Category 4). Suppressed, not
        # zeroed: a number that cannot be misread is worth more than one
        # that can.
        if base.get("rent_ramp_excluded"):
            doc.add_paragraph(
                f"Rent ramp excluded — no market-rent data. This case is an "
                f"occupancy ramp only, from {_fmt_pct(current_occ)} to a "
                f"target of {_fmt_pct(target_occ)}; no upside from pushing "
                f"rents toward market is in the returns below."
            )
        elif in_place and market:
            doc.add_paragraph(
                f"In-place rent of ${in_place:.2f}/SF/mo vs market of ${market:.2f}/SF/mo "
                f"represents a {((market - in_place) / in_place):.0%} rent gap. "
                f"Current occupancy of {_fmt_pct(current_occ)} with "
                f"target stabilization at {_fmt_pct(target_occ)}."
            )

        # VA scenario comparison table
        doc.add_heading("Value-Add Returns (Unlevered)", level=3)
        va_table = doc.add_table(rows=1, cols=4)
        va_table.style = "Light Grid Accent 1"
        vh = va_table.rows[0].cells
        vh[0].text = "Metric"
        vh[1].text = "Bear"
        vh[2].text = "Base"
        vh[3].text = "Bull"

        for label, key, fmt_fn in [
            ("Months to Stabilize", "months_to_stabilize",
             lambda v: str(int(v)) if v else "N/A"),
            ("Stabilized NOI", "stabilized_noi",
             lambda v: _fmt_currency(v)),
            ("Stabilized Yield/Cost", "yield_on_cost",
             lambda v: _fmt_pct(v)),
            ("Unlevered IRR (net of costs)", "irr",
             lambda v: _fmt_pct(v)),
            ("MOIC", "moic",
             lambda v: f"{v:.2f}x" if v else "N/A"),
            ("Development Spread", "development_spread",
             lambda v: f"{v*100:.0f} bps" if v else "N/A"),
        ]:
            row = va_table.add_row().cells
            row[0].text = label
            for i, scen in enumerate(["bear", "base", "bull"]):
                val = va_results.get(scen, {}).get(key)
                row[i + 1].text = fmt_fn(val) if val is not None else "N/A"

        # VA max offer
        if va_max_offer and va_max_offer.get("max_price"):
            doc.add_paragraph()
            doc.add_paragraph(
                f"Value-Add Maximum Offer Price (for "
                f"{va_max_offer.get('target_irr', cfg.SOLVER_TARGET_IRR):.0%} "
                f"IRR): "
                f"{_fmt_currency(va_max_offer['max_price'])} "
                f"(implied entry cap: {_fmt_pct(va_max_offer.get('implied_entry_cap'))})"
            ).bold = True


def _add_section_8(doc, risk_analysis):
    doc.add_heading("8. Risk Analysis", level=1)

    # Why this deal could fail
    why_fail = risk_analysis.get("why_deal_could_fail", [])
    if why_fail:
        doc.add_heading("Why This Deal Could Fail", level=2)
        for r in why_fail:
            doc.add_paragraph(
                f"{r.get('risk', 'Unknown')}: {r.get('description', '')}",
                style="List Bullet",
            )

    # Full risk register
    risks = risk_analysis.get("risks", [])
    if risks:
        doc.add_heading("Risk Register", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdrs = table.rows[0].cells
        hdrs[0].text = "Risk"
        hdrs[1].text = "Category"
        hdrs[2].text = "Severity"
        hdrs[3].text = "Mitigation"

        for r in risks:
            row = table.add_row().cells
            row[0].text = r.get("risk") or ""
            row[1].text = r.get("category") or ""
            row[2].text = r.get("severity") or ""
            row[3].text = r.get("mitigation") or ""

    doc.add_paragraph(f"\nOverall Risk Rating: {risk_analysis.get('risk_rating', 'TBD')}")


def _add_section_9(doc):
    doc.add_heading("9. Due Diligence Items", level=1)
    items = [
        "Obtain and review actual T-12 P&L (not broker pro forma)",
        "Verify 3-mile population and demographics via census data",
        "Confirm new supply pipeline with local planning/permitting records",
        "Conduct physical property inspection and condition assessment",
        "Obtain rent roll with move-in dates and rate history",
        "Reconcile economic vs physical occupancy from the rent roll "
        "(concessions, delinquency, below-street in-place rates)",
        "Verify competitor street-rate trend over trailing 6-12 months — "
        "confirm the in-place-to-market gap is closing from below (market "
        "rising), not from above (street rates falling)",
        "Verify property tax assessment and potential reassessment at sale price",
        "Review competitor rent survey (independent of CIM data)",
        "Confirm insurance quotes for the specific property",
        "Review lease/rental agreement terms",
        "Environmental Phase I assessment",
        "Title and survey review",
        "Verify zoning and entitlements",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _add_section_10(doc, gate_results, scenario_results, max_offer, risk_analysis, cim_data):
    doc.add_heading("10. Recommendation", level=1)

    # Determine recommendation
    failed = [g for g in gate_results if g["result"] == "FAIL"]
    tbd = [g for g in gate_results if g["result"] == "TBD"]
    base_irr = scenario_results.get("base", {}).get("irr") if scenario_results else None

    if failed:
        rec = "DECLINE"
        rationale = "One or more screening gates have failed:"
    elif tbd:
        rec = "PURSUE CONTINGENT ON"
        rationale = "Screening gates passed but the following require verification:"
    elif base_irr and base_irr >= cfg.GATES["min_irr_5yr"]:
        rec = "PURSUE"
        rationale = (f"All screening gates passed and base case returns meet "
                     f"the {cfg.GATES['min_irr_5yr']:.0%} IRR target.")
    else:
        rec = "PURSUE CONTINGENT ON"
        rationale = (f"Screening gates passed but base case IRR is below "
                     f"{cfg.GATES['min_irr_5yr']:.0%} target.")

    p = doc.add_paragraph()
    run = p.add_run(f"RECOMMENDATION: {rec}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = (
        RGBColor(0, 128, 0) if rec == "PURSUE" else
        RGBColor(204, 102, 0) if "CONTINGENT" in rec else
        RGBColor(204, 0, 0)
    )

    doc.add_paragraph(rationale)

    if failed:
        for g in failed:
            doc.add_paragraph(
                f"Gate {g['gate']} ({g['name']}): {g.get('note', '')}",
                style="List Bullet",
            )

    if tbd:
        for g in tbd:
            doc.add_paragraph(
                f"Gate {g['gate']} ({g['name']}): {g.get('note', '')}",
                style="List Bullet",
            )

    # Pricing guidance
    if max_offer and max_offer.get("max_price") and cim_data.asking_price:
        mp = max_offer["max_price"]
        asking = cim_data.asking_price
        discount = (asking - mp) / asking if asking else 0
        doc.add_paragraph()
        doc.add_paragraph(
            f"Maximum Offer: {_fmt_currency(mp)} "
            f"({discount:.1%} discount to asking price of {_fmt_currency(asking)})"
        )

    doc.add_paragraph()
    doc.add_paragraph(
        "Note: This analysis is based on CIM-provided data supplemented by "
        "benchmark assumptions. All figures should be verified during due diligence."
    ).italic = True


# ── LP-Facing Investor Summary (item G) ─────────────────────────────
#
# A two-page document for a sophisticated family office. The operator's
# framing sets the structure: a clear target return, a description of the
# PLAN to achieve it, and risks acknowledged WITH mitigants.
#
# It is a SECOND RENDERING, never a second computation. Every figure is
# read off the same result dicts `generate_memo` receives. The only
# arithmetic permitted here is differences and ratios between figures
# that are already published elsewhere — the leverage delta, the
# yield-on-cost spread. Never a re-derivation of a modelled quantity;
# that is the second-source-of-truth defect the rule exists to prevent.
#
# **The two-page mechanism CONSTRAINS Word rather than predicting it.**
# Geometry is pinned, styles use EXACTLY line spacing, and every block is
# a table row with an EXACTLY height, so Word cannot reflow. What remains
# is a content budget, enforced by `output.page_budget`, which raises
# `InvestorSummaryOverflow` rather than silently shrinking anything. Read
# that module's docstring for the conditions under which the guarantee is
# wrong — it is a content budget, not a page count.
#
# **Write and measure are one call, never two.** `_is_para` and
# `_is_table` render the text AND charge the budget for that same
# string. An earlier draft measured in one place and rendered in
# another, which is a divergence waiting to happen: the budget would
# pass on text the document did not contain.
#
# **Distribution is gated, the build is not.** A document aimed at
# prospective investors edges toward securities marketing, which sits
# behind the operator's General Counsel gate. `_SUMMARY_LEGEND` states
# that on the page itself; do not remove it, and route wording past GC
# before this leaves the firm.

_IS_BODY_PT = 9
_IS_HEAD_PT = 12
_IS_MICRO_PT = 7.5

_IS_MAX_THESIS = 3
_IS_MAX_RISKS = 3
_IS_MAX_PLAN_ROWS = 3

# Field caps, each asserted directly by the tests.
_IS_MAX_NAME_CHARS = 60
_IS_MAX_MITIGANT_CHARS = 180
_IS_MAX_THESIS_CHARS = 180
_IS_MAX_PLAN_CHARS = 120

# Risk severities, most severe first. `identify_risks` emits title case
# and pre-sorts; this orders the same vocabulary rather than inventing a
# second one.
_IS_SEVERITY_ORDER = {"High": 0, "Medium": 1, "Low": 2}

# Typographic punctuation this document and broker PDFs actually use,
# folded to ASCII so `page_budget` can measure it. Accented Latin is
# handled by NFKD below. Anything still non-ASCII after both — CJK,
# emoji — deliberately reaches `text_width_pt` and RAISES: the width
# table has no entry for it, and measuring it at Latin width is how a
# two-page guarantee silently becomes three pages.
_IS_PUNCT = {"’": "'", "‘": "'", "“": '"', "”": '"',
             "—": "-", "–": "-", "…": "...", "·": "-",
             "→": "->", "×": "x", " ": " ", "•": "*"}

#: Shown while `config.INVESTOR_SUMMARY_GC_CLEARED` is False, on the
#: document's own first line. It is NOT a substitute for
#: `_SUMMARY_LEGEND`, which is permanent and states what the document is;
#: this states that nobody has yet cleared it to leave the firm.
#:
#: On the page rather than only on screen because the failure this guards
#: is a file already detached from the app — attached to an email, in a
#: data room. A caveat that lives beside the download button is invisible
#: the moment the .docx moves, which is precisely when it matters.
#:
#: ASCII only, deliberately. `_is_para` folds typographic punctuation to
#: ASCII so `page_budget` can measure it, so an em dash here would make
#: the constant differ from the rendered text — and a test asserting
#: `_GC_PENDING_NOTICE in body` would fail while the notice was on the
#: page. `_SUMMARY_LEGEND` is ASCII for the same reason.
_GC_PENDING_NOTICE = (
    "INTERNAL DRAFT - NOT CLEARED FOR EXTERNAL DISTRIBUTION. This summary "
    "has not been reviewed by counsel. Do not send it to any prospective "
    "investor or third party."
)

_SUMMARY_LEGEND = (
    "Prepared by CIM Analyst from the seller's Confidential Information "
    "Memorandum supplemented by benchmark assumptions. Figures are "
    "underwriting estimates, not results, and are subject to due "
    "diligence. Past or projected performance is not a guarantee of "
    "future results. This document is for internal and prospect "
    "discussion only. It is not an offer to sell or a solicitation of an "
    "offer to buy any security, and it is not investment advice."
)


def generate_investor_summary(property_name: str, cim_data,
                              market_analysis: dict, physical_analysis: dict,
                              scenario_results: dict, risk_analysis: dict,
                              rent_analysis: dict = None,
                              value_add: dict = None,
                              va_results: dict = None,
                              gate_results: list = None,
                              gate_summary: dict = None,
                              check_summary: dict = None,
                              sources_uses: dict = None,
                              levered: dict = None, debt: dict = None,
                              assumption_fill_log: list = None,
                              thesis: list = None,
                              output_dir: str = ".") -> str:
    """Generate the 2-page LP-facing investor summary .docx.

    Args:
        thesis: operator override for the three thesis bullets. None
            derives them from the priority ladder in `_derive_thesis`,
            which is the normal path — there is no thesis field on any
            result object, and adding a DB column for prose the model can
            already justify would be a second source of truth.

    Raises:
        InvestorSummaryOverflow: content exceeds a page's budget. Loud on
            purpose — a silently shrunk document is one nobody notices is
            wrong, and this one goes to investors.

    Returns: path to the generated file.
    """
    doc, page1, page2 = _is_build(
        property_name, cim_data, market_analysis, physical_analysis,
        scenario_results, risk_analysis, rent_analysis, value_add,
        va_results, gate_results, gate_summary, check_summary,
        sources_uses, levered, debt, assumption_fill_log, thesis)

    page1.check()
    # The floor guards the opposite defect: a page 2 the truncation
    # ladder emptied out is as wrong as one that overflows, and a
    # one-sided assert never catches it. It applies ONLY to a deal that
    # had that content in the first place — a thin early-look CIM with
    # no scenarios and no risks is legitimately short, and failing it
    # would make the degraded path impossible rather than safe.
    complete = bool(scenario_results and (risk_analysis or {}).get("risks"))
    page2.check(floor_pt=PAGE_MIN_PT if complete else None)

    safe_name = safe_filename(property_name or "Unknown_Property")
    filepath = os.path.join(output_dir, f"Investor_Summary_{safe_name}.docx")
    doc.save(filepath)
    return filepath


def _is_build(property_name, cim_data, market_analysis, physical_analysis,
              scenario_results, risk_analysis, rent_analysis, value_add,
              va_results, gate_results, gate_summary, check_summary,
              sources_uses, levered, debt, assumption_fill_log, thesis):
    """Compose the document and return it with both page budgets.

    Split out so the tests can assert the budget NUMBERS rather than only
    that composing did not raise. A test that just checks "no exception"
    cannot tell a page that fits comfortably from one a byte away from
    overflowing.
    """
    doc = Document()
    _is_pin_geometry(doc)

    profile = (physical_analysis or {}).get("property_profile") or {}
    base = (scenario_results or {}).get("base") or {}
    lev_base = (levered or {}).get("base") or {}
    # CLAUDE.md decision 7: no LP net figure without its assumption stamp.
    # The figures and the stamp were gated independently, so a payload with
    # a levered scenario but no stamp printed an LP net IRR bare. Coupling
    # them in ONE predicate is how memo section 6 does it (`if not base:
    # return`), and it makes the invariant structural rather than a
    # property of whatever the caller happened to pass.
    levered = levered if (lev_base.get("assumption_stamp")) else None
    lev_base = lev_base if levered else {}

    page1 = PageBudget("Page 1")
    page2 = PageBudget("Page 2")

    _is_gc_notice(doc, page1)
    _is_portfolio_notice(doc, page1, cim_data)
    _is_header(doc, page1, cim_data, profile, property_name)
    _is_target_return(doc, page1, scenario_results, levered)
    _is_assumption_stamp(doc, page1, lev_base)
    _is_thesis(doc, page1, thesis, cim_data, physical_analysis,
               rent_analysis, value_add, market_analysis)
    _is_key_metrics(doc, page1, cim_data, base, physical_analysis,
                    sources_uses, debt)
    _is_sources_uses(doc, page1, sources_uses, lev_base)

    _is_plan_to_achieve(doc, page2, value_add, va_results, base)
    _is_scenarios(doc, page2, scenario_results)
    _is_risks(doc, page2, risk_analysis)
    _is_market(doc, page2, market_analysis, gate_results)
    _is_footer(doc, page2, gate_summary, check_summary,
               assumption_fill_log)

    return doc, page1, page2


# ── Geometry, styles, and the write-and-measure helpers ──────────────

def _is_pin_geometry(doc):
    """Pin the page and define the three styles the budget assumes.

    EXACTLY line spacing and `widow_control = False` are not cosmetic:
    Word's widow/orphan control is the classic invisible line-pusher,
    and "at least" spacing lets one tall glyph grow a row. Either would
    break a budget computed from nominal metrics.
    """
    from docx.enum.text import WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE

    section = doc.sections[0]
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.left_margin = section.right_margin = Inches(MARGIN_X_IN)
    section.top_margin = section.bottom_margin = Inches(MARGIN_Y_IN)

    for name, size in (("LPBody", _IS_BODY_PT), ("LPHead", _IS_HEAD_PT),
                       ("LPMicro", _IS_MICRO_PT)):
        try:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        fmt = style.paragraph_format
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(size * 1.15)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(2)
        fmt.widow_control = False


def _ascii(text) -> str:
    """Fold typographic punctuation and accents to ASCII.

    Deliberately does NOT strip what it cannot fold. Dropping a CJK
    glyph here would let an unmeasurable string through silently; leaving
    it makes `page_budget.text_width_pt` raise, which is the contract.
    """
    import unicodedata

    s = str(text if text is not None else "")
    for bad, good in _IS_PUNCT.items():
        s = s.replace(bad, good)
    s = unicodedata.normalize("NFKD", s)
    return "".join(c for c in s if not unicodedata.combining(c))


_IS_STYLE_PT = {"LPBody": _IS_BODY_PT, "LPHead": _IS_HEAD_PT,
                "LPMicro": _IS_MICRO_PT}


def _is_para(doc, budget, label, text, style="LPBody", bold=False,
             italic=False, color=None):
    """Render a paragraph AND charge the budget for the same string."""
    text = _ascii(text)
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if budget is not None:
        budget.add(label, paragraph_pt(text, _IS_STYLE_PT[style], bold=bold))
    return para


def _is_table(doc, budget, label, rows, widths_in, header=False):
    """Render a fixed-layout table AND charge the budget for it.

    `autofit=False` plus explicit cell widths is what stops Word
    re-deciding column widths, which would invalidate every wrap the
    budget computed. Row heights are EXACTLY for the same reason.
    """
    from docx.enum.table import WD_ROW_HEIGHT_RULE

    rows = [[_ascii(c) for c in row] for row in rows]
    table = doc.add_table(rows=0, cols=len(widths_in))
    table.autofit = False
    if header:
        table.style = "Table Grid"
    for cells in rows:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = Pt(_IS_BODY_PT * 1.15)
        for i, value in enumerate(cells):
            cell = row.cells[i]
            cell.width = Inches(widths_in[i])
            para = cell.paragraphs[0]
            para.style = doc.styles["LPBody"]
            para.add_run(value)
    if budget is not None:
        budget.add(label, table_pt(rows, _IS_BODY_PT,
                                   [w * 72 for w in widths_in]))
    return table


def _truncate(text, limit: int) -> str:
    text = str(text if text is not None else "")
    return text if len(text) <= limit else text[:limit - 3].rstrip() + "..."


# ── Page 1 — "What you make" ─────────────────────────────────────────

def _is_gc_notice(doc, budget):
    """The distribution gate, on the first line of page 1 (item G).

    Read at CALL time via `cfg.` rather than imported by name, so
    flipping the flag takes effect without a reimport — the scalar rule
    recorded in CLAUDE.md's config-reads note.

    Charged to the page budget like every other block, because it is
    real content: a notice that renders without being measured is how a
    document that fits in the tests overflows in Word.
    """
    if cfg.INVESTOR_SUMMARY_GC_CLEARED:
        return
    _is_para(doc, budget, "gc/pending", _GC_PENDING_NOTICE,
             style="LPMicro", bold=True)


def _is_portfolio_notice(doc, budget, cim_data):
    """The portfolio caveat on page 1 of the LP document (no-op otherwise).

    This is the one surface that leaves the firm, so it is the one where
    the caveat binds hardest (the decision-7 posture). One fixed sentence,
    charged to the budget like every block; the detection EVIDENCE is
    deliberately not enumerated — that is analyst diligence detail, and
    the memo carries it.
    """
    if not getattr(cim_data, "portfolio_signal", None):
        return
    _is_para(doc, budget, "portfolio/notice",
             "This document may describe more than one property; figures "
             "may combine portfolio-level and per-property values. Refer "
             "to the full investment memo before relying on any figure.",
             style="LPMicro", bold=True)


def _is_header(doc, budget, cim_data, profile, property_name):
    name = _truncate(property_name or profile.get("property_name")
                     or "Property", _IS_MAX_NAME_CHARS)
    _is_para(doc, budget, "header/name", name, style="LPHead", bold=True)

    bits = [profile.get("city_state"),
            f"Built {profile.get('year_built')}" if profile.get("year_built") else None,
            _fmt_number(profile.get("nrsf"), suffix=" NRSF") if profile.get("nrsf") else None,
            f"{profile.get('total_units')} units" if profile.get("total_units") else None,
            f"{profile.get('cc_pct'):.0%} climate-controlled" if profile.get("cc_pct") else None]
    line = " - ".join(b for b in bits if b)
    if line:
        _is_para(doc, budget, "header/profile", line)

    # Physical and economic occupancy as a PAIR, always. A broker quoting
    # one occupancy number is almost always quoting physical, and the
    # spread between the two is this fund's entire value-add thesis.
    phys, econ = profile.get("physical_occupancy"), profile.get("economic_occupancy")
    if phys is not None or econ is not None:
        occ = f"Occupancy: {_fmt_pct(phys)} physical / {_fmt_pct(econ)} economic"
        if phys is not None and econ is not None:
            occ += f"  ({(phys - econ) * 100:.0f}-point spread)"
        _is_para(doc, budget, "header/occupancy", occ)


def _is_target_return(doc, budget, scenario_results, levered):
    """Unlevered and LP net, side by side at equal weight, with the
    leverage effect between them.

    Equal weight is the operator's call and it is the honest
    presentation: on config defaults the loan constant frequently sits
    above yield on cost, so LP net can land BELOW unlevered. Leading
    with the levered number alone would bury that, and a sophisticated
    allocator who spots it stops trusting the whole document.
    """
    if not scenario_results:
        return
    _is_para(doc, budget, "target/head", "TARGET RETURN", style="LPHead",
             bold=True)

    base = scenario_results.get("base") or {}
    bear = scenario_results.get("bear") or {}
    bull = scenario_results.get("bull") or {}
    lev = levered or {}

    def band(lo, hi, fmt):
        return "-" if lo is None or hi is None else f"{fmt(lo)} to {fmt(hi)}"

    rows = [["", "Unlevered (property)", "LP Net (after fees & promote)"],
            ["IRR", _fmt_pct(base.get("irr")),
             _fmt_pct((lev.get("base") or {}).get("lp_net_irr"))],
            ["MOIC", _fmt_x(base.get("moic")),
             _fmt_x((lev.get("base") or {}).get("lp_moic"))],
            ["Bear to Bull", band(bear.get("irr"), bull.get("irr"), _fmt_pct),
             band((lev.get("bear") or {}).get("lp_net_irr"),
                  (lev.get("bull") or {}).get("lp_net_irr"), _fmt_pct)]]
    _is_table(doc, budget, "target/table", rows, [1.3, 3.0, 3.0], header=True)

    # Leverage effect — a DIFFERENCE between two published figures, which
    # is the only arithmetic this document is allowed to do.
    unlev, lp = base.get("irr"), (lev.get("base") or {}).get("lp_net_irr")
    if unlev is not None and lp is not None:
        bps = (lp - unlev) * 10_000
        verdict = "accretive" if bps >= 0 else "DILUTIVE"
        _is_para(doc, budget, "target/leverage",
                 f"Leverage effect: {bps:+,.0f} bps ({verdict}). Debt is "
                 f"{verdict} to the LP at the modelled terms.",
                 bold=bps < 0)


def _is_assumption_stamp(doc, budget, lev_base):
    """CLAUDE.md key design decision 7 — no LP net IRR without its stamp.

    This is the only surface that leaves the firm, so the rule binds
    hardest here: the reader has no memo section 6 and no Returns tab to
    cross-check which conventions produced the number.
    """
    stamp = (lev_base or {}).get("assumption_stamp") or []
    labels = " - ".join(r.get("label", "") for r in stamp if r.get("label"))
    if not labels:
        # Unreachable from `_is_build`, which nulls `levered` outright when
        # the stamp is missing so no levered FIGURE prints either. Kept as
        # a guard for direct callers.
        return
    # "Proposed, subject to the final partnership agreement" is the right
    # sentence for a convention nobody has read the LPA on, and the WRONG
    # one for a convention the operator confirmed against it — this
    # document goes to investors, so overstating and understating are both
    # costly. Rows carry their own status; the caveat follows the rows
    # that still need it rather than blanketing all five.
    # Counted in three, not two. A `moot` row was NOT confirmed — nobody
    # read the LPA on it; it simply stopped being able to move the number.
    # Folding it into the confirmed count would claim more of the document
    # had been read than has been, which is the overstatement this whole
    # sentence exists to avoid.
    confirmed = sum(1 for r in stamp if r.get("status") == "confirmed")
    moot = sum(1 for r in stamp if r.get("status") == "moot")
    total = len(stamp)
    if confirmed + moot == 0:
        caveat = ("These are proposed terms, subject to the final "
                  "partnership agreement.")
    else:
        parts = []
        if confirmed:
            parts.append(f"{confirmed} of {total} confirmed against the "
                         f"executed partnership agreement")
        if moot:
            parts.append(f"{moot} made moot by it")
        settled = ", ".join(parts)
        caveat = (f"{settled}."
                  if confirmed + moot == total
                  else f"{settled}; the rest are proposed terms, subject to it.")
        caveat = caveat[0].upper() + caveat[1:]
    _is_para(doc, budget, "stamp",
             f"LP net returns are computed under: {labels}. {caveat}",
             style="LPMicro", italic=True, color=RGBColor(0x55, 0x55, 0x55))


def _derive_thesis(cim_data, physical_analysis, rent_analysis, value_add,
                   market_analysis) -> list:
    """Top three, in fixed priority order.

    Each candidate is GATED on a value the pipeline already produced and
    restates that value, so the thesis cannot claim something the model
    did not find. Priority is fixed rather than scored: a ranking
    function would be a judgment this document is not allowed to make.
    """
    out = []

    pvr = (physical_analysis or {}).get("price_vs_replacement") or {}
    discount = pvr.get("discount_to_replacement")
    if discount is not None and discount > 0:
        out.append(f"Basis {discount:.0%} below replacement cost "
                   f"({_fmt_currency(pvr.get('asking_price'))} against "
                   f"{_fmt_currency(pvr.get('replacement_cost'))} to build) - "
                   f"a structural entry advantage a competing buyer cannot "
                   f"underwrite away.")

    # `is not None`, not truthiness: a 0% economic occupancy is a
    # property collecting nothing, which is the LOUDEST version of the
    # profile this bullet exists to surface.
    phys = getattr(cim_data, "physical_occupancy", None)
    econ = getattr(cim_data, "economic_occupancy", None)
    flag = cfg.GATES.get("econ_phys_spread_flag")
    if (phys is not None and econ is not None and flag
            and (phys - econ) >= flag):
        out.append(f"Economic occupancy trails physical by "
                   f"{(phys - econ) * 100:.0f} points ({_fmt_pct(econ)} "
                   f"against {_fmt_pct(phys)}) - the asset is full and "
                   f"under-collecting, a management problem rather than a "
                   f"demand problem.")

    gap = (rent_analysis or {}).get("rent_gap_pct")
    if gap is not None and gap > 0:
        out.append(f"In-place rents sit {gap:.0%} below market, closable "
                   f"through the existing tenant base rather than new "
                   f"lease-up.")

    revenue_ops = (value_add or {}).get("revenue_opportunities") or []
    if revenue_ops:
        top = revenue_ops[0]
        out.append(f"{top.get('category')}: {top.get('description')}")

    demos = (market_analysis or {}).get("demographics") or {}
    msa = (market_analysis or {}).get("msa_info") or {}
    pop = demos.get("population_3mi")
    if pop:
        line = f"{pop:,.0f} people within three miles"
        if msa.get("is_top_50"):
            line += f"; {msa.get('msa_name')} is a top-50 MSA"
        out.append(line + ".")

    return out[:_IS_MAX_THESIS]


def _is_thesis(doc, budget, thesis, cim_data, physical_analysis,
               rent_analysis, value_add, market_analysis):
    bullets = thesis if thesis is not None else _derive_thesis(
        cim_data, physical_analysis, rent_analysis, value_add, market_analysis)
    bullets = [_truncate(b, _IS_MAX_THESIS_CHARS)
               for b in (bullets or [])][:_IS_MAX_THESIS]
    if not bullets:
        return
    _is_para(doc, budget, "thesis/head", "Investment Thesis", style="LPHead",
             bold=True)
    for b in bullets:
        _is_para(doc, budget, "thesis/bullet", f"* {b}")


def _is_key_metrics(doc, budget, cim_data, base, physical_analysis,
                    sources_uses, debt):
    pvr = (physical_analysis or {}).get("price_vs_replacement") or {}
    rows = [["Asking Price", _fmt_currency(getattr(cim_data, "asking_price", None))],
            ["Price / SF", _fmt_currency(getattr(cim_data, "price_per_sf", None))]]

    discount = pvr.get("discount_to_replacement")
    if discount is not None:
        rows.append(["Discount to Replacement", f"{discount:.1%}"])

    rows.append(["Entry Cap", _fmt_cap(base.get("entry_cap"))])
    exit_label = _fmt_cap(base.get("exit_cap"))
    if base.get("exit_cap_coerced"):
        exit_label += " (floored at entry)"
    rows.append(["Exit Cap (Base)", exit_label])
    rows.append(["Yr-1 Yield on Cost", _fmt_pct(base.get("yield_on_cost"))])

    # Total basis is the UNLEVERED definition and excludes financing costs
    # (CLAUDE.md decision 3). `sources_uses["total_uses"]` is documented as
    # `total_basis + financing_costs`, so printing it under this label put a
    # financing-inflated number on the one document built to keep the
    # unlevered lens financing-free. `output/excel_writer.py` labels the
    # same row off `scen["total_basis"]`; this reads the same field.
    if base.get("total_basis") is not None:
        rows.append(["Total Basis", _fmt_currency(base.get("total_basis"))])
    if sources_uses:
        rows.append(["Equity Required",
                     _fmt_currency(sources_uses.get("total_equity"))])
    if debt and debt.get("loan"):
        from model.debt import binding_constraint_label
        rows.append(["Senior Debt", f"{_fmt_currency(debt.get('loan'))} "
                                    f"({binding_constraint_label(debt)})"])

    # Gaps are SHOWN, not dropped (operator's call): an LP cannot tell a
    # metric that is missing from one that was never part of the analysis
    # if the row simply vanishes.
    _is_para(doc, budget, "metrics/head", "Key Metrics", style="LPHead",
             bold=True)
    _is_table(doc, budget, "metrics/table", rows, [2.6, 4.7])


def _is_sources_uses(doc, budget, sources_uses, lev_base):
    """Sources & Uses, plus the fee and promote load.

    Fee transparency is volunteered rather than waited for. This audience
    asks what was deducted to reach an LP net number within ten minutes;
    printing it converts an interrogation into a checked box.
    """
    if not sources_uses:
        return
    _is_para(doc, budget, "su/head", "Capital Stack", style="LPHead", bold=True)

    total = sources_uses.get("total_uses") or 0

    def share(amount):
        return f"{amount / total:.0%}" if total and amount is not None else "-"

    rows = [["Total Uses", _fmt_currency(total), ""],
            ["Senior Debt", _fmt_currency(sources_uses.get("senior_debt")),
             share(sources_uses.get("senior_debt"))],
            ["Total Equity", _fmt_currency(sources_uses.get("total_equity")),
             share(sources_uses.get("total_equity"))]]
    coinvest = sources_uses.get("gp_coinvest_pct")
    if coinvest is not None:
        rows.append(["  of which GP co-invest",
                     _fmt_currency(sources_uses.get("gp_equity")),
                     f"{coinvest:.0%} of equity"])
    if lev_base:
        fee_pct = lev_base.get("am_fee_pct")
        rows.append(["Asset-management fee (hold total)",
                     _fmt_currency(lev_base.get("am_fee_total")),
                     f"{fee_pct:.2%}/yr" if fee_pct is not None else ""])
        rows.append(["GP promote (hold total)",
                     _fmt_currency(lev_base.get("gp_promote")), ""])

    _is_table(doc, budget, "su/table", rows, [3.3, 2.2, 1.8])


# ── Page 2 — "How we get there, and what breaks it" ──────────────────

def _is_plan_to_achieve(doc, budget, value_add, va_results, base):
    """The centre of the operator's framing: how the return is produced.

    Three sources, in order of how concrete they are. A value-add deal
    gets its named opportunities and, when the monthly engine ran, the
    quantified bridge. A stabilized deal has neither, so it falls back to
    where the return comes from arithmetically — spreads between figures
    already published on page 1.
    """
    head = _is_para(doc, budget, "plan/head", "Plan to Achieve the Return",
                    style="LPHead", bold=True)
    # `page_break_before` rather than `add_page_break()`, which injects a
    # stray empty paragraph — a line of height the budget never charged.
    head.paragraph_format.page_break_before = True

    ops = list((value_add or {}).get("revenue_opportunities") or [])
    ops += list((value_add or {}).get("expense_opportunities") or [])
    uplift = (value_add or {}).get("estimated_noi_uplift")

    if ops:
        if uplift:
            _is_para(doc, budget, "plan/uplift",
                     f"Identified NOI uplift: {_fmt_currency(uplift)} per "
                     f"year at stabilization.", bold=True)
        rows = [["Initiative", "Impact / yr", "Timeline"]]
        for op in ops[:_IS_MAX_PLAN_ROWS]:
            rows.append([
                _truncate(f"{op.get('category')}: {op.get('description')}",
                          _IS_MAX_PLAN_CHARS),
                _fmt_currency(op.get("est_annual_impact")),
                op.get("timeline") or "-"])
        _is_table(doc, budget, "plan/table", rows, [4.3, 1.5, 1.5], header=True)

    va = (va_results or {}).get("base") or {}
    if va:
        bridge = []
        cur, tgt = va.get("current_occupancy"), va.get("target_occupancy")
        months = va.get("months_to_stabilize")
        if cur is not None and tgt is not None:
            bridge.append(f"occupancy {cur:.0%} -> {tgt:.0%}"
                          + (f" over {months} months" if months else ""))
        ip, tp = va.get("in_place_rent_psf"), va.get("target_rent_psf")
        mp = va.get("market_rent_psf")
        # The LP-facing document is the one that leaves the firm, so the
        # rule binds hardest here: with no market-rent data there is no
        # rent leg to the bridge and no "market" to quote — the copied
        # in-place rent is not one (item T Category 4).
        if va.get("rent_ramp_excluded"):
            bridge.append("no rent ramp (no market-rent data)")
        elif ip and tp:
            leg = f"rent ${ip:.2f} -> ${tp:.2f}/SF"
            if mp:
                leg += f" against a ${mp:.2f} market"
            bridge.append(leg)
        if bridge:
            _is_para(doc, budget, "plan/bridge",
                     "Underwritten bridge: " + "; ".join(bridge) + ".")
    elif not ops:
        # Stabilized deal: no value-add narrative exists, so say where the
        # return actually comes from. All three are differences between
        # figures printed on page 1.
        parts = []
        yoc, entry = base.get("yield_on_cost"), base.get("entry_cap")
        if yoc is not None and entry is not None:
            parts.append(f"a {(yoc - entry) * 10_000:+,.0f} bp spread of "
                         f"Year-1 yield on cost over the entry cap")
        noi = base.get("noi_projection") or []
        if len(noi) >= 2 and noi[0]:
            parts.append(f"NOI growth of {(noi[-1] / noi[0] - 1):.0%} across "
                         f"the hold")
        exit_cap = base.get("exit_cap")
        if exit_cap is not None and entry is not None:
            parts.append(f"an exit underwritten "
                         f"{(exit_cap - entry) * 10_000:+,.0f} bp "
                         f"{'wider' if exit_cap >= entry else 'tighter'} "
                         f"than entry")
        if parts:
            _is_para(doc, budget, "plan/sources",
                     "This is a stabilized asset; the return comes from "
                     + ", ".join(parts) + ".")


def _is_scenarios(doc, budget, scenario_results):
    if not scenario_results:
        return
    hold = _hold_years(scenario_results)
    _is_para(doc, budget, "scen/head", f"Scenario Returns ({hold}-Year Hold)",
             style="LPHead", bold=True)
    rows = [["", "Bear", "Base", "Bull"]]
    for label, key, fmt in (("Unlevered IRR", "irr", _fmt_pct),
                            ("MOIC", "moic", _fmt_x),
                            ("Yr-1 Yield on Cost", "yield_on_cost", _fmt_pct),
                            ("Exit Cap", "exit_cap", _fmt_cap)):
        rows.append([label] + [fmt((scenario_results.get(s) or {}).get(key))
                               for s in ("bear", "base", "bull")])
    _is_table(doc, budget, "scen/table", rows, [2.5, 1.6, 1.6, 1.6],
              header=True)


def _is_risks(doc, budget, risk_analysis):
    """Top three by severity, each WITH its mitigant.

    Mitigants are never dropped by the truncation ladder. A risk printed
    without one reads as an unanswered objection, which is worse for this
    audience than not raising the risk at all.
    """
    risks = list((risk_analysis or {}).get("risks") or [])
    if not risks:
        return
    risks.sort(key=lambda r: _IS_SEVERITY_ORDER.get(r.get("severity"), 99))

    _is_para(doc, budget, "risk/head", "Principal Risks & Mitigants",
             style="LPHead", bold=True)
    for r in risks[:_IS_MAX_RISKS]:
        _is_para(doc, budget, "risk/item",
                 f"{r.get('severity', '')} - {r.get('risk', '')}", bold=True)
        mitigant = _truncate(r.get("mitigation") or "Under diligence.",
                             _IS_MAX_MITIGANT_CHARS)
        _is_para(doc, budget, "risk/mitigant", f"Mitigant: {mitigant}",
                 italic=True)

    failure = ((risk_analysis or {}).get("why_deal_could_fail") or [None])[0]
    if failure:
        _is_para(doc, budget, "risk/failure",
                 f"Primary failure mode: {failure}")


def _is_market(doc, budget, market_analysis, gate_results):
    demos = (market_analysis or {}).get("demographics") or {}
    msa = (market_analysis or {}).get("msa_info") or {}
    by_gate = {g.get("gate"): g for g in (gate_results or [])}

    rows = []
    if demos.get("population_3mi"):
        rows.append(["Population (3-mile)", f"{demos['population_3mi']:,.0f}"])
    if demos.get("median_hhi_3mi"):
        rows.append(["Median HHI (3-mile)",
                     _fmt_currency(demos["median_hhi_3mi"])])
    if msa.get("msa_name"):
        # `is_top_50` is a bool and there is no MSA rank number anywhere
        # in this codebase. Never print one.
        rows.append(["Market", msa["msa_name"]
                     + (" - top-50 MSA" if msa.get("is_top_50") else "")])
    # SF/capita is TBD whenever competitive supply is unentered, which is
    # the common case. Conditional, not required.
    sf_gate = by_gate.get(5)
    if sf_gate and "TBD" not in str(sf_gate.get("actual", "")):
        rows.append(["Supply (3-mile)", str(sf_gate.get("actual"))])

    # Unlike Key Metrics, a market section with NO rows is omitted
    # entirely: a heading over three blank rows conveys nothing, where a
    # named metric reading N/A is itself information.
    if not rows:
        return
    _is_para(doc, budget, "mkt/head", "Market Snapshot", style="LPHead",
             bold=True)
    _is_table(doc, budget, "mkt/table", rows, [2.6, 4.7])


def _is_footer(doc, budget, gate_summary, check_summary,
               assumption_fill_log=None):
    bits = []
    rec = (gate_summary or {}).get("recommendation")
    if rec:
        bits.append(f"Screening result: {rec}")
    blocking = (check_summary or {}).get("blocking_failed")
    if blocking:
        bits.append(f"{blocking} blocking model check(s) unresolved")
    # The COUNT, never the table (item T Category 4). Every block on
    # these two pages is charged against a fixed page budget that RAISES
    # on overflow, and a fill log is variable-length by nature — one deal
    # with eleven filled inputs would cost an LP the document. The count
    # is one clause on a paragraph already budgeted, and the IC memo's
    # Appendix A carries the detail for anyone who asks.
    filled = len(assumption_fill_log or [])
    if filled:
        bits.append(f"{filled} assumption(s) filled from defaults")
    if bits:
        _is_para(doc, budget, "footer/status", " - ".join(bits), bold=True)

    _is_para(doc, budget, "footer/legend", _SUMMARY_LEGEND, style="LPMicro",
             italic=True, color=RGBColor(0x55, 0x55, 0x55))


# ── Formatting Helpers ──────────────────────────────────────────────

def _fmt_currency(val) -> str:
    if val is None:
        return "N/A"
    if abs(val) >= 1_000_000:
        return f"${val:,.0f}"
    return f"${val:,.0f}"


def _fmt_pct(val) -> str:
    if val is None:
        return "N/A"
    return f"{val:.1%}"


def _fmt_cap(val) -> str:
    """Cap rates to three places. The obsolescence drift is 5–10 bp/yr, so
    at `_fmt_pct`'s single decimal the printed build-up would not add up."""
    if val is None:
        return "N/A"
    return f"{val:.3%}"


def _exit_cap_derivation(scen: dict) -> str:
    """One sentence retracing a scenario's exit cap to its market anchor.

    Empty string when the scenario carries no derivation — a stored run
    from before the cap became derived must still render.
    """
    d = (scen or {}).get("exit_cap_detail") or {}
    if not d or d.get("market_cap") is None:
        return ""
    band = d.get("age_band") or "—"
    if d.get("age_band_known") is False:
        band += ", year built unknown"
    from analysis.valuation import describe_market_cap
    txt = (f"= {_fmt_cap(d['market_cap'])} market cap "
           f"({d.get('asset_class') or 'asset'}, {band}, "
           f"{describe_market_cap(d)})"
           f" {d.get('scenario_spread_bps', 0):+g} bp scenario spread"
           f" {d.get('drift_total_bps', 0):+g} bp obsolescence drift"
           f" ({d.get('drift_bps_per_year')} bp/yr × "
           f"{d.get('hold_years')} yrs)")
    if scen.get("exit_cap_coerced"):
        txt += (f" = {_fmt_cap(scen.get('requested_exit_cap'))}, then raised "
                f"to the entry cap to hold exit ≥ entry")
    return txt


def _fmt_number(val, suffix="") -> str:
    if val is None:
        return "N/A"
    return f"{val:,.0f}{suffix}"


def _fmt_x(val) -> str:
    """Multiples and coverage ratios. NOT `_fmt_number(v, "x")`: that
    rounds to whole numbers, which prints a 1.25x DSCR as "1x" and a
    1.39x MOIC as "1x" — the two decimals are the entire content."""
    if val is None:
        return "N/A"
    return f"{val:.2f}x"



