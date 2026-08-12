"""Item E3 — the levered seam (`model/levered.py`) and its surfaces.

E3a is the arithmetic; E3b is the assumptions inputs, the results-page
lens, the memo section and the Excel sheet, in a section of its own at
the bottom of this file. One module because it is one item, and because
a surface test that renders an oracle's real numbers is worth more than
one built on a fixture invented for it.

Every oracle here was derived from scratch in `Decimal`, importing
nothing from the repo, BEFORE `model/levered.py` existed — so these
assert an independently computed answer rather than a snapshot of the
implementation. The derivation cross-validates against E1: oracle A's
annual debt service (455,088.98) and payoff (5,616,658.65) reproduce
design-doc oracle 5 to the cent from an independent monthly roll-forward.

Three oracles, each exercising a path the others do not:

* **A** — debt yield binds, the deal never clears the pref, promote is
  zero, and leverage is DILUTIVE (7.1479% LP net against 7.3031%
  unlevered). A levered lens that always prints a bigger number than the
  unlevered screen has a bug; this is the regression that says so.
* **B** — LTV binds on a 24-month-IO loan, so the IO→amort transition is
  visible in the debt service, tier 1 clears and a promote is paid.
* **C** — levered cash flow is negative for four years: the funded
  reserve is drawn to zero, then capital is CALLED, and no period ever
  carries a negative distribution.

THE EXIT-NOI CONVENTION — now a config name, `EXIT_NOI_CONVENTION`
(decision 5, settled 2026-08-10), default "trailing": the projection
capitalizes the terminal hold year's OWN NOI. The design doc — and
therefore `tests/test_debt.py`'s oracle 5 — capitalizes year 6, which is
the FORWARD convention in a single-rate fixture, about 3% higher there.
Both remain deliberate; see CLAUDE.md's design-decisions block. These
oracles are computed on the TRAILING default, because that is the
projection the wiring actually consumes.
"""

import pytest

import config as cfg
from analysis.valuation import project_cash_flows
from model.debt import DebtTerms, build_debt_schedule
from model.levered import build_levered_returns
from model.returns_model import build_sources_uses
from model.waterfall import resolve_waterfall_terms

CENT = 0.005          # "to the cent" — half a cent of slack for float noise
BP = 1e-6             # IRRs to four decimal places, as a rate

COSTS = {"acquisition_closing_pct": 0.01, "disposition_cost_pct": 0.01}


def _params(growth):
    """A scenario parameter dict whose NOI grows at exactly `growth`.

    Revenue and expenses grow at the same rate, so the NOI margin holds
    and the series is `y1 * (1 + growth) ** k` — which is what the
    oracles were derived on. Both growth bands carry the same rate so a
    hold longer than three years does not change convention mid-stream.

    The exit cap is NOT in here. PR #31 made it an argument derived from
    a market anchor rather than a scenario parameter; these oracles pin
    it explicitly so the levered arithmetic is tested against a fixed
    exit, not against whatever the market table says this week.
    """
    return {"yr1_noi_bump": 0.0, "stabilized_occ": 0.88,
            "rev_cagr_yr1_3": growth, "rev_cagr_yr4_5": growth,
            "exp_growth": growth}


def _build(*, price, y1_noi, growth, exit_cap, terms, hold_years=5,
           reserve=0.0, gp_coinvest_pct=0.10, am_fee_pct=0.01):
    """Assemble one levered deal end to end, the way the engine does."""
    projection = project_cash_flows(
        y1_noi, price, 0.0, _params(growth),
        hold_years=hold_years, expense_ratio=0.40, costs=COSTS,
        reserve=reserve, exit_cap=exit_cap)
    debt = build_debt_schedule(price, projection["noi"][0], terms,
                               hold_years=hold_years)
    sources_uses = build_sources_uses(
        price=price, capex=0.0,
        acquisition_cost=projection["acquisition_cost"],
        reserve=reserve, financing_costs=debt["financing_costs"],
        senior_debt=debt["loan"], gp_coinvest_pct=gp_coinvest_pct)
    terms_w = resolve_waterfall_terms(
        capital_structure={"gp_coinvest_pct": gp_coinvest_pct})
    levered = build_levered_returns(
        projection, sources_uses=sources_uses, debt=debt,
        waterfall_terms=terms_w, am_fee_pct=am_fee_pct)
    return projection, debt, sources_uses, levered


# ── Oracle A: debt yield binds, the pref is never cleared ────────────

ORACLE_A = DebtTerms(rate=0.065, amort_years=30, io_months=0, max_ltv=0.65,
                     min_dscr=1.25, min_debt_yield=0.10, orig_fee_pct=0.01,
                     exit_fee_pct=0.0)


@pytest.fixture(scope="module")
def oracle_a():
    return _build(price=10_000_000, y1_noi=600_000, growth=0.03,
                  exit_cap=0.0625, terms=ORACLE_A)


def test_oracle_a_sizes_on_the_debt_yield(oracle_a):
    _, debt, _, _ = oracle_a
    assert debt["binding_constraint"] == "debt_yield"
    assert debt["loan"] == pytest.approx(6_000_000.00, abs=CENT)
    assert debt["financing_costs"] == pytest.approx(60_000.00, abs=CENT)
    assert debt["payoff_balance"] == pytest.approx(5_616_658.65, abs=CENT)
    assert debt["annual_debt_service"][0] == pytest.approx(455_088.98,
                                                           abs=CENT)


def test_oracle_a_stack_ties_to_the_basis_plus_financing(oracle_a):
    projection, debt, su, _ = oracle_a
    assert projection["total_basis"] == pytest.approx(10_100_000.00, abs=CENT)
    assert su["total_uses"] == pytest.approx(10_160_000.00, abs=CENT)
    assert su["total_uses"] == pytest.approx(
        projection["total_basis"] + debt["financing_costs"], abs=CENT)
    assert su["total_equity"] == pytest.approx(4_160_000.00, abs=CENT)


def test_oracle_a_levered_cash_flow(oracle_a):
    _, _, _, lev = oracle_a
    assert [r["am_fee"] for r in lev["years"]] == pytest.approx(
        [41_600.00] * 5, abs=CENT)
    assert [r["levered_cf"] for r in lev["years"]] == pytest.approx(
        [103_311.02, 121_311.02, 139_851.02, 158_947.22, 5_258_793.39],
        abs=CENT)
    assert lev["am_fee_total"] == pytest.approx(208_000.00, abs=CENT)


def test_oracle_a_never_clears_the_pref_so_no_promote(oracle_a):
    _, _, _, lev = oracle_a
    wf = lev["waterfall"]
    assert wf["tier1_current"] is False
    assert wf["unreturned_capital"] + wf["unpaid_pref"] == pytest.approx(
        225_455.71, abs=CENT)
    assert wf["gp"]["promote"] == pytest.approx(0.0, abs=CENT)
    assert lev["lp_net_irr"] == pytest.approx(0.071479, abs=BP)
    assert lev["lp_moic"] == pytest.approx(1.3900, abs=1e-4)


def test_oracle_a_leverage_is_dilutive_and_the_model_can_say_so(oracle_a):
    """The regression that catches a levered lens which only ever flatters.

    Oracle A borrows at a 7.585% constant against a deal yielding less
    than that on cost, so the LP's net return is BELOW the unlevered IRR.
    A model that cannot represent negative leverage is a model that will
    recommend debt on every deal.
    """
    projection, _, _, lev = oracle_a
    assert projection["irr"] == pytest.approx(0.073031, abs=BP)
    assert lev["lp_net_irr"] < projection["irr"]


# ── Oracle B: LTV binds, IO rolls to amortizing, a promote is paid ───

ORACLE_B = DebtTerms(rate=0.0625, amort_years=30, io_months=24, max_ltv=0.65,
                     min_dscr=1.25, min_debt_yield=0.10, orig_fee_pct=0.01,
                     exit_fee_pct=0.0)


@pytest.fixture(scope="module")
def oracle_b():
    return _build(price=10_000_000, y1_noi=750_000, growth=0.04,
                  exit_cap=0.0775, terms=ORACLE_B)


def test_oracle_b_io_rolls_to_an_amortizing_payment(oracle_b):
    """The design doc's named error, caught: "IO→amort transition payment
    not recomputed". Years 1-2 pay interest only on 6,500,000 at 6.25%;
    year 3 is the first full amortizing year."""
    _, debt, _, _ = oracle_b
    assert debt["binding_constraint"] == "ltv"
    assert debt["loan"] == pytest.approx(6_500_000.00, abs=CENT)
    assert debt["annual_debt_service"][:2] == pytest.approx(
        [406_250.00, 406_250.00], abs=CENT)
    assert debt["annual_debt_service"][2:] == pytest.approx(
        [480_259.42, 480_259.42, 480_259.42], abs=CENT)
    assert debt["payoff_balance"] == pytest.approx(6_256_487.17, abs=CENT)


def test_oracle_b_pays_a_promote_on_the_lp_attributable_residual(oracle_b):
    projection, _, su, lev = oracle_b
    assert su["total_equity"] == pytest.approx(3_665_000.00, abs=CENT)
    assert [r["levered_cf"] for r in lev["years"]] == pytest.approx(
        [307_100.00, 337_100.00, 294_290.58, 326_738.58, 5_311_997.08],
        abs=CENT)
    wf = lev["waterfall"]
    assert wf["tier1_current"] is True
    assert wf["gp"]["promote"] == pytest.approx(263_790.53, abs=CENT)
    assert lev["lp_net_irr"] == pytest.approx(0.129941, abs=BP)
    assert lev["lp_moic"] == pytest.approx(1.7146, abs=1e-4)
    assert wf["gp"]["irr"] == pytest.approx(0.224806, abs=BP)
    # Positive leverage here, unlike oracle A.
    assert projection["irr"] == pytest.approx(0.097899, abs=BP)
    assert lev["lp_net_irr"] > projection["irr"]


# ── Oracle C: the shortfall path — draw the reserve, then call ───────

ORACLE_C = DebtTerms(rate=0.075, amort_years=25, io_months=0, max_ltv=0.75,
                     min_dscr=0.0, min_debt_yield=0.0, orig_fee_pct=0.01,
                     exit_fee_pct=0.005)


@pytest.fixture(scope="module")
def oracle_c():
    return _build(price=10_000_000, y1_noi=400_000, growth=0.12,
                  exit_cap=0.0600, terms=ORACLE_C, reserve=150_000.0)


def test_oracle_c_draws_the_reserve_before_calling_capital(oracle_c):
    """Item D funds the reserve at close and puts it in `total_basis`, so
    a shortfall it covers is not a waterfall event at all — the period
    distributes 0 and nothing new is called. Only the remainder is a
    capital call."""
    _, _, _, lev = oracle_c
    year_1 = lev["years"][0]
    assert year_1["levered_cf"] == pytest.approx(-293_342.06, abs=CENT)
    assert year_1["reserve_drawn"] == pytest.approx(150_000.00, abs=CENT)
    assert year_1["capital_call"] == pytest.approx(143_342.06, abs=CENT)
    assert lev["reserve_funded"] == pytest.approx(150_000.00, abs=CENT)
    assert lev["reserve_remaining"] == pytest.approx(0.0, abs=CENT)
    # The reserve is exhausted in year 1, so years 2-4 call the full gap.
    assert [r["reserve_drawn"] for r in lev["years"][1:]] == pytest.approx(
        [0.0] * 4, abs=CENT)
    assert [r["capital_call"] for r in lev["years"][1:4]] == pytest.approx(
        [246_775.48, 195_483.24, 137_226.87], abs=CENT)
    # The exact total, not the sum of the four ROUNDED calls above — those
    # differ in the third decimal and summing them is how a "to the cent"
    # assertion quietly acquires a cent of error.
    assert lev["capital_calls_total"] == pytest.approx(722_827.6438, abs=CENT)


def test_oracle_c_never_distributes_a_negative_number(oracle_c):
    """E2's `_align_series` REJECTS a negative distribution, deliberately:
    netting one against the pref accrual pays the LP a reduced preferred
    return for the privilege of the deal losing money."""
    _, _, _, lev = oracle_c
    assert [r["distribution"] for r in lev["years"][:4]] == pytest.approx(
        [0.0] * 4, abs=CENT)
    assert all(d >= 0 for d in lev["distributions"])


def test_oracle_c_am_fee_follows_the_calls_one_period_behind(oracle_c):
    """The fee is 1% of equity outstanding at the START of the period, so
    year 2's base includes year 1's call and excludes year 2's. An
    end-of-period base would be a loop with no fixed point: a call raises
    the fee, which deepens the shortfall, which raises the call."""
    _, _, su, lev = oracle_c
    assert su["total_equity"] == pytest.approx(2_825_000.00, abs=CENT)
    assert [r["am_fee"] for r in lev["years"]] == pytest.approx(
        [28_250.00, 29_683.42, 32_151.18, 34_106.01, 35_478.28], abs=CENT)
    # Year 2's fee is exactly 1% of equity plus year 1's call.
    assert lev["years"][1]["am_fee"] == pytest.approx(
        0.01 * (2_825_000.00 + 143_342.06), abs=CENT)


def test_oracle_c_contributions_carry_the_calls_at_their_own_periods(oracle_c):
    _, _, _, lev = oracle_c
    assert lev["contributions"] == pytest.approx(
        [2_825_000.00, 143_342.06, 246_775.48, 195_483.24, 137_226.87, 0.0],
        abs=CENT)
    assert lev["lp_net_irr"] == pytest.approx(-0.009442, abs=BP)
    wf = lev["waterfall"]
    assert wf["gp"]["promote"] == pytest.approx(0.0, abs=CENT)
    # No promote was paid before the last call, so nothing is unrecovered.
    assert wf["unrecovered_promote"] == pytest.approx(0.0, abs=CENT)


def test_oracle_c_charges_the_exit_fee_at_sale_not_as_a_use_of_funds(oracle_c):
    """`build_debt_schedule` puts only origination in `financing_costs`.
    The exit fee is paid out of proceeds; carrying it as a use would
    inflate the basis and understate the return at both ends."""
    projection, debt, su, lev = oracle_c
    assert debt["exit_fee"] == pytest.approx(34_399.71, abs=CENT)
    assert su["total_uses"] == pytest.approx(
        projection["total_basis"] + debt["financing_costs"], abs=CENT)
    assert debt["exit_fee"] not in [u["amount"] for u in su["uses"]]


# ── The four traps E1 and E2 flagged for this item ───────────────────

def test_the_distribution_series_starts_at_period_zero(oracle_a):
    """Trap 3. E2's scalar-contribution shorthand cannot tell a period-0
    distribution from a series starting at year 1, and `cash_flows[1:]`
    is exactly `hold_years` long — the shape that measured LP IRR
    14.1563% against a correct 11.2437%. Both series are spelled out."""
    _, _, _, lev = oracle_a
    assert len(lev["distributions"]) == 6          # hold_years + 1
    assert len(lev["contributions"]) == 6
    assert lev["distributions"][0] == 0.0
    assert lev["contributions"][0] == pytest.approx(4_160_000.00, abs=CENT)


def test_the_waterfall_reads_the_deals_coinvest_not_the_config_default():
    """Trap 2. GP co-invest is per-deal. Called without
    `capital_structure`, `resolve_waterfall_terms` falls back to
    `config.GP_COINVEST_PCT`, and a deal edited to 25% would print a
    stack split 25/75 beside an LP net IRR computed on 10/90."""
    assert cfg.GP_COINVEST_PCT == 0.10, "fixture assumes the shipped default"
    _, _, su, lev = _build(price=10_000_000, y1_noi=750_000, growth=0.04,
                           exit_cap=0.0775, terms=ORACLE_B,
                           gp_coinvest_pct=0.25)
    assert su["gp_coinvest_pct"] == 0.25
    assert lev["waterfall"]["terms"]["gp_coinvest_pct"] == 0.25
    # The stack and the waterfall agree about whose equity it is.
    assert lev["waterfall"]["gp"]["contributions"] == pytest.approx(
        su["gp_equity"], abs=CENT)
    assert lev["waterfall"]["lp"]["contributions"] == pytest.approx(
        su["lp_equity"], abs=CENT)


def test_the_am_fee_stamp_row_carries_the_rate_and_the_base(oracle_a):
    """Trap 4. E2's stamp row says "rate and base set by the caller, not
    this module" because E2 does not charge the fee. E3a does, so the row
    must name what it charged — otherwise the stamp reads complete beside
    an LP *net* IRR while omitting the input that makes it net."""
    _, _, _, lev = oracle_a
    row = next(r for r in lev["assumption_stamp"]
               if r["key"] == "am_fee_treatment")
    assert "1.00%" in row["label"]
    assert "invested equity" in row["label"].lower()
    assert "set by the caller" not in row["label"]
    assert row["rate"] == pytest.approx(0.01)
    assert row["base"] == "invested_equity"
    # The other five questions still get answered.
    assert {r["key"] for r in lev["assumption_stamp"]} == {
        "pref_compounding", "accrual_base", "ordering", "am_fee_treatment",
        "promote_basis", "catch_up"}


def test_a_shortfall_is_reported_not_swallowed(oracle_c):
    """Trap 1's reporting half. A deal that called capital four times has
    to say so where a consumer cannot miss it."""
    _, _, _, lev = oracle_c
    assert lev["capital_calls_total"] > 0
    assert lev["called_capital_after_close"] is True


# ── The unlevered screen must not have moved ────────────────────────

def test_debt_does_not_touch_the_unlevered_projection():
    """The operator's decision on 2026-08-01: financing costs stay OUT of
    `total_basis`, so the primary 10% IRR screen is identical whether or
    not the deal carries a loan. E1's handoff prescribed the opposite;
    the tie moved instead of the projection."""
    unlevered = project_cash_flows(
        600_000, 10_000_000, 0.0, _params(0.03),
        hold_years=5, expense_ratio=0.40, costs=COSTS, exit_cap=0.0625)
    projection, debt, _, _ = _build(price=10_000_000, y1_noi=600_000,
                                    growth=0.03, exit_cap=0.0625,
                                    terms=ORACLE_A)
    assert debt["loan"] > 0, "fixture must actually carry debt"
    assert projection["total_basis"] == pytest.approx(
        unlevered["total_basis"], abs=CENT)
    assert projection["irr"] == pytest.approx(unlevered["irr"], abs=1e-12)
    assert projection["moic"] == pytest.approx(unlevered["moic"], abs=1e-12)
    assert projection["yield_on_cost"] == pytest.approx(
        unlevered["yield_on_cost"], abs=1e-12)


def test_the_exit_capitalizes_the_trailing_year_not_the_forward_year(oracle_a):
    """The DEFAULT convention (`config.EXIT_NOI_CONVENTION =
    "trailing"`, decision 5 — settled 2026-08-10). Year 5's own NOI is
    capitalized, NOT year 6 — the design doc and `tests/test_debt.py`'s
    oracle 5 use the forward convention and are about 3% higher. Neither
    is a bug; they are different conventions, the config names them, and
    `tests/test_exit_noi_convention.py` pins the forward branch."""
    projection, _, _, _ = oracle_a
    trailing = 600_000 * 1.03 ** 4
    assert projection["noi"][-1] == pytest.approx(trailing, abs=CENT)
    assert projection["exit_value"] == pytest.approx(trailing / 0.0625,
                                                     abs=CENT)
    forward = 600_000 * 1.03 ** 5
    assert projection["exit_value"] < forward / 0.0625


# ── Edges ───────────────────────────────────────────────────────────

def test_a_deal_that_supports_no_debt_is_the_unlevered_deal():
    """Covenants can bind to nothing. A zero loan must not produce a
    divide-by-zero, a NaN or a promote — it produces the unlevered deal
    with an AM fee."""
    no_debt = DebtTerms(rate=0.065, amort_years=30, max_ltv=0.0,
                        min_dscr=1.25, min_debt_yield=0.10)
    projection, debt, su, lev = _build(price=10_000_000, y1_noi=600_000,
                                       growth=0.03, exit_cap=0.0625,
                                       terms=no_debt)
    assert debt["loan"] == 0.0
    assert debt["financing_costs"] == 0.0
    assert su["senior_debt"] == 0.0
    assert su["total_equity"] == pytest.approx(projection["total_basis"],
                                               abs=CENT)
    assert [r["debt_service"] for r in lev["years"]] == pytest.approx(
        [0.0] * 5, abs=CENT)
    assert lev["lp_net_irr"] is not None
    assert lev["capital_calls_total"] == pytest.approx(0.0, abs=CENT)


@pytest.mark.parametrize("hold_years", [1, 10])
def test_hold_period_extremes_stay_aligned(hold_years):
    """`hold_years` is editable 1-10. The two series must stay
    `hold_years + 1` long at both ends, with the exit in the last year."""
    _, debt, _, lev = _build(price=10_000_000, y1_noi=600_000, growth=0.03,
                             exit_cap=0.0625, terms=ORACLE_A,
                             hold_years=hold_years)
    assert len(lev["years"]) == hold_years
    assert len(lev["contributions"]) == hold_years + 1
    assert len(lev["distributions"]) == hold_years + 1
    assert len(debt["annual_debt_service"]) == hold_years
    # Only the final year carries the sale.
    assert lev["years"][-1]["exit_proceeds"] != 0
    assert all(r["exit_proceeds"] == 0 for r in lev["years"][:-1])


def test_an_unsupported_am_fee_base_raises_rather_than_defaulting():
    """The same discipline E2 applied to its four convention fields: a
    convention this module does not implement raises, because silently
    substituting one produces a confident wrong LP net IRR — which is the
    failure mode the whole item exists to remove. "1% of asset value" is
    a real convention and roughly 2.4x the equity base on this fixture,
    so defaulting it would be a quiet 2.4x error in the fee."""
    projection = project_cash_flows(
        600_000, 10_000_000, 0.0, _params(0.03),
        hold_years=5, expense_ratio=0.40, costs=COSTS, exit_cap=0.0625)
    debt = build_debt_schedule(10_000_000, 600_000, ORACLE_A, hold_years=5)
    su = build_sources_uses(price=10_000_000, capex=0.0,
                            acquisition_cost=projection["acquisition_cost"],
                            financing_costs=debt["financing_costs"],
                            senior_debt=debt["loan"], gp_coinvest_pct=0.10)
    with pytest.raises(ValueError, match="am_fee_base"):
        build_levered_returns(projection, sources_uses=su, debt=debt,
                              waterfall_terms=resolve_waterfall_terms(),
                              am_fee_base="asset_value")


def test_config_defaults_are_what_the_module_uses():
    """No second copy of the AM fee. `config.AM_FEE_PCT` is the source and
    a drifting mirror is the failure this pins."""
    assert cfg.AM_FEE_PCT == 0.01
    assert cfg.AM_FEE_BASE == "invested_equity"
    projection, _, _, lev = _build(price=10_000_000, y1_noi=600_000,
                                   growth=0.03, exit_cap=0.0625,
                                   terms=ORACLE_A, am_fee_pct=None)
    assert lev["am_fee_pct"] == cfg.AM_FEE_PCT
    assert lev["am_fee_base"] == cfg.AM_FEE_BASE
    assert [r["am_fee"] for r in lev["years"]] == pytest.approx(
        [41_600.00] * 5, abs=CENT)


# ── One loan, three scenarios ───────────────────────────────────────

def test_the_loan_is_sized_once_off_the_base_case():
    """Sizing per scenario would hand the bear case a SMALLER loan and
    flatten its own downside — the model would understate exactly the risk
    the bear case exists to show. A lender sizes on one underwriting."""
    from model.returns_model import build_returns_model
    from registry import ScenarioType

    model = build_returns_model(
        adjusted_ttm_noi=600_000, asking_price=10_000_000, nrsf=60_000,
        expense_ratio=0.40, debt_terms=ORACLE_A)

    assert set(model["levered"]) == set(model["scenarios"])
    loans = {name: lev["debt"]["loan"] for name, lev in model["levered"].items()}
    assert len(set(loans.values())) == 1, loans
    schedules = {name: tuple(lev["debt"]["annual_debt_service"])
                 for name, lev in model["levered"].items()}
    assert len(set(schedules.values())) == 1
    assert model["debt"]["binding_constraint"] == "debt_yield"

    # The bear case still produces a WORSE LP outcome than the bull — the
    # loan being constant is what lets the scenarios differ honestly.
    bear = model["levered"][ScenarioType.BEAR]["lp_net_irr"]
    bull = model["levered"][ScenarioType.BULL]["lp_net_irr"]
    assert bear is None or bull is None or bear < bull


def test_a_projection_with_no_noi_series_raises_instead_of_sizing_on_zero():
    """Found while wiring: the scenario API relabels `noi` as
    `noi_projection`, so reading only `noi` and defaulting to 0.0 sized
    the loan on a Year 1 NOI of zero — producing a $0 debt-yield-bound
    loan and a levered lens identical to the unlevered one, with nothing
    saying the input was missing. Both names are accepted; NEITHER
    raises."""
    from model.levered import noi_series

    assert noi_series({"noi": [1.0, 2.0]}) == [1.0, 2.0]
    assert noi_series({"noi_projection": [3.0]}) == [3.0]
    with pytest.raises(ValueError, match="no NOI series"):
        noi_series({"total_basis": 10_000_000})


def test_the_capital_stack_carries_the_sized_loan():
    """`build_sources_uses` gets the real `senior_debt` and
    `financing_costs`, so debt DISPLACES equity rather than inflating
    uses — item D's promise, now that a loan actually exists."""
    from model.returns_model import build_returns_model

    model = build_returns_model(
        adjusted_ttm_noi=600_000, asking_price=10_000_000, nrsf=60_000,
        expense_ratio=0.40, debt_terms=ORACLE_A)
    su, debt = model["sources_uses"], model["debt"]
    assert su["senior_debt"] == pytest.approx(debt["loan"], abs=CENT)
    assert su["financing_costs"] == pytest.approx(debt["financing_costs"],
                                                  abs=CENT)
    base = next(s for s in model["scenarios"].values() if isinstance(s, dict))
    assert su["total_uses"] == pytest.approx(
        base["total_basis"] + debt["financing_costs"], abs=CENT)
    assert su["total_uses"] == pytest.approx(su["total_sources"], abs=CENT)


def test_the_blocking_tie_check_passes_on_a_levered_run():
    """The check item A made blocking, on a deal that now carries debt."""
    from analysis import checks
    from model.returns_model import build_returns_model

    model = build_returns_model(
        adjusted_ttm_noi=600_000, asking_price=10_000_000, nrsf=60_000,
        expense_ratio=0.40, debt_terms=ORACLE_A)
    result = next(r for r in checks.run_checks(
        checks.CheckInput(scenarios=model["scenarios"],
                          sources_uses=model["sources_uses"]),
        only={"sources_uses_ties"}))
    assert result.status == checks.PASS, result.message


# ── Repairs from the E3a audit ──────────────────────────────────────

def test_each_scenario_gets_its_own_debt_dict_not_a_shared_alias():
    """One loan is sized for the whole deal and every scenario embeds it.
    Returning the object itself made all three the same dict, so the first
    consumer to annotate per-scenario debt data in place would corrupt the
    other two."""
    from model.returns_model import build_returns_model
    from registry import ScenarioType

    model = build_returns_model(
        adjusted_ttm_noi=600_000, asking_price=10_000_000, nrsf=60_000,
        expense_ratio=0.40, debt_terms=ORACLE_A)
    bear = model["levered"][ScenarioType.BEAR]["debt"]
    bull = model["levered"][ScenarioType.BULL]["debt"]
    assert bear is not bull
    assert bear["annual_debt_service"] is not bull["annual_debt_service"]
    # Same VALUES — the loan really is sized once.
    assert bear["loan"] == bull["loan"]
    # Mutating one leaves the other untouched.
    bear["loan"] = -1
    assert bull["loan"] != -1


def test_a_debt_payload_missing_its_payoff_raises():
    """Defaulting a missing payoff to 0.0 computes the exit as though the
    loan were forgiven at sale, and reports an LP net IRR that is too
    HIGH with no error anywhere. A zero VALUE is fine; a missing KEY is a
    broken payload."""
    projection = project_cash_flows(
        600_000, 10_000_000, 0.0, _params(0.03),
        hold_years=5, expense_ratio=0.40, costs=COSTS, exit_cap=0.0625)
    debt = build_debt_schedule(10_000_000, 600_000, ORACLE_A, hold_years=5)
    su = build_sources_uses(price=10_000_000, capex=0.0,
                            acquisition_cost=projection["acquisition_cost"],
                            financing_costs=debt["financing_costs"],
                            senior_debt=debt["loan"], gp_coinvest_pct=0.10)
    broken = {k: v for k, v in debt.items() if k != "payoff_balance"}
    with pytest.raises(ValueError, match="payoff_balance"):
        build_levered_returns(projection, sources_uses=su, debt=broken,
                              waterfall_terms=resolve_waterfall_terms())
    # A zero exit fee is legitimate and must still work.
    zeroed = {**debt, "exit_fee": 0.0}
    assert build_levered_returns(
        projection, sources_uses=su, debt=zeroed,
        waterfall_terms=resolve_waterfall_terms())["lp_net_irr"] is not None


def test_a_blank_am_fee_falls_back_to_config_instead_of_crashing():
    """An HTML form posts "" for a cleared numeric field. Every sibling
    resolver in this repo tests `not in (None, "")`; `is None` would make
    E3b's form field raise on `float("")` the day it lands."""
    projection = project_cash_flows(
        600_000, 10_000_000, 0.0, _params(0.03),
        hold_years=5, expense_ratio=0.40, costs=COSTS, exit_cap=0.0625)
    debt = build_debt_schedule(10_000_000, 600_000, ORACLE_A, hold_years=5)
    su = build_sources_uses(price=10_000_000, capex=0.0,
                            acquisition_cost=projection["acquisition_cost"],
                            financing_costs=debt["financing_costs"],
                            senior_debt=debt["loan"], gp_coinvest_pct=0.10)
    lev = build_levered_returns(projection, sources_uses=su, debt=debt,
                                waterfall_terms=resolve_waterfall_terms(),
                                am_fee_pct="", am_fee_base="")
    assert lev["am_fee_pct"] == cfg.AM_FEE_PCT
    assert lev["am_fee_base"] == cfg.AM_FEE_BASE


def test_sizing_raises_when_the_base_scenario_is_missing():
    """The loan is sized off the BASE case. Falling back to whichever
    scenario computed first could price the debt on the bull case's
    richer NOI, and `sources_uses_ties` cannot catch that — it checks the
    stack's internal arithmetic, not which NOI justified the loan."""
    from model.returns_model import build_returns_model
    from registry import ScenarioType

    real = build_returns_model(
        adjusted_ttm_noi=600_000, asking_price=10_000_000, nrsf=60_000,
        expense_ratio=0.40, debt_terms=ORACLE_A)
    assert isinstance(real["scenarios"][ScenarioType.BASE], dict)

    import model.returns_model as rm

    def _no_base(*args, **kwargs):
        # The key REMOVED, not set to None: `_build_summary_table` does
        # `scenarios.get(name, {})` and so tolerates an absent key while
        # crashing on an explicit None. Absent is therefore the shape
        # that actually reaches the debt-sizing fallback.
        scen = dict(real["scenarios"])
        scen.pop(ScenarioType.BASE)
        return scen

    original = rm.run_scenarios
    rm.run_scenarios = _no_base
    try:
        with pytest.raises(ValueError, match="base scenario is missing"):
            build_returns_model(
                adjusted_ttm_noi=600_000, asking_price=10_000_000,
                nrsf=60_000, expense_ratio=0.40, debt_terms=ORACLE_A)
    finally:
        rm.run_scenarios = original


def test_the_tie_check_catches_a_stack_built_without_the_debt_modules_fee():
    """The failure the self-referential version could not see. A caller
    that forgets `financing_costs=debt["financing_costs"]` produces a
    `total_uses` missing the fee AND a `financing_costs` of 0 — both wrong
    the same way, so `uses == basis + 0` reconciled and the check PASSED
    on a deal underfunded by the entire origination fee."""
    from analysis import checks
    from model.returns_model import build_returns_model

    model = build_returns_model(
        adjusted_ttm_noi=600_000, asking_price=10_000_000, nrsf=60_000,
        expense_ratio=0.40, debt_terms=ORACLE_A)
    debt = model["debt"]
    assert debt["financing_costs"] > 0

    base = next(s for s in model["scenarios"].values()
                if isinstance(s, dict))
    forgot_the_fee = build_sources_uses(
        price=10_000_000, capex=0.0,
        acquisition_cost=base["acquisition_cost"],
        senior_debt=debt["loan"], gp_coinvest_pct=0.10)   # no financing_costs

    # Self-consistent, and wrong: it ties to the DCF basis exactly.
    assert forgot_the_fee["total_uses"] == pytest.approx(
        base["total_basis"], abs=CENT)

    without_debt = next(r for r in checks.run_checks(
        checks.CheckInput(scenarios=model["scenarios"],
                          sources_uses=forgot_the_fee),
        only={"sources_uses_ties"}))
    assert without_debt.status == checks.PASS      # the old blind spot

    with_debt = next(r for r in checks.run_checks(
        checks.CheckInput(scenarios=model["scenarios"],
                          sources_uses=forgot_the_fee, debt=debt),
        only={"sources_uses_ties"}))
    assert with_debt.status == checks.FAIL
    assert with_debt.severity == checks.BLOCKING
    assert "financing costs" in with_debt.message


def test_a_loan_maturing_inside_the_hold_reaches_the_check_register():
    """Previously a `logger.warning` nobody reads, while the results page
    showed a levered IRR computed as though the loan amortized past its
    own maturity. Item E3a put a sized loan on every deal, so the
    condition went live."""
    from analysis import checks

    short_term = DebtTerms(rate=0.065, amort_years=30, term_years=3,
                           max_ltv=0.65, min_dscr=1.25, min_debt_yield=0.10)
    debt = build_debt_schedule(10_000_000, 600_000, short_term, hold_years=5)
    assert debt["matures_before_exit"] is True

    result = next(r for r in checks.run_checks(
        checks.CheckInput(debt=debt), only={"loan_matures_before_exit"}))
    assert result.status == checks.FAIL
    assert result.severity == checks.ADVISORY
    assert "matures in year 3" in result.message
    assert "5 years" in result.message

    ok = build_debt_schedule(10_000_000, 600_000, ORACLE_A, hold_years=5)
    passing = next(r for r in checks.run_checks(
        checks.CheckInput(debt=ok), only={"loan_matures_before_exit"}))
    assert passing.status == checks.PASS

    # A deal with no debt is skipped, not passed — we did not look.
    no_debt = build_debt_schedule(
        10_000_000, 600_000,
        DebtTerms(rate=0.065, amort_years=30, max_ltv=0.0, min_dscr=0.0,
                  min_debt_yield=0.0), hold_years=5)
    skipped = next(r for r in checks.run_checks(
        checks.CheckInput(debt=no_debt), only={"loan_matures_before_exit"}))
    assert skipped.status == checks.SKIPPED


def test_the_result_is_json_safe_all_the_way_down(oracle_a):
    """These results are persisted to Postgres JSONB.
    `webapp.services.json_safe` falls back to `str(obj)`, so a frozen
    dataclass anywhere in the payload persists as an unqueryable string —
    the bug E2 found in its own `terms` key. Debt terms are the same
    shape and would have repeated it."""
    import json

    from webapp.services import json_safe

    _, _, _, lev = oracle_a
    assert isinstance(lev["debt"]["terms"], dict)
    assert lev["debt"]["terms"]["rate"] == pytest.approx(0.065)
    assert isinstance(lev["waterfall"]["terms"], dict)
    encoded = json.dumps(json_safe(lev))
    assert "DebtTerms(" not in encoded
    assert "WaterfallTerms(" not in encoded


# ════════════════════════════════════════════════════════════════════
# Item E3b — the surfaces
#
# E3a computed and persisted the levered lens; nothing rendered it.
# These cover the four places it now surfaces (assumptions inputs,
# results page, memo, workbook) and the traps E3a flagged for this item.
# They are here rather than in a new module because they are the same
# item's tests and this file is already where item E3 lives.
# ════════════════════════════════════════════════════════════════════

from django.http import QueryDict            # noqa: E402


def _levered_form(**vals):
    """An AssumptionsForm carrying only the levered block, submitted the
    way a browser does — every value a string."""
    from webapp.forms import AssumptionsForm
    return AssumptionsForm(data={k: str(v) for k, v in vals.items()})


def _levered_run(oracle):
    """The `result_json` shape the results page, memo and workbook read:
    one `debt` block and one `levered` entry per scenario, exactly as
    `webapp.services` persists it."""
    projection, debt, _, lev = oracle
    return {
        "debt": debt,
        "levered": {sc: lev for sc in ("bear", "base", "bull")},
        "scenario_results": {
            sc: {"irr": projection["irr"], "hold_years": 5,
                 "noi_projection": projection["noi"]}
            for sc in ("bear", "base", "bull")},
    }


# ── The assumptions inputs ──────────────────────────────────────────

def test_every_levered_percent_field_converts_to_a_decimal_on_save():
    """The percent-vs-decimal boundary E3a flagged as the thing that
    breaks a naive form. `rate=6.5` meaning 6.5% priced a $6.5M loan at
    $3,520,833/mo against a correct $43,888/mo — eighty times wrong, in a
    payment that feeds every levered return. `DebtTerms` now RAISES above
    1.0, so a naive form fails every save instead; either way the
    conversion belongs here, as it does for every other percentage on
    this page."""
    from webapp.forms import (submitted_am_fee_pct, submitted_debt_terms,
                              submitted_waterfall_terms)

    form = _levered_form(debt_rate=6.5, debt_max_ltv=65,
                         debt_min_debt_yield=10, debt_orig_fee_pct=1,
                         debt_exit_fee_pct=0.5, debt_min_dscr=1.25,
                         debt_amort_years=25, debt_io_months=24,
                         debt_term_years=10, wf_pref_rate=8,
                         wf_promote_split=20, am_fee_pct=1)
    assert form.is_valid(), form.errors

    debt = submitted_debt_terms(form.cleaned_data)
    assert debt["rate"] == pytest.approx(0.065)
    assert debt["max_ltv"] == pytest.approx(0.65)
    assert debt["min_debt_yield"] == pytest.approx(0.10)
    assert debt["orig_fee_pct"] == pytest.approx(0.01)
    assert debt["exit_fee_pct"] == pytest.approx(0.005)
    # A coverage RATIO, not a percentage — 1.25x must survive untouched,
    # which is also why DebtTerms exempts it from its own >1.0 guard.
    assert debt["min_dscr"] == pytest.approx(1.25)
    # Integers stay integers: a 24.0-month IO would not compare equal to
    # the config default and would write a pointless override row.
    assert debt["amort_years"] == 25 and isinstance(debt["amort_years"], int)
    assert debt["io_months"] == 24
    assert debt["term_years"] == 10

    wf = submitted_waterfall_terms(form.cleaned_data)
    assert wf["pref_rate"] == pytest.approx(0.08)
    assert wf["promote_split"] == pytest.approx(0.20)
    assert submitted_am_fee_pct(form.cleaned_data) == pytest.approx(0.01)


@pytest.mark.django_db
def test_the_levered_block_round_trips_through_a_save():
    """Save → store → redisplay must land on the same whole numbers the
    analyst typed. A one-way conversion is how a 6.25% loan redisplays as
    0.0625% and gets "corrected" to 6.25 basis points on the next save."""
    from webapp.forms import build_initial, build_overrides
    from webapp.models import Deal

    deal = Deal.objects.create(deal_id="lev-rt", property_name="RT")
    # Every value differs from its config default on purpose: this section
    # stores DELTAS, so a field set to the default writes no row and the
    # assertion below would pass for the wrong reason.
    form = _levered_form(debt_rate=6.75, debt_max_ltv=70, debt_min_dscr=1.30,
                         debt_amort_years=25, debt_io_months=0,
                         debt_term_years=7, debt_min_debt_yield=9,
                         debt_orig_fee_pct=1, debt_exit_fee_pct=0,
                         wf_pref_rate=9, wf_promote_split=25,
                         wf_pref_compounding="annual", wf_ordering="roc_first",
                         am_fee_pct=1.5)
    assert form.is_valid(), form.errors
    deal.assumption_overrides = build_overrides(form.cleaned_data,
                                                QueryDict(""), deal)
    deal.save()

    stored = deal.assumption_overrides
    assert stored["debt_terms"]["rate"] == pytest.approx(0.0675)
    assert stored["debt_terms"]["max_ltv"] == pytest.approx(0.70)
    assert stored["debt_terms"]["min_dscr"] == pytest.approx(1.30)
    assert stored["debt_terms"]["term_years"] == 7
    assert stored["waterfall_terms"]["pref_rate"] == pytest.approx(0.09)
    assert stored["am_fee_pct"] == pytest.approx(0.015)

    initial = build_initial(deal)
    assert initial["debt_rate"] == pytest.approx(6.75)
    assert initial["debt_max_ltv"] == pytest.approx(70)
    assert initial["debt_min_dscr"] == pytest.approx(1.30)
    assert initial["debt_term_years"] == 7
    assert initial["wf_pref_rate"] == pytest.approx(9)
    assert initial["wf_promote_split"] == pytest.approx(25)
    assert initial["am_fee_pct"] == pytest.approx(1.5)


@pytest.mark.django_db
def test_submitting_the_defaults_writes_no_override_rows():
    """Deltas only, like every other section. A block that stores its own
    defaults makes `applied_overrides` claim an analyst decision nobody
    made, and freezes this deal against a later config change."""
    from webapp.forms import build_initial, build_overrides
    from webapp.models import Deal

    deal = Deal.objects.create(deal_id="lev-def", property_name="Def")
    initial = build_initial(deal)
    posted = {k: v for k, v in initial.items()
              if v is not None and (k.startswith(("debt_", "wf_"))
                                    or k == "am_fee_pct")}
    form = _levered_form(**posted)
    assert form.is_valid(), form.errors
    out = build_overrides(form.cleaned_data, QueryDict(""), deal)
    assert "debt_terms" not in out
    assert "waterfall_terms" not in out
    assert "am_fee_pct" not in out


@pytest.mark.django_db
def test_an_unlevered_deal_prefills_the_unlevered_pref_and_stores_no_override():
    """The pref default follows the deal's leverage — 8% / 6% (LPA,
    2026-08-12) — and the two halves of the form must agree on which.

    The failure this pins is quiet in both directions. If the prefill
    kept showing 8% on a deal with no debt, the operator would look at a
    pref the fund does not charge. If the prefill showed 6% but the
    override diff still measured against 8%, merely OPENING and saving
    the page would store 6% as an analyst decision nobody made —
    freezing this deal against a later change to the rate and claiming
    an override in `applied_overrides` that was really just the default.
    """
    from webapp.forms import build_initial, build_overrides
    from webapp.models import Deal

    deal = Deal.objects.create(
        deal_id="lev-unlev", property_name="All Equity",
        assumption_overrides={"debt_terms": {"max_ltv": 0}})

    initial = build_initial(deal)
    assert initial["wf_pref_rate"] == pytest.approx(
        cfg.PREF_RATE_UNLEVERED * 100)

    posted = {k: v for k, v in initial.items()
              if v is not None and (k.startswith(("debt_", "wf_"))
                                    or k == "am_fee_pct")}
    out = build_overrides(_valid(_levered_form(**posted)), QueryDict(""), deal)
    assert "waterfall_terms" not in out
    # The deal's own max_ltv=0 is still a real override and must survive.
    assert out["debt_terms"]["max_ltv"] == 0


def _valid(form):
    assert form.is_valid(), form.errors
    return form.cleaned_data


@pytest.mark.django_db
def test_unexposed_debt_and_waterfall_keys_survive_a_save():
    """The trap this form's shape creates. Six keys have no field —
    `loan_type`, `index_rate`, `spread`, `accrual_base`,
    `am_fee_treatment`, `catch_up` — so rebuilding each section purely
    from the form would DELETE a CLI-set floating rate on the first
    unrelated save here. The deal keeps running, at a different cost of
    debt, with nothing anywhere saying so. Every capital-block key had a
    field, so this could not arise before item E3b."""
    from webapp.forms import build_overrides
    from webapp.models import Deal

    deal = Deal.objects.create(
        deal_id="lev-carry", property_name="Carry",
        assumption_overrides={
            "debt_terms": {"index_rate": 0.045, "spread": 0.025,
                           "loan_type": "bridge_floating"},
            "waterfall_terms": {"accrual_base": "contributed",
                                "am_fee_treatment": "above_waterfall"}})
    # An edit that touches only the LTV — nothing about the rate.
    form = _levered_form(debt_max_ltv=60)
    assert form.is_valid(), form.errors
    out = build_overrides(form.cleaned_data, QueryDict(""), deal)

    assert out["debt_terms"]["index_rate"] == pytest.approx(0.045)
    assert out["debt_terms"]["spread"] == pytest.approx(0.025)
    assert out["debt_terms"]["loan_type"] == "bridge_floating"
    assert out["debt_terms"]["max_ltv"] == pytest.approx(0.60)
    assert out["waterfall_terms"]["accrual_base"] == "contributed"
    assert out["waterfall_terms"]["am_fee_treatment"] == "above_waterfall"


@pytest.mark.django_db
def test_a_floating_rate_deal_prefills_a_blank_rate_and_stays_floating():
    """Why the prefill reads the RESOLVED terms and not
    `{**config, **saved}`. `resolve_debt_terms` clears the seeded fixed
    rate for a floating override; merging config in directly would
    prefill 6.25%, post it, and the resolver's "both named, fixed wins"
    branch would silently convert the deal to fixed paper."""
    from model.debt import resolve_debt_terms
    from webapp.forms import build_initial, build_overrides
    from webapp.models import Deal

    deal = Deal.objects.create(
        deal_id="lev-float", property_name="Float",
        assumption_overrides={"debt_terms": {"index_rate": 0.045,
                                             "spread": 0.025}})
    assert build_initial(deal)["debt_rate"] is None

    form = _levered_form(debt_amort_years=25)      # rate left blank
    assert form.is_valid(), form.errors
    out = build_overrides(form.cleaned_data, QueryDict(""), deal)
    resolved = resolve_debt_terms(out["debt_terms"])
    assert resolved.rate is None
    assert resolved.all_in_rate() == pytest.approx(0.070)


@pytest.mark.parametrize("bad", [
    # The percent-vs-decimal typo E1's review measured: `rate=6.5`
    # meaning 6.5% priced a $6.5M loan eighty times too expensively.
    {"debt_rate": 650},
    # Would fall through `monthly_payment`'s degenerate branch to a
    # sizing constant of 1200%/yr — a loan two orders of magnitude too
    # small, reported without complaint.
    {"debt_amort_years": 0},
    # Switches off both the full-IO test and the maturity warning.
    {"debt_term_years": 0},
    {"debt_io_months": -1},
    {"debt_min_dscr": -1},
    # NaN passes every `< 0` guard untouched (every comparison against
    # NaN is False) and then poisons the arithmetic downstream.
    {"debt_min_dscr": "nan"},
    {"wf_promote_split": 150},
    # The BOUNDARY, and the case that proves the field bounds are not the
    # whole story (review finding). `DebtTerms` rejects a decimal above
    # 1.0, so `max_value=100` covers it; `WaterfallTerms` requires
    # STRICTLY below 1.0, because a 100% promote hands the GP the entire
    # residual. 100 is in bounds, converts to exactly 1.0, and is caught
    # only by the resolver backstop.
    {"wf_promote_split": 100},
    {"wf_pref_rate": 100},
])
def test_a_save_can_never_store_terms_the_model_will_reject(bad):
    """Each case is a `DebtTerms` / `WaterfallTerms` raise, refused while
    a human is looking at the page instead of surfacing as a failed run
    twenty minutes later. Some are caught by a field bound and some only
    by the resolver backstop — which is the point: the form does not
    re-list the model's rules, it runs them."""
    assert not _levered_form(**bad).is_valid()


def test_the_resolver_backstop_catches_what_no_field_bound_can():
    """`clean()` validates by calling the REAL resolvers rather than
    re-listing their rules — the duplicated-constant divergence this repo
    has a rule against, and one that would go stale the first time a
    guard is added to either dataclass.

    Every case above is reachable today only through a field, so this
    exercises the backstop directly: no field can produce this input now,
    which is exactly why the guard is worth keeping. Adding a rule to
    `DebtTerms` must not silently turn a save into a failed run.

    Reported as a NON-field error, so `cleaned_data` keeps its keys —
    `assumptions_preview` reads them on an invalid form by design (the
    item D lesson)."""
    from webapp.forms import AssumptionsForm

    form = AssumptionsForm(data={})
    assert form.is_valid(), form.errors
    form._validate_levered_terms({"debt_min_dscr": float("nan")})
    assert "Debt terms" in str(form.non_field_errors())

    wf = AssumptionsForm(data={})
    assert wf.is_valid(), wf.errors
    wf._validate_levered_terms({"wf_pref_compounding": "quarterly"})
    assert "Waterfall terms" in str(wf.non_field_errors())


def test_the_waterfall_selectors_offer_only_implemented_values():
    """`accrual_base="committed"`, `am_fee_treatment="netted_from_lp"`
    and `catch_up=True` are real conventions this repo does NOT
    implement — `WaterfallTerms` raises on each. A dropdown whose second
    option crashes the run is a trap, so they get no field at all and
    stay in the assumption stamp instead."""
    from webapp.forms import WF_FORM_KEYS
    from webapp.forms import AssumptionsForm

    assert "accrual_base" not in WF_FORM_KEYS
    assert "am_fee_treatment" not in WF_FORM_KEYS
    assert "catch_up" not in WF_FORM_KEYS

    form = AssumptionsForm()
    for name, allowed in (("wf_pref_compounding", {"annual", "simple"}),
                          ("wf_ordering", {"roc_first", "pref_first"})):
        offered = {value for value, _ in form.fields[name].choices}
        assert offered == allowed


@pytest.mark.django_db
def test_an_edited_loan_term_makes_the_maturity_advisory_fire():
    """`loan_matures_before_exit` could not fire before this item: the
    config term is 10 years and the hold is capped at 10. Making
    `term_years` editable is what switches it on, so the form path is
    asserted end to end rather than only the model path above."""
    from analysis import checks
    from model.debt import build_debt_schedule, resolve_debt_terms
    from webapp.forms import build_overrides
    from webapp.models import Deal

    deal = Deal.objects.create(deal_id="lev-mat", property_name="Mat")
    form = _levered_form(debt_term_years=3, hold_years=5)
    assert form.is_valid(), form.errors
    out = build_overrides(form.cleaned_data, QueryDict(""), deal)
    assert out["debt_terms"]["term_years"] == 3

    debt = build_debt_schedule(10_000_000, 600_000,
                               resolve_debt_terms(out["debt_terms"]),
                               hold_years=5)
    result = next(r for r in checks.run_checks(
        checks.CheckInput(debt=debt), only={"loan_matures_before_exit"}))
    assert result.status == checks.FAIL
    assert result.severity == checks.ADVISORY


# ── The results page ────────────────────────────────────────────────

def test_the_results_lens_is_absent_on_a_run_that_predates_the_levered_layer():
    """Old runs carry no `levered` key. The whole block is gated rather
    than degrading into a table of N/A, which would read as "this deal
    supports no debt" instead of "this run was computed before the lens
    existed"."""
    from webapp.results import levered_context

    assert levered_context({})["has_levered"] is False
    assert levered_context({"debt": {}, "levered": {}})["has_levered"] is False


def test_the_results_lens_renders_the_headline_the_loan_and_the_stamp(oracle_a):
    from webapp.results import levered_context

    ctx = levered_context(_levered_run(oracle_a))
    assert ctx["has_levered"] is True

    labels = {row["label"]: row["cells"] for row in ctx["levered_rows"]}
    assert labels["5-Year LP Net IRR"] == ["7.1%", "7.1%", "7.1%"]
    assert labels["5-Year LP MOIC"] == ["1.39x", "1.39x", "1.39x"]
    assert labels["GP Promote"] == ["$0", "$0", "$0"]

    loan = dict(ctx["loan_rows"])
    assert loan["Loan Amount"] == "$6,000,000"
    # The debt module's own label table, never a second copy of it.
    assert loan["Bound By"] == "Min Debt Yield"
    assert loan["All-In Rate"] == "6.50%"
    assert loan["Interest-Only"] == "0 mos"      # zero is an answer here
    assert loan["Equity Required"] == "$4,160,000"

    assert len(ctx["levered_years"]) == 5
    assert ctx["levered_years"][0]["am_fee"] == "$41,600"

    # No LP net IRR leaves the building without its stamp — six rows
    # since the catch-up confirmation, and each one changes the number.
    stamp = {row["label"] for row in ctx["levered_stamp"]}
    assert len(ctx["levered_stamp"]) == 6
    assert any("of invested equity" in label for label in stamp)


def test_the_results_lens_says_when_leverage_is_dilutive(oracle_a):
    """Oracle A prints 7.1479% LP net against a 7.3031% unlevered IRR.
    A reader who assumes levered must beat unlevered reads that as a bug,
    so the page says it first."""
    from webapp.results import levered_context

    ctx = levered_context(_levered_run(oracle_a))
    assert ctx["levered_dilutive"] == "Bear, Base, Bull"


def test_the_results_lens_survives_a_non_converging_irr(oracle_a):
    """`run_waterfall` returns None rather than NaN when the IRR does not
    converge, because `json.dumps(nan)` is invalid JSON. The page must
    print N/A, not raise, and must not report a None IRR as dilutive."""
    from webapp.results import levered_context

    run = _levered_run(oracle_a)
    run["levered"] = {sc: {**run["levered"][sc], "lp_net_irr": None,
                           "lp_moic": None}
                      for sc in ("bear", "base", "bull")}
    ctx = levered_context(run)
    labels = {row["label"]: row["cells"] for row in ctx["levered_rows"]}
    assert labels["5-Year LP Net IRR"] == ["N/A", "N/A", "N/A"]
    assert ctx["levered_dilutive"] == ""


def test_every_display_surface_reads_one_rate_and_one_covenant_label():
    """`DebtTerms.all_in_rate()` RAISES on an unresolvable pair, which is
    right when a rate is about to price a loan and wrong when a page is
    rendering a run stored months ago — a display surface that 500s on a
    stored payload loses the rest of the page. So `displayed_rate` is the
    None-returning twin, and it lives in `model.debt` BESIDE the formula
    it mirrors: the first draft of this item had the same three lines
    copied into the results page, the memo and the workbook (audit
    finding), which is three chances to disagree about one loan's rate.
    """
    import inspect

    from model.debt import binding_constraint_label, displayed_rate
    from output import excel_writer, memo_writer
    from webapp import results

    assert displayed_rate({"rate": 0.0625}) == pytest.approx(0.0625)
    assert displayed_rate({"rate": None, "index_rate": 0.045,
                           "spread": 0.025}) == pytest.approx(0.070)
    # Half a floating pair is not a rate. None, not the bare index — a
    # spread-only loan priced at 2.25% is a confident wrong number.
    assert displayed_rate({"rate": None, "spread": 0.025}) is None
    assert displayed_rate({}) is None
    assert displayed_rate(None) is None

    assert binding_constraint_label({"binding_constraint": "dscr"}) \
        == "Min DSCR"
    assert binding_constraint_label({}) == "N/A"
    assert binding_constraint_label(None) == "N/A"

    # No surface may grow its own copy back. `index_rate` appearing in a
    # display module means someone re-derived the rate locally.
    for module in (results, memo_writer, excel_writer):
        assert "index_rate" not in inspect.getsource(module), (
            f"{module.__name__} re-derives the all-in rate — call "
            f"model.debt.displayed_rate instead")


def test_the_results_lens_handles_a_priced_deal_that_supports_no_debt(oracle_a):
    """Covenants can size a real deal to zero debt. The lens still has an
    AM fee, a waterfall and an LP net IRR to show, so it renders — with
    the flag saying the figures are the all-equity case. Every rate and
    ratio the loan strip would print is None here, and none of them may
    render as a number (audit coverage gap)."""
    from webapp.results import levered_context

    run = _levered_run(oracle_a)
    run["debt"] = {**run["debt"], "loan": 0.0, "ltv": None,
                   "debt_yield": None, "dscr_year_1": None,
                   "payoff_balance": 0.0, "origination_fee": 0.0}
    ctx = levered_context(run)
    assert ctx["has_levered"] is True
    assert ctx["levered_no_loan"] is True
    loan = dict(ctx["loan_rows"])
    assert loan["Loan Amount"] == "$0"
    assert loan["LTV"] == "N/A"
    assert loan["Year-1 DSCR"] == "N/A"
    assert loan["Debt Yield"] == "N/A"


# ── The memo and the workbook ───────────────────────────────────────

def test_the_levered_lens_reaches_the_memo(tmp_path, mock_cim_data, oracle_a):
    from docx import Document

    from output.memo_writer import generate_memo

    _, debt, _, lev = oracle_a
    path = generate_memo(
        property_name="Levered Memo", cim_data=mock_cim_data,
        gate_results=[], market_analysis={}, physical_analysis={},
        financial_analysis={}, rent_analysis={},
        scenario_results=_levered_run(oracle_a)["scenario_results"],
        value_add={}, risk_analysis={}, max_offer={},
        levered={sc: lev for sc in ("bear", "base", "bull")}, debt=debt,
        output_dir=str(tmp_path))
    text = "\n".join(p.text for p in Document(path).paragraphs)
    tables = [c.text for t in Document(path).tables for r in t.rows
              for c in r.cells]

    assert "Levered Returns (LP Net)" in text
    assert "Levered Assumptions" in text
    assert "$6,000,000" in text                 # the sized loan
    assert "Min Debt Yield" in text             # what bound it
    assert "DILUTIVE" in text                   # stated, not left to infer
    assert "LP Net IRR" in tables
    assert "7.1%" in tables
    # The stamp is not optional: it is what makes "net" mean anything.
    assert "1.00% of invested equity" in text


def test_the_memo_and_workbook_build_with_no_levered_payload(tmp_path,
                                                             mock_cim_data):
    """A deal with no NOI or no asking price prices no loan. Both writers
    must still build — the same contract `sources_uses` and `checks`
    already follow."""
    from docx import Document
    from openpyxl import load_workbook

    from output.excel_writer import generate_excel
    from output.memo_writer import generate_memo

    memo = generate_memo(
        property_name="No Debt", cim_data=mock_cim_data, gate_results=[],
        market_analysis={}, physical_analysis={}, financial_analysis={},
        rent_analysis={}, scenario_results={}, value_add={},
        risk_analysis={}, max_offer={}, levered=None, debt=None,
        output_dir=str(tmp_path))
    assert "Levered Returns (LP Net)" not in "\n".join(
        p.text for p in Document(memo).paragraphs)

    xlsx = generate_excel(
        property_name="No Debt", cim_data=mock_cim_data,
        financial_analysis={}, scenario_results={}, sensitivity={},
        max_offer={}, levered=None, debt=None, output_dir=str(tmp_path))
    assert "Levered Returns" not in load_workbook(xlsx).sheetnames


def test_the_levered_lens_reaches_the_workbook(tmp_path, mock_cim_data,
                                               oracle_a):
    from openpyxl import load_workbook

    from output.excel_writer import generate_excel

    _, debt, _, lev = oracle_a
    path = generate_excel(
        property_name="Levered Model", cim_data=mock_cim_data,
        financial_analysis={},
        scenario_results=_levered_run(oracle_a)["scenario_results"],
        sensitivity={}, max_offer={},
        levered={sc: lev for sc in ("bear", "base", "bull")}, debt=debt,
        output_dir=str(tmp_path))
    wb = load_workbook(path)
    assert "Levered Returns" in wb.sheetnames
    ws = wb["Levered Returns"]
    labels = {ws.cell(row=r, column=1).value for r in range(1, ws.max_row + 1)}
    assert "Loan Amount" in labels
    assert "LP Net IRR" in labels
    assert "Base-Case Equity Cash Flow" in labels
    # The stamp travels with the number here too.
    assert any(isinstance(v, str) and "AM fee above" in v
               for v in labels if v)

    values = {ws.cell(row=r, column=1).value: ws.cell(row=r, column=2).value
              for r in range(1, ws.max_row + 1)}
    assert values["Loan Amount"] == pytest.approx(6_000_000.00, abs=CENT)
    assert values["Bound By"] == "Min Debt Yield"


def test_the_writers_do_not_mutate_the_payload_they_are_handed(tmp_path,
                                                               mock_cim_data,
                                                               oracle_a):
    """This item is presentation only, and this is the assertion that
    keeps it that way. One loan is sized for the whole deal and every
    scenario's result embeds it, so a writer annotating the dict in place
    would corrupt what the other surfaces read — the aliasing hazard E3a
    deep-copied `debt` to prevent."""
    import copy

    from output.excel_writer import generate_excel
    from output.memo_writer import generate_memo

    _, debt, _, lev = oracle_a
    levered = {sc: lev for sc in ("bear", "base", "bull")}
    scen = _levered_run(oracle_a)["scenario_results"]
    before = copy.deepcopy((debt, levered, scen))

    generate_memo(
        property_name="Immutable", cim_data=mock_cim_data, gate_results=[],
        market_analysis={}, physical_analysis={}, financial_analysis={},
        rent_analysis={}, scenario_results=scen, value_add={},
        risk_analysis={}, max_offer={}, levered=levered, debt=debt,
        output_dir=str(tmp_path))
    generate_excel(
        property_name="Immutable", cim_data=mock_cim_data,
        financial_analysis={}, scenario_results=scen, sensitivity={},
        max_offer={}, levered=levered, debt=debt, output_dir=str(tmp_path))

    assert (debt, levered, scen) == before


@pytest.mark.django_db
def test_the_returns_tab_template_actually_renders_the_lens(oracle_a):
    """A context builder that returns the right dict proves nothing about
    the page. Django resolves missing variables to the empty string, so a
    template can silently render NOTHING and every context test still
    passes — and the tag errors that DO raise (a bad tuple unpack in the
    loan strip, for one) only surface at render time. So render it."""
    from django.template.loader import render_to_string

    from webapp.results import levered_context

    html = render_to_string("webapp/_tab_returns.html",
                            levered_context(_levered_run(oracle_a)))
    assert "Levered Returns" in html
    assert "$6,000,000" in html                  # the sized loan
    assert "Min Debt Yield" in html              # the binding covenant
    assert "7.1%" in html                        # LP net IRR
    assert "$41,600" in html                     # year-1 AM fee
    assert "dilutive" in html                    # said, not left to infer
    assert "1.00% of invested equity" in html    # the stamp travels


@pytest.mark.django_db
def test_the_returns_tab_renders_unchanged_without_a_levered_payload():
    """The unlevered screen is the primary gate and must not depend on
    the lens existing — a run stored before item E3a still renders."""
    from django.template.loader import render_to_string

    from webapp.results import levered_context

    html = render_to_string("webapp/_tab_returns.html",
                            {**levered_context({}),
                             "scenario_rows": [{"label": "5-Year IRR",
                                                "cells": ["1%", "2%", "3%"]}]})
    assert "Static Returns" in html
    assert "Levered Returns" not in html


@pytest.mark.django_db
def test_the_assumptions_page_renders_every_levered_input(client,
                                                          django_user_model,
                                                          settings, tmp_path):
    """Fourteen fields declared on a form prove nothing about the page.
    This asserts each one reaches the DOM, prefilled from config, with the
    percents shown as WHOLE numbers — which is the boundary that priced a
    $6.5M loan eighty times too expensively when it was got wrong."""
    from webapp.models import Deal

    settings.CIM_DEALS_DIR = str(tmp_path)
    user = django_user_model.objects.create_user(username="lev", password="x")
    client.force_login(user)
    deal = Deal.objects.create(
        deal_id="lev-page", property_name="Lev Page", extract_status="done",
        cim_json={"property_name": "Lev Page", "asking_price": 10_000_000.0,
                  "nrsf": 50_000.0, "ttm_noi": 600_000.0},
        extraction_report={"missing": []})

    html = client.get(f"/deals/{deal.pk}/assumptions/").content.decode()
    assert "Debt &amp; Waterfall" in html
    for name in ("debt_rate", "debt_amort_years", "debt_io_months",
                 "debt_term_years", "debt_max_ltv", "debt_min_dscr",
                 "debt_min_debt_yield", "debt_orig_fee_pct",
                 "debt_exit_fee_pct", "wf_pref_rate", "wf_promote_split",
                 "wf_pref_compounding", "wf_ordering", "am_fee_pct"):
        assert f'name="{name}"' in html, f"{name} never reached the page"
    # Whole numbers, not decimals: 6.25 for a 0.0625 rate, 65.0 for a 0.65
    # LTV. `_pct_display` returns a float, so a round percentage renders
    # with a trailing .0 — the same as the GP co-invest field beside it,
    # and the number a browser posts back is identical either way.
    assert 'value="6.25"' in html
    assert 'value="65.0"' in html
    # A coverage ratio passes through untouched.
    assert 'value="1.25"' in html
    # The two conventions that have no second implemented value get no
    # field — a dropdown whose other option crashes the run is a trap.
    assert 'name="wf_accrual_base"' not in html
    assert 'name="wf_catch_up"' not in html


@pytest.mark.django_db
def test_a_carried_forward_unimplemented_convention_does_not_500_the_page():
    """`WaterfallTerms` splits its refusals across TWO exception types: a
    bad NUMBER raises ValueError, an unimplemented CONVENTION raises
    NotImplementedError. The three conventions this form deliberately
    does not expose are exactly the NotImplementedError ones, and exactly
    the ones that reach the prefill by being carried forward from a
    stored override — so catching only ValueError made the fallback a lie
    for the single case most likely to hit it (review finding).
    """
    from webapp.forms import AssumptionsForm, build_initial
    from webapp.models import Deal

    deal = Deal.objects.create(
        deal_id="lev-notimpl", property_name="NotImpl",
        assumption_overrides={"waterfall_terms": {"catch_up": True}})

    # Prefill falls back to the raw merge instead of raising.
    initial = build_initial(deal)
    assert initial["wf_pref_rate"] == pytest.approx(8.0)

    # And the same split is honoured on the way in, so a convention that
    # reached cleaned_data becomes a form error rather than a 500.
    form = AssumptionsForm(data={})
    assert form.is_valid(), form.errors
    form._validate_levered_terms({"wf_pref_compounding": "quarterly"})
    assert "Waterfall terms" in str(form.non_field_errors())


# ── Item E4 — the levered max offer reaches all three surfaces ───────


def _solved_levered_offer():
    from model.solver import solve_max_price_levered
    return solve_max_price_levered(adjusted_ttm_noi=300_000)


def test_the_levered_max_offer_reaches_the_memo(tmp_path, mock_cim_data,
                                                oracle_a):
    from docx import Document

    from output.memo_writer import generate_memo

    offer = _solved_levered_offer()
    # Section 6 returns early without scenarios, and the engine never
    # produces a levered max offer without them either — so a fixture
    # pairing an offer with no scenarios tests a state that cannot exist.
    path = generate_memo(
        property_name="E4 Memo", cim_data=mock_cim_data, gate_results=[],
        market_analysis={}, physical_analysis={}, financial_analysis={},
        rent_analysis={},
        scenario_results=_levered_run(oracle_a)["scenario_results"],
        value_add={}, risk_analysis={}, max_offer={},
        levered_max_offer=offer, output_dir=str(tmp_path))
    text = "\n".join(p.text for p in Document(path).paragraphs)

    assert "LP NET IRR" in text
    assert f"{offer['max_price']:,.0f}" in text
    # The stack at that price, not just the price — a max offer nobody can
    # finance is not an offer.
    assert f"{offer['senior_debt']:,.0f}" in text
    assert f"{offer['total_equity']:,.0f}" in text
    # Same rule as every other LP net figure: the stamp travels with it.
    assert "pref" in text.lower()


def test_the_memo_omits_the_levered_max_offer_when_no_loan_was_priced(
        tmp_path, mock_cim_data, oracle_a):
    """A block of N/A reads as a failed calculation rather than an absent
    one — the same contract `_add_levered_returns` already keeps.

    Scenarios ARE supplied, so section 6 renders in full and the absence
    of the levered paragraph is the thing under test rather than an early
    return hiding it.
    """
    from docx import Document

    from output.memo_writer import generate_memo

    path = generate_memo(
        property_name="E4 None", cim_data=mock_cim_data, gate_results=[],
        market_analysis={}, physical_analysis={}, financial_analysis={},
        rent_analysis={},
        scenario_results=_levered_run(oracle_a)["scenario_results"],
        value_add={}, risk_analysis={}, max_offer={},
        levered_max_offer=None, output_dir=str(tmp_path))
    text = "\n".join(p.text for p in Document(path).paragraphs)
    assert "6. Valuation & Returns" in text     # section 6 really rendered
    assert "LP NET IRR" not in text


def test_the_levered_max_offer_reaches_the_workbook(tmp_path, mock_cim_data):
    from openpyxl import load_workbook

    from output.excel_writer import generate_excel

    offer = _solved_levered_offer()
    path = generate_excel(
        property_name="E4 Model", cim_data=mock_cim_data,
        financial_analysis={}, scenario_results={}, sensitivity={},
        max_offer={}, levered_max_offer=offer, output_dir=str(tmp_path))
    ws = load_workbook(path)["Max Offer"]
    cells = {ws.cell(row=r, column=1).value: ws.cell(row=r, column=2).value
             for r in range(1, ws.max_row + 1)}

    assert cells["Maximum Purchase Price"] == pytest.approx(
        offer["max_price"], abs=CENT)
    assert cells["Achieved LP Net IRR"] == pytest.approx(offer["lp_net_irr"])
    assert cells["Senior Debt at Max Price"] == pytest.approx(
        offer["senior_debt"], abs=CENT)
    # Both basis definitions, because subtracting debt from the wrong one
    # gives the wrong equity (CLAUDE.md key design decision 3).
    assert cells["Total Basis (excl. financing)"] == pytest.approx(
        offer["total_basis"], abs=CENT)
    assert cells["Total Uses (incl. financing)"] == pytest.approx(
        offer["total_uses"], abs=CENT)
    # The unlevered header must still say it is unlevered — two prices on
    # one sheet solved to different bars.
    headers = {ws.cell(row=r, column=1).value for r in range(1, ws.max_row + 1)}
    assert "Maximum Offer Price Derivation (Unlevered)" in headers
    assert "Maximum Offer Price Derivation (Levered — LP Net)" in headers


def test_the_workbook_omits_the_levered_max_offer_when_absent(tmp_path,
                                                              mock_cim_data):
    from openpyxl import load_workbook

    from output.excel_writer import generate_excel

    path = generate_excel(
        property_name="E4 No Offer", cim_data=mock_cim_data,
        financial_analysis={}, scenario_results={}, sensitivity={},
        max_offer={}, levered_max_offer=None, output_dir=str(tmp_path))
    ws = load_workbook(path)["Max Offer"]
    headers = {ws.cell(row=r, column=1).value for r in range(1, ws.max_row + 1)}
    assert "Maximum Offer Price Derivation (Levered — LP Net)" not in headers


@pytest.mark.django_db
def test_the_returns_tab_renders_the_levered_max_offer_card():
    """Render it, do not trust the context dict — Django resolves a
    missing variable to the empty string, so a mis-named key renders a
    blank card and every context assertion still passes."""
    from django.template.loader import render_to_string

    from webapp.results import returns_context

    offer = _solved_levered_offer()
    html = render_to_string("webapp/_tab_returns.html",
                            returns_context({"levered_max_offer": offer}))
    assert "Max Offer — LP Net" in html
    assert f"${offer['max_price']:,.0f}" in html
    # The target is on the card because it is NOT the unlevered target.
    assert "15.0%" in html
    # `coerced_region` is ordinary and must not raise a badge on its own.
    assert "CHECK EXIT CAP" not in html


@pytest.mark.django_db
def test_the_returns_tab_omits_the_card_on_runs_stored_before_e4():
    from django.template.loader import render_to_string

    from webapp.results import returns_context

    html = render_to_string("webapp/_tab_returns.html", returns_context({}))
    assert "Max Offer — LP Net" not in html
    assert "Max Offer — Static" in html      # the unlevered card is untouched


@pytest.mark.django_db
def test_the_returns_tab_flags_an_observed_monotonicity_inversion():
    """The badge is driven by `monotonicity_warning` only. If it were
    driven by `coerced_region` it would fire on most deals and be
    ignored."""
    from django.template.loader import render_to_string

    from webapp.results import returns_context

    offer = {**_solved_levered_offer(),
             "monotonicity_warning": "LP net IRR RISES with price ..."}
    html = render_to_string("webapp/_tab_returns.html",
                            returns_context({"levered_max_offer": offer}))
    assert "CHECK EXIT CAP" in html
