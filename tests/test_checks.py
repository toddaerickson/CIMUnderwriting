"""Model error-check register (analysis/checks.py).

Every check gets a passing case, a failing case, and a boundary case at its
own tolerance — the boundary is the assertion that matters, since a check
whose threshold is off by an epsilon still passes a pass/fail pair.
"""

import pytest

from analysis import checks as C


def _run(check_id, **kw):
    """Evaluate exactly one check against a CheckInput built from kw."""
    results = C.run_checks(C.CheckInput(**kw), only={check_id})
    assert len(results) == 1
    return results[0]


# ── 1. Income identity (blocking) ───────────────────────────────────

INCOME = {"ttm_total_revenue": 560_000.0, "ttm_total_expenses": 220_000.0}


def test_income_identity_passes_on_consistent_triple():
    r = _run("income_identity", **INCOME, ttm_noi=340_000.0)
    assert r.status == C.PASS
    assert r.severity == C.BLOCKING


def test_income_identity_fails_beyond_tolerance():
    r = _run("income_identity", **INCOME, ttm_noi=400_000.0)
    assert r.status == C.FAIL
    assert r.blocks
    # Wording webapp.forms raises verbatim — see test_web_deals.py.
    assert "off by $60,000" in r.message
    assert r.values["delta"] == -60_000.0


def test_income_identity_boundary_at_tolerance_passes():
    # tolerance = max($1k, 1% × $560k) = $5,600 — exactly at it passes.
    r = _run("income_identity", **INCOME, ttm_noi=340_000.0 - 5_600.0)
    assert r.status == C.PASS
    assert r.values["tolerance"] == 5_600.0

    r = _run("income_identity", **INCOME, ttm_noi=340_000.0 - 5_600.01)
    assert r.status == C.FAIL


def test_income_identity_tolerance_floor_is_one_thousand():
    # 1% of $10k revenue is $100, but the floor holds it at $1,000.
    assert C.noi_recon_tolerance(10_000.0) == 1_000.0
    assert C.noi_recon_tolerance(560_000.0) == 5_600.0


def test_income_identity_skips_when_incomplete():
    r = _run("income_identity", ttm_total_revenue=560_000.0)
    assert r.status == C.SKIPPED
    assert not r.blocks


# ── 4. Occupancy sanity (blocking) ──────────────────────────────────

def test_occupancy_passes_with_normal_spread():
    r = _run("occupancy_sanity", physical_occupancy=0.92,
             economic_occupancy=0.80)
    assert r.status == C.PASS


def test_occupancy_fails_when_economic_exceeds_physical():
    r = _run("occupancy_sanity", physical_occupancy=0.80,
             economic_occupancy=0.92)
    assert r.status == C.FAIL
    assert r.blocks
    assert "exceeds physical" in r.message


def test_occupancy_boundary_equal_passes():
    r = _run("occupancy_sanity", physical_occupancy=0.88,
             economic_occupancy=0.88)
    assert r.status == C.PASS


def test_occupancy_boundary_full_occupancy_passes():
    r = _run("occupancy_sanity", physical_occupancy=1.0,
             economic_occupancy=1.0)
    assert r.status == C.PASS


def test_occupancy_fails_when_out_of_range():
    # 92 instead of 0.92 — the classic percent-scale slip.
    r = _run("occupancy_sanity", physical_occupancy=92.0)
    assert r.status == C.FAIL
    assert "outside 0–100%" in r.message


def test_occupancy_skips_when_absent():
    assert _run("occupancy_sanity").status == C.SKIPPED


# ── 5. EGR ≤ GPR (blocking) ─────────────────────────────────────────

def test_egr_le_gpr_passes():
    r = _run("egr_le_gpr", ttm_gpr=800_000.0, ttm_egr=720_000.0)
    assert r.status == C.PASS


def test_egr_le_gpr_fails():
    r = _run("egr_le_gpr", ttm_gpr=720_000.0, ttm_egr=800_000.0)
    assert r.status == C.FAIL
    assert r.blocks


def test_egr_le_gpr_boundary_equal_passes():
    r = _run("egr_le_gpr", ttm_gpr=800_000.0, ttm_egr=800_000.0)
    assert r.status == C.PASS
    # $1 of float noise is absorbed; $2 is not.
    assert _run("egr_le_gpr", ttm_gpr=800_000.0,
                ttm_egr=800_001.0).status == C.PASS
    assert _run("egr_le_gpr", ttm_gpr=800_000.0,
                ttm_egr=800_002.0).status == C.FAIL


def test_egr_le_gpr_skips_when_absent():
    assert _run("egr_le_gpr", ttm_gpr=800_000.0).status == C.SKIPPED


# ── 2. Unit mix SF vs NRSF (advisory) ───────────────────────────────

def _mix(count, sf, rate=100.0):
    return ({"count": count, "sf": sf, "rate": rate},)


def test_unit_mix_sf_passes():
    r = _run("unit_mix_sf", nrsf=50_000.0, unit_mix=_mix(500, 100.0))
    assert r.status == C.PASS
    assert r.severity == C.ADVISORY


def test_unit_mix_sf_fails_and_never_blocks():
    # 400 × 100 = 40,000 SF against a stated 50,000 — a partial extraction.
    r = _run("unit_mix_sf", nrsf=50_000.0, unit_mix=_mix(400, 100.0))
    assert r.status == C.FAIL
    assert not r.blocks          # advisory by deviation — see the plan doc
    assert "20.0%" in r.message


def test_unit_mix_sf_boundary_at_two_percent():
    # tolerance = 2% × 50,000 = 1,000 SF
    r = _run("unit_mix_sf", nrsf=50_000.0, unit_mix=_mix(490, 100.0))
    assert r.status == C.PASS    # off by exactly 1,000
    r = _run("unit_mix_sf", nrsf=50_000.0, unit_mix=_mix(489, 100.0))
    assert r.status == C.FAIL    # off by 1,100


def test_unit_mix_sf_skips_without_a_mix():
    assert _run("unit_mix_sf", nrsf=50_000.0).status == C.SKIPPED


# ── 3. Unit mix rents vs GPR (advisory) ─────────────────────────────

def test_unit_mix_gpr_passes():
    # 500 units × $100/mo × 12 = $600,000
    r = _run("unit_mix_gpr", ttm_gpr=600_000.0, unit_mix=_mix(500, 100.0))
    assert r.status == C.PASS


def test_unit_mix_gpr_fails():
    r = _run("unit_mix_gpr", ttm_gpr=800_000.0, unit_mix=_mix(500, 100.0))
    assert r.status == C.FAIL
    assert "vacancy and concessions do not explain" in r.message


def test_unit_mix_gpr_boundary_at_three_percent():
    # tolerance = 3% × 600,000 = 18,000
    r = _run("unit_mix_gpr", ttm_gpr=618_000.0, unit_mix=_mix(500, 100.0))
    assert r.status == C.PASS
    r = _run("unit_mix_gpr", ttm_gpr=619_000.0, unit_mix=_mix(500, 100.0))
    assert r.status == C.FAIL


# ── 6/7. OpEx bands (advisory) ──────────────────────────────────────

def test_opex_ratio_band_passes_inside_band():
    r = _run("opex_ratio_band", opex_revenue_ratio=0.45)
    assert r.status == C.PASS


def test_opex_ratio_band_fails_below_band():
    r = _run("opex_ratio_band", opex_revenue_ratio=0.20)
    assert r.status == C.FAIL
    assert "below" in r.message


def test_opex_ratio_band_boundary_at_both_edges():
    assert _run("opex_ratio_band", opex_revenue_ratio=0.35).status == C.PASS
    assert _run("opex_ratio_band", opex_revenue_ratio=0.55).status == C.PASS
    assert _run("opex_ratio_band", opex_revenue_ratio=0.3499).status == C.FAIL
    assert _run("opex_ratio_band", opex_revenue_ratio=0.5501).status == C.FAIL


def test_opex_ratio_band_derives_ratio_when_not_precomputed():
    # The form path has the income triple but no expense_ratio_check.
    r = _run("opex_ratio_band", ttm_total_expenses=252_000.0,
             ttm_total_revenue=560_000.0)
    assert r.status == C.PASS
    assert r.values["opex_revenue_ratio"] == pytest.approx(0.45)


def test_opex_ratio_band_honours_a_caller_supplied_band():
    r = _run("opex_ratio_band", opex_revenue_ratio=0.30,
             benchmarks={"opex_revenue_ratio": (0.25, 0.60)})
    assert r.status == C.PASS


def test_opex_per_nrsf_band_pass_fail_boundary():
    assert _run("opex_per_nrsf_band", opex_per_nrsf=4.00).status == C.PASS
    assert _run("opex_per_nrsf_band", opex_per_nrsf=1.00).status == C.FAIL
    assert _run("opex_per_nrsf_band", opex_per_nrsf=3.00).status == C.PASS
    assert _run("opex_per_nrsf_band", opex_per_nrsf=5.50).status == C.PASS
    assert _run("opex_per_nrsf_band", opex_per_nrsf=2.99).status == C.FAIL


def test_opex_per_nrsf_band_derives_from_totals():
    r = _run("opex_per_nrsf_band", ttm_total_expenses=200_000.0,
             nrsf=50_000.0)
    assert r.values["opex_per_nrsf"] == pytest.approx(4.0)
    assert r.status == C.PASS


# ── 8. Expense line floors (advisory, loud) ─────────────────────────

def _line(category, per_nrsf, low=1.20, high=2.50):
    return {"category": category, "benchmark_key": "property_tax",
            "per_nrsf": per_nrsf, "benchmark_range": (low, high)}


def test_expense_line_floor_passes_when_every_line_is_real():
    r = _run("expense_line_floor",
             expense_lines=(_line("Property Taxes", 1.80),
                            _line("Insurance", 0.20, 0.12, 0.25)))
    assert r.status == C.PASS


def test_expense_line_floor_catches_the_dollar_property_tax():
    # $1 of tax on 50,000 NRSF is $0.00002/SF — the Abilene case.
    r = _run("expense_line_floor",
             expense_lines=(_line("Property Taxes", 0.00002),))
    assert r.status == C.FAIL
    assert "Property Taxes" in r.message
    assert "below half the benchmark floor" in r.message


def test_expense_line_floor_boundary_at_half_the_low():
    # low = $1.20 → half is $0.60. At the half it passes; below it fails.
    assert _run("expense_line_floor",
                expense_lines=(_line("Property Taxes", 0.60),)).status == C.PASS
    assert _run("expense_line_floor",
                expense_lines=(_line("Property Taxes", 0.59),)).status == C.FAIL


def test_expense_line_floor_reports_zero_and_missing_separately():
    r = _run("expense_line_floor",
             expense_lines=(_line("Payroll", 0.0, 0.30, 0.60),
                            _line("Utilities", None, 0.08, 0.18)))
    assert r.status == C.FAIL
    assert "zero: Payroll" in r.message
    assert "not stated in the CIM: Utilities" in r.message


def test_expense_line_floor_handles_the_percent_shaped_mgmt_fee_line():
    r = _run("expense_line_floor",
             expense_lines=({"category": "Management Fee",
                             "benchmark_key": "mgmt_fee_pct",
                             "cim_pct": 0.01,
                             "benchmark_range_pct": (0.03, 0.06)},))
    assert r.status == C.FAIL
    assert "Management Fee" in r.message


def test_expense_line_floor_skips_without_lines():
    assert _run("expense_line_floor").status == C.SKIPPED


# ── 9. Exit cap coercion (advisory) ─────────────────────────────────

def test_exit_cap_coercion_passes_when_nothing_was_raised():
    r = _run("exit_cap_coercion",
             scenarios={"base": {"exit_cap": 0.075, "exit_cap_coerced": False,
                                 "requested_exit_cap": 0.075}})
    assert r.status == C.PASS


def test_exit_cap_coercion_surfaces_the_silent_raise():
    r = _run("exit_cap_coercion",
             scenarios={"base": {"exit_cap": 0.085, "exit_cap_coerced": True,
                                 "requested_exit_cap": 0.065}})
    assert r.status == C.FAIL
    assert "6.5%" in r.message and "8.5%" in r.message
    assert r.values["base"]["requested_exit_cap"] == 0.065


def test_exit_cap_coercion_skips_without_scenarios():
    assert _run("exit_cap_coercion").status == C.SKIPPED


# ── 10. Price vs replacement cost (advisory) ────────────────────────

def test_price_vs_replacement_passes_at_a_discount():
    r = _run("price_vs_replacement",
             price_vs_replacement={"comparable": True, "passes_gate": True,
                                   "asking_per_sf": 70.0,
                                   "replacement_per_sf": 95.0,
                                   "discount_to_replacement": 0.26})
    assert r.status == C.PASS


def test_price_vs_replacement_fails_at_a_premium():
    r = _run("price_vs_replacement",
             price_vs_replacement={"comparable": True, "passes_gate": False,
                                   "asking_per_sf": 120.0,
                                   "replacement_per_sf": 95.0,
                                   "discount_to_replacement": -0.26})
    assert r.status == C.FAIL
    assert "premium" in r.message


def test_price_vs_replacement_skips_when_not_comparable():
    r = _run("price_vs_replacement",
             price_vs_replacement={"comparable": False})
    assert r.status == C.SKIPPED


# ── Registry mechanics ──────────────────────────────────────────────

def test_registry_ids_are_unique_and_severities_are_valid():
    assert len(set(C.CHECK_IDS)) == len(C.CHECK_IDS)
    assert all(s.severity in (C.BLOCKING, C.ADVISORY) for s in C.CHECKS)


def test_run_checks_evaluates_every_check_on_an_empty_input():
    results = C.run_checks(C.CheckInput())
    assert len(results) == len(C.CHECKS)
    # Nothing to look at → nothing claimed. No check may report a pass.
    assert {r.status for r in results} == {C.SKIPPED}
    assert C.blocking_failures(results) == []


def test_summarize_counts_each_bucket():
    results = C.run_checks(C.CheckInput(
        ttm_total_revenue=560_000.0, ttm_total_expenses=220_000.0,
        ttm_noi=400_000.0,                       # blocking fail
        opex_revenue_ratio=0.10,                 # advisory fail
        physical_occupancy=0.92, economic_occupancy=0.80))   # pass
    s = C.summarize(results)
    assert s["total"] == len(C.CHECKS)
    assert s["blocking_failed"] == 1
    assert s["advisory_failed"] >= 1
    assert s["passed"] >= 1
    assert s["total"] == s["passed"] + s["failed"] + s["skipped"]


def test_round_trip_through_dicts_is_lossless():
    results = C.run_checks(C.CheckInput(**INCOME, ttm_noi=400_000.0))
    assert C.from_dicts(C.to_dicts(results)) == results


def test_from_dicts_tolerates_an_older_stored_shape():
    restored = C.from_dicts([{"id": "income_identity", "label": "x",
                              "severity": C.BLOCKING, "status": C.FAIL,
                              "message": "m", "values": {},
                              "source": "s", "unexpected_future_key": 1}])
    assert restored[0].id == "income_identity"
    assert restored[0].blocks


def test_to_dicts_is_json_serializable():
    import json
    results = C.run_checks(C.CheckInput(
        **INCOME, ttm_noi=400_000.0, nrsf=50_000.0,
        unit_mix=_mix(400, 100.0),
        expense_lines=(_line("Property Taxes", 0.00002),),
        scenarios={"base": {"exit_cap": 0.085, "exit_cap_coerced": True,
                            "requested_exit_cap": 0.065}}))
    json.dumps(C.to_dicts(results))    # must not raise


# ── input_from_cim adapter ──────────────────────────────────────────

def test_input_from_cim_maps_unit_types_and_analysis_outputs():
    from extract.parser import CIMData, UnitType

    cim = CIMData()
    cim.nrsf = 50_000.0
    cim.ttm_gpr = 600_000.0
    cim.ttm_egr = 540_000.0
    cim.physical_occupancy = 0.90
    cim.economic_occupancy = 0.82
    cim.unit_mix = [UnitType(size_label="10x10", count=500, sf=100.0,
                             rate=100.0)]
    fin = {
        "expense_analysis": {"lines": [_line("Property Taxes", 1.80)]},
        "expense_ratio_check": {"opex_revenue_ratio": 0.44,
                                "opex_per_nrsf": 4.10,
                                "benchmark_opex_range": (3.10, 5.60),
                                "benchmark_ratio_range": (0.35, 0.55)},
    }
    phys = {"price_vs_replacement": {"comparable": True, "passes_gate": True,
                                     "asking_per_sf": 70.0,
                                     "replacement_per_sf": 95.0,
                                     "discount_to_replacement": 0.26}}

    inp = C.input_from_cim(cim, fin, phys, {"base": {"exit_cap": 0.075}})

    assert inp.unit_mix == ({"count": 500, "sf": 100.0, "rate": 100.0},)
    assert inp.benchmarks["total_opex"] == (3.10, 5.60)   # state-adjusted
    assert inp.opex_per_nrsf == 4.10
    assert inp.price_vs_replacement["passes_gate"] is True

    results = {r.id: r for r in C.run_checks(inp)}
    assert results["unit_mix_sf"].status == C.PASS
    assert results["unit_mix_gpr"].status == C.PASS
    assert results["occupancy_sanity"].status == C.PASS
    assert results["price_vs_replacement"].status == C.PASS


def test_input_from_cim_survives_a_bare_cim():
    from extract.parser import CIMData

    inp = C.input_from_cim(CIMData())
    assert inp.unit_mix == ()
    assert inp.benchmarks is None
    assert C.blocking_failures(C.run_checks(inp)) == []


# ── The $1 property-tax line, end to end ────────────────────────────
# Acceptance criterion from the scope contract: a run with a $1 property-tax
# line must produce check 8 in the memo, the Excel Checks sheet and the
# results page. The first two are asserted here against the real writers;
# the results page is asserted in tests/test_web_runs.py.

def _dollar_tax_cim():
    """A CIM whose property-tax line reads $1 for the year — the Abilene
    case. Oklahoma, so the $/SF benchmark path runs rather than the Texas
    income-capitalization formula."""
    from extract.parser import CIMData, FinancialLine

    cim = CIMData()
    cim.property_name = "Dollar Tax Storage"
    cim.city, cim.state = "Tulsa", "OK"
    cim.nrsf = 50_000.0
    cim.total_units = 400
    cim.asking_price = 5_000_000.0
    cim.physical_occupancy = 0.90
    cim.economic_occupancy = 0.82
    cim.ttm_gpr = 600_000.0
    cim.ttm_egr = 550_000.0
    cim.ttm_total_revenue = 560_000.0
    cim.ttm_total_expenses = 220_000.0
    cim.ttm_noi = 340_000.0
    cim.mgmt_fee_pct = 0.05
    cim.expense_lines = [
        FinancialLine(label="Property Taxes", t12=1.0),      # the bad line
        FinancialLine(label="Insurance", t12=9_000.0),
        FinancialLine(label="Utilities", t12=6_000.0),
        FinancialLine(label="Repairs & Maintenance", t12=15_000.0),
        FinancialLine(label="Advertising", t12=5_000.0),
        FinancialLine(label="Payroll", t12=22_000.0),
        FinancialLine(label="General & Administrative", t12=7_000.0),
        FinancialLine(label="Replacement Reserve", t12=10_000.0),
    ]
    return cim


def _dollar_tax_register():
    from analysis.financials import analyze_financials
    from analysis.physical import analyze_physical

    cim = _dollar_tax_cim()
    fin = analyze_financials(cim)
    phys = analyze_physical(cim)
    return cim, fin, phys, C.run_checks(C.input_from_cim(cim, fin, phys))


def test_dollar_property_tax_is_flagged_by_the_register():
    _, _, _, results = _dollar_tax_register()
    by_id = {r.id: r for r in results}
    assert by_id["expense_line_floor"].status == C.FAIL
    assert "Property Taxes" in by_id["expense_line_floor"].message
    # ... and it does not block: the analyst is told, not stopped.
    assert C.blocking_failures(results) == []


def test_dollar_property_tax_reaches_the_excel_checks_sheet(tmp_path):
    from openpyxl import load_workbook
    from output.excel_writer import generate_excel

    cim, fin, _, results = _dollar_tax_register()
    path = generate_excel(
        property_name=cim.property_name, cim_data=cim,
        financial_analysis=fin, scenario_results={}, sensitivity={},
        max_offer={}, checks=C.to_dicts(results), output_dir=str(tmp_path))

    wb = load_workbook(path)
    assert "Checks" in wb.sheetnames
    text = "\n".join(str(c.value) for row in wb["Checks"].iter_rows()
                     for c in row if c.value)
    assert "Property Taxes" in text
    assert "FAIL" in text
    # The whole register is written, not only the findings.
    assert "SKIPPED" in text or "PASS" in text
    for spec in C.CHECKS:
        assert spec.label in text


def test_dollar_property_tax_reaches_the_memo(tmp_path):
    from docx import Document
    from output.memo_writer import generate_memo

    cim, fin, phys, results = _dollar_tax_register()
    path = generate_memo(
        property_name=cim.property_name, cim_data=cim, gate_results=[],
        market_analysis={}, physical_analysis=phys, financial_analysis=fin,
        rent_analysis={}, scenario_results={}, value_add={}, risk_analysis={},
        max_offer={}, checks=C.to_dicts(results), output_dir=str(tmp_path))

    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    text += "\n".join(c.text for t in doc.tables for r in t.rows
                      for c in r.cells)
    assert "Model Checks" in text
    assert "Property Taxes" in text


def test_memo_and_excel_are_unchanged_when_no_checks_are_supplied(tmp_path):
    """The CLI and any older caller pass no register — the writers must
    still produce their files, without an empty Checks sheet."""
    from openpyxl import load_workbook
    from output.excel_writer import generate_excel

    cim, fin, _, _ = _dollar_tax_register()
    path = generate_excel(
        property_name=cim.property_name, cim_data=cim,
        financial_analysis=fin, scenario_results={}, sensitivity={},
        max_offer={}, output_dir=str(tmp_path))
    assert "Checks" not in load_workbook(path).sheetnames
