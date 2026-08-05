"""Item D — Sources & Uses + capital stack, with item H's CapEx basis.

Three kinds of test, same discipline as tests/test_transaction_costs.py:

1. **No-op proofs.** The defaults (`reserve=0`, `capex_basis="amount"`) must
   reproduce the pre-item-D numbers exactly, which is what makes this a
   safe change to every published IRR rather than a silent re-baselining.
   Asserted at an absolute 1e-9, not a relative tolerance.
2. **Hand-computed oracles.** Built from the same degenerate flat case the
   item B suite uses, where the right answer is derivable on paper:

       ttm_noi 1,000,000 · price 10,000,000 · capex 0 · expense_ratio 0.50
       zero growth · zero costs · exit cap 0.10

   → NOI is 1,000,000 flat, exit value is 10,000,000, and with a reserve R
     you invest 10,000,000 + R to collect 5,000,000 of NOI plus a
     10,000,000 sale. MOIC is therefore exactly 15,000,000 / (10,000,000+R).
3. **Identity checks.** Uses = Sources = the DCF basis, in every scenario,
   for every reserve — the invariant the whole item exists to establish.
"""

import numpy_financial as npf
import pytest

from analysis import checks
from analysis.valuation import (project_cash_flows, resolve_exit_cap,
                                resolve_market_cap, run_scenarios)
from model.debt import DebtTerms
from model.returns_model import (BASIS_AMOUNT, BASIS_PCT_PRICE, BASIS_PER_SF,
                                 BASIS_PER_UNIT, build_returns_model,
                                 build_sources_uses, resolve_capital_amount,
                                 resolve_capital_structure)
from model.solver import solve_max_price
from registry import ScenarioType

NO_COSTS = {"acquisition_closing_pct": 0.0, "disposition_cost_pct": 0.0}
FLAT_PARAMS = {
    "yr1_noi_bump": 0.0, "stabilized_occ": 0.88,
    "rev_cagr_yr1_3": 0.0, "rev_cagr_yr4_5": 0.0,
    "exp_growth": 0.0,
}
#: The flat case buys and sells at a 10 cap. It is now passed as an
#: argument rather than carried in the params dict: the exit cap is
#: derived from a market anchor and is no longer a scenario parameter.
FLAT_EXIT_CAP = 0.10
FLAT_NOI = 1_000_000
FLAT_PRICE = 10_000_000

#: The solver round-trip tests re-run the solved price forward through the
#: projection, so both sides must price the exit identically. Pinning the
#: anchor here is what makes that an equality rather than a coincidence of
#: both happening to read the same config table.
PINNED_ANCHOR = resolve_market_cap(market_cap=0.0625)
PINNED_BASE_EXIT_CAP = resolve_exit_cap(
    PINNED_ANCHOR["market_cap"], ScenarioType.BASE)["exit_cap"]


def flat(**kw):
    args = dict(ttm_noi=FLAT_NOI, price=FLAT_PRICE, capex=0,
                params=FLAT_PARAMS, hold_years=5, expense_ratio=0.50,
                costs=NO_COSTS, exit_cap=FLAT_EXIT_CAP)
    args.update(kw)
    return project_cash_flows(**args)


# ── 1. The defaults are a no-op ──────────────────────────────────────

@pytest.mark.parametrize("scen", [s.value for s in ScenarioType])
def test_zero_reserve_reproduces_the_pre_item_d_projection(scen):
    """Passing reserve=0 explicitly and omitting it entirely must produce
    bit-identical results, or this change re-baselined every stored run."""
    common = dict(adjusted_ttm_noi=300_000, asking_price=4_000_000,
                  nrsf=50_000, capex=250_000, expense_ratio=0.42)
    without = run_scenarios(**common)[ScenarioType(scen)]
    with_zero = run_scenarios(**common, reserve=0.0)[ScenarioType(scen)]
    for key in ("irr", "moic", "yield_on_cost", "total_basis", "exit_value"):
        assert without[key] == pytest.approx(with_zero[key], abs=1e-9)


def test_zero_reserve_solver_reproduces_the_item_b_pin():
    """The pin captured at commit bb15311 in tests/test_transaction_costs.

    Re-baselined 3,625,000 -> 3,609,375 by item T Category 3, which gave
    the three solvers one shared bracket at the wider dear cap. The
    reasoning is recorded once, beside `_SOLVER_PINS` in
    tests/test_transaction_costs.py — this is the same case with a reserve
    argument, and what it is here to prove is that reserve=0 changes
    nothing, not what the price is.
    """
    result = solve_max_price(adjusted_ttm_noi=300_000, capex=200_000,
                             transaction_costs=NO_COSTS)
    assert result["max_price"] == pytest.approx(3_609_375.0, abs=1e-6)
    assert result["reserve"] == 0.0


def test_capex_amount_basis_is_the_historical_reading():
    assert resolve_capital_amount(250_000, BASIS_AMOUNT) == 250_000
    assert resolve_capital_amount(250_000, None) == 250_000


# ── 2. resolve_capital_amount ────────────────────────────────────────

def test_per_sf_and_amount_agree_on_the_same_money():
    """The contract's acceptance criterion: the same figure entered two
    ways must produce one number."""
    assert resolve_capital_amount(0.50, BASIS_PER_SF, nrsf=50_000) == 25_000.0
    assert (resolve_capital_amount(0.50, BASIS_PER_SF, nrsf=50_000)
            == resolve_capital_amount(25_000, BASIS_AMOUNT))


def test_per_unit_basis():
    assert resolve_capital_amount(125, BASIS_PER_UNIT, units=400) == 50_000.0


def test_pct_price_basis_reads_a_decimal_fraction():
    """Canonical units, matching every other percentage in the model
    layer — webapp.forms owns the whole-number boundary."""
    assert resolve_capital_amount(0.02, BASIS_PCT_PRICE,
                                  price=5_000_000) == 100_000.0


def test_blank_amount_is_zero_not_an_error():
    assert resolve_capital_amount(None, BASIS_PER_SF, nrsf=50_000) == 0.0
    assert resolve_capital_amount("", BASIS_AMOUNT) == 0.0


@pytest.mark.parametrize("basis,kwargs", [
    (BASIS_PER_SF, {"nrsf": None}),
    (BASIS_PER_SF, {"nrsf": 0}),
    (BASIS_PER_UNIT, {"units": None}),
    (BASIS_PCT_PRICE, {"price": 0}),
])
def test_missing_driver_contributes_zero_rather_than_a_wrong_magnitude(
        basis, kwargs, caplog):
    """`$0.50/SF` with no NRSF is not fifty cents. Returning the raw
    number would put a figure three orders of magnitude off into the
    capital stack; returning 0 says "could not compute" and says it in
    the log."""
    with caplog.at_level("WARNING"):
        assert resolve_capital_amount(0.50, basis, **kwargs) == 0.0
    assert "missing or zero" in caplog.text


def test_unknown_basis_falls_back_to_dollars_with_a_warning(caplog):
    """A stored override written by a future version must not take down a
    run on an older one."""
    with caplog.at_level("WARNING"):
        assert resolve_capital_amount(1_234, "per_acre") == 1_234
    assert "unknown capital basis" in caplog.text


# ── 3. resolve_capital_structure ─────────────────────────────────────

def test_capital_structure_defaults_come_from_config():
    import config as cfg

    resolved = resolve_capital_structure(None)
    assert resolved["capex_basis"] == cfg.DEFAULT_CAPEX_BASIS
    assert resolved["operating_reserve"] == cfg.DEFAULT_OPERATING_RESERVE
    assert resolved["gp_coinvest_pct"] == cfg.GP_COINVEST_PCT


def test_omitting_a_key_means_default_not_zero():
    """Same contract as resolve_transaction_costs — a silent zero here is
    how a reserve or a co-invest share disappears from a deal."""
    import config as cfg

    resolved = resolve_capital_structure({"operating_reserve": 75_000})
    assert resolved["operating_reserve"] == 75_000
    assert resolved["gp_coinvest_pct"] == cfg.GP_COINVEST_PCT


def test_an_unsupported_basis_is_replaced_not_raised(caplog):
    with caplog.at_level("WARNING"):
        resolved = resolve_capital_structure(
            {"operating_reserve_basis": BASIS_PCT_PRICE})   # not offered
    assert resolved["operating_reserve_basis"] == BASIS_AMOUNT


# ── 4. build_sources_uses ────────────────────────────────────────────

def test_uses_sum_every_line_and_sources_equal_them():
    su = build_sources_uses(price=5_000_000, capex=250_000,
                            acquisition_cost=50_000, reserve=75_000)
    assert su["total_uses"] == 5_375_000
    assert su["total_sources"] == pytest.approx(5_375_000)
    assert su["balanced"] is True
    assert su["delta"] == pytest.approx(0.0, abs=1e-9)


def test_all_equity_stack_when_there_is_no_debt():
    """The contract's acceptance criterion: debt at 0 → 100% equity."""
    su = build_sources_uses(price=5_000_000, capex=250_000,
                            acquisition_cost=50_000, gp_coinvest_pct=0.10)
    assert su["senior_debt"] == 0.0
    assert su["total_equity"] == su["total_uses"]
    assert su["gp_equity"] == pytest.approx(530_000)
    assert su["lp_equity"] == pytest.approx(4_770_000)
    assert su["gp_equity"] + su["lp_equity"] == pytest.approx(su["total_uses"])


def test_debt_displaces_equity_rather_than_adding_to_uses():
    """The property that makes this schema correct on the day item E1
    sizes a loan: uses are unchanged, equity shrinks by the loan."""
    no_debt = build_sources_uses(price=5_000_000, acquisition_cost=50_000)
    levered = build_sources_uses(price=5_000_000, acquisition_cost=50_000,
                                 senior_debt=3_000_000)
    assert levered["total_uses"] == no_debt["total_uses"]
    assert levered["total_equity"] == no_debt["total_equity"] - 3_000_000
    assert levered["total_sources"] == pytest.approx(levered["total_uses"])
    assert levered["ltv"] == pytest.approx(3_000_000 / 5_050_000)


def test_gp_coinvest_pct_defaults_to_config_read_at_call_time(monkeypatch):
    import config as cfg

    monkeypatch.setattr(cfg, "GP_COINVEST_PCT", 0.25)
    su = build_sources_uses(price=1_000_000)
    assert su["gp_coinvest_pct"] == 0.25
    assert su["gp_equity"] == pytest.approx(250_000)


# ── 5. The invariant: Uses = Sources = the DCF basis ─────────────────

@pytest.mark.parametrize("reserve", [0, 75_000, 250_000])
@pytest.mark.parametrize("capex", [0, 250_000])
def test_sources_uses_ties_to_total_basis_in_every_scenario(reserve, capex):
    """Item D's invariant, restated for item E3a.

    `build_returns_model` now sizes a loan on every deal, so Total Uses
    exceeds the DCF basis by exactly the origination fee — financing
    costs are a use of funds and are deliberately NOT in `total_basis`,
    which is what keeps the unlevered screen financing-free. The tie is
    still exact; it just carries the term."""
    model = build_returns_model(
        adjusted_ttm_noi=400_000, asking_price=5_000_000, nrsf=50_000,
        capex=capex, expense_ratio=0.42, reserve=reserve)
    su = model["sources_uses"]
    financing = su["financing_costs"]
    assert financing > 0, "the config default charges a point at close"
    for name, scen in model["scenarios"].items():
        assert scen["total_basis"] + financing == pytest.approx(
            su["total_uses"], abs=0.01), name
        assert scen["total_basis"] + financing == pytest.approx(
            su["total_sources"], abs=0.01), name


def test_equity_equals_the_dcf_year_zero_outflow_plus_financing_less_debt():
    """The other half of the contract's acceptance criterion, restated for
    item E3a: the loan DISPLACES equity, so the year-zero equity cheque is
    the DCF outflow plus the financing cost, less the loan proceeds."""
    model = build_returns_model(
        adjusted_ttm_noi=400_000, asking_price=5_000_000, nrsf=50_000,
        capex=250_000, expense_ratio=0.42, reserve=75_000)
    base = model["scenarios"][ScenarioType.BASE]
    su = model["sources_uses"]
    assert su["total_equity"] == pytest.approx(
        -base["cash_flows"][0] + su["financing_costs"] - su["senior_debt"],
        abs=0.01)
    # And with no debt sized, it collapses to item D's original identity.
    unlevered = build_returns_model(
        adjusted_ttm_noi=400_000, asking_price=5_000_000, nrsf=50_000,
        capex=250_000, expense_ratio=0.42, reserve=75_000,
        debt_terms=DebtTerms(rate=0.0625, max_ltv=0.0, min_dscr=0.0,
                             min_debt_yield=0.0, orig_fee_pct=0.01))
    base_u = unlevered["scenarios"][ScenarioType.BASE]
    assert unlevered["sources_uses"]["total_equity"] == pytest.approx(
        -base_u["cash_flows"][0], abs=0.01)


def test_reserve_is_entered_as_dollars_or_per_sf_to_the_same_effect():
    """Reserve entered as $/NRSF and as $ produce identical results."""
    per_sf = resolve_capital_amount(1.50, BASIS_PER_SF, nrsf=50_000)
    dollars = resolve_capital_amount(75_000, BASIS_AMOUNT)
    assert per_sf == dollars
    a = build_returns_model(adjusted_ttm_noi=400_000, asking_price=5_000_000,
                            nrsf=50_000, reserve=per_sf)
    b = build_returns_model(adjusted_ttm_noi=400_000, asking_price=5_000_000,
                            nrsf=50_000, reserve=dollars)
    assert (a["scenarios"][ScenarioType.BASE]["irr"]
            == b["scenarios"][ScenarioType.BASE]["irr"])


# ── 6. The reserve inside the DCF (hand-computed) ────────────────────

def test_reserve_enters_the_basis_and_is_not_released_at_exit():
    p = flat(reserve=500_000)
    assert p["total_basis"] == 10_500_000
    assert p["cash_flows"][0] == -10_500_000
    # Nothing is added back at exit: the final flow is NOI + the SAME net
    # sale proceeds as the no-reserve case.
    assert p["cash_flows"][-1] == pytest.approx(11_000_000)
    # MOIC = (5 × 1M NOI + 10M sale) / 10.5M, on paper.
    assert p["moic"] == pytest.approx(15_000_000 / 10_500_000)
    assert p["yield_on_cost"] == pytest.approx(1_000_000 / 10_500_000)


def test_reserve_lowers_the_irr_and_the_irr_still_zeroes_its_npv():
    """IRR has no closed form here, so assert the thing that defines it."""
    without = flat()
    with_reserve = flat(reserve=500_000)
    assert with_reserve["irr"] < without["irr"]
    assert npf.npv(with_reserve["irr"], with_reserve["cash_flows"]) == \
        pytest.approx(0, abs=1e-6)


# ── 7. The solver ────────────────────────────────────────────────────

def test_solved_price_with_a_reserve_reproduces_the_target_irr():
    import config as cfg

    solved = solve_max_price(adjusted_ttm_noi=300_000, capex=0,
                             expense_ratio=0.42, reserve=200_000,
                             market_cap=PINNED_ANCHOR)
    forward = project_cash_flows(
        ttm_noi=300_000, price=solved["max_price"], capex=0,
        params=cfg.SCENARIO_DEFAULTS[ScenarioType.BASE],
        expense_ratio=0.42, reserve=200_000,
        exit_cap=PINNED_BASE_EXIT_CAP)
    assert forward["irr"] == pytest.approx(solved["target_irr"], abs=0.001)
    assert solved["total_basis"] == pytest.approx(forward["total_basis"])


def test_pct_of_price_capex_scales_inside_the_bisection():
    """The defect item B fixed for closing costs, applied to CapEx: hold
    it at the asking-price dollars while solving a different price and the
    answer describes a deal nobody is buying."""
    import config as cfg

    pct = 0.05
    solved = solve_max_price(adjusted_ttm_noi=300_000, capex=0,
                             expense_ratio=0.42, capex_pct_of_price=pct,
                             market_cap=PINNED_ANCHOR)
    assert solved["capex"] == pytest.approx(solved["max_price"] * pct)
    forward = project_cash_flows(
        ttm_noi=300_000, price=solved["max_price"],
        capex=solved["max_price"] * pct,
        params=cfg.SCENARIO_DEFAULTS[ScenarioType.BASE], expense_ratio=0.42,
        exit_cap=PINNED_BASE_EXIT_CAP)
    assert forward["irr"] == pytest.approx(solved["target_irr"], abs=0.001)
    # And it is a different answer from freezing CapEx at the asking price.
    frozen = solve_max_price(adjusted_ttm_noi=300_000,
                             capex=4_000_000 * pct, expense_ratio=0.42,
                             market_cap=PINNED_ANCHOR)
    assert solved["max_price"] != pytest.approx(frozen["max_price"], abs=1.0)


# ── 8. Check 11 ──────────────────────────────────────────────────────

def _check(inp) -> checks.CheckResult:
    return next(r for r in checks.run_checks(inp, only={"sources_uses_ties"}))


def test_check_passes_when_the_stack_ties():
    model = build_returns_model(
        adjusted_ttm_noi=400_000, asking_price=5_000_000, nrsf=50_000,
        capex=250_000, expense_ratio=0.42, reserve=75_000)
    r = _check(checks.CheckInput(scenarios=model["scenarios"],
                                sources_uses=model["sources_uses"]))
    assert r.status == checks.PASS
    assert r.severity == checks.BLOCKING


def test_check_fails_when_uses_and_the_dcf_basis_disagree():
    model = build_returns_model(
        adjusted_ttm_noi=400_000, asking_price=5_000_000, nrsf=50_000,
        expense_ratio=0.42)
    broken = dict(model["sources_uses"])
    broken["total_uses"] = broken["total_uses"] + 1_000
    broken["total_sources"] = broken["total_uses"]
    r = _check(checks.CheckInput(scenarios=model["scenarios"],
                                sources_uses=broken))
    assert r.status == checks.FAIL
    assert "DCF basis" in r.message


# ── 9. The identity item E3a moved ───────────────────────────────────
# E1 measured that financing costs break the old `Uses == total_basis`
# tie by exactly the origination fee. The operator's call on 2026-08-01
# was to move the IDENTITY rather than the projection, so the unlevered
# screen stays financing-free. These pin the new identity from both
# directions.

def _levered_model(financing_costs=60_000.0, senior_debt=6_000_000.0):
    """A returns model whose stack carries real debt and a real fee."""
    model = build_returns_model(
        adjusted_ttm_noi=400_000, asking_price=5_000_000, nrsf=50_000,
        capex=250_000, expense_ratio=0.42, reserve=75_000)
    base = next(s for s in model["scenarios"].values() if isinstance(s, dict))
    model["sources_uses"] = build_sources_uses(
        price=5_000_000, capex=250_000,
        acquisition_cost=base["acquisition_cost"], reserve=75_000,
        financing_costs=financing_costs, senior_debt=senior_debt)
    return model


def test_check_passes_when_uses_exceed_the_basis_by_exactly_the_fee():
    model = _levered_model()
    su = model["sources_uses"]
    base = next(s for s in model["scenarios"].values() if isinstance(s, dict))
    assert su["financing_costs"] == pytest.approx(60_000.0)
    assert su["total_uses"] == pytest.approx(base["total_basis"] + 60_000.0)
    r = _check(checks.CheckInput(scenarios=model["scenarios"],
                                sources_uses=su))
    assert r.status == checks.PASS
    assert r.severity == checks.BLOCKING
    # The message names both halves, so a reader can see WHY Uses is
    # bigger than the basis rather than wondering whether it is a bug.
    assert "financing costs" in r.message


def test_check_still_fails_when_the_gap_is_not_the_financing_cost():
    """The check is no weaker for carrying a financing term: a stack that
    is off by anything OTHER than the fee still fails loudly."""
    model = _levered_model()
    broken = dict(model["sources_uses"])
    broken["total_uses"] += 1_000
    broken["total_sources"] = broken["total_uses"]
    r = _check(checks.CheckInput(scenarios=model["scenarios"],
                                sources_uses=broken))
    assert r.status == checks.FAIL


def test_check_fails_when_the_fee_is_dropped_from_the_stack():
    """The other direction: a stack reporting no financing cost while its
    Uses still contain one is the E1 failure mode, and it still FAILs."""
    model = _levered_model()
    broken = dict(model["sources_uses"])
    broken["financing_costs"] = 0.0
    r = _check(checks.CheckInput(scenarios=model["scenarios"],
                                sources_uses=broken))
    assert r.status == checks.FAIL


def test_debt_displaces_equity_rather_than_inflating_uses():
    """Item D's promise, now that E3a supplies a real loan."""
    levered = _levered_model()["sources_uses"]
    unlevered = _levered_model(financing_costs=0.0,
                               senior_debt=0.0)["sources_uses"]
    assert levered["total_uses"] - unlevered["total_uses"] == pytest.approx(
        60_000.0)
    assert levered["total_equity"] == pytest.approx(
        unlevered["total_equity"] + 60_000.0 - 6_000_000.0)
    assert levered["total_uses"] == pytest.approx(levered["total_sources"])


def test_check_fails_when_scenarios_disagree_on_the_basis():
    model = build_returns_model(
        adjusted_ttm_noi=400_000, asking_price=5_000_000, nrsf=50_000,
        expense_ratio=0.42)
    scenarios = {k: dict(v) for k, v in model["scenarios"].items()}
    scenarios[ScenarioType.BULL]["total_basis"] += 5_000
    r = _check(checks.CheckInput(scenarios=scenarios,
                                sources_uses=model["sources_uses"]))
    assert r.status == checks.FAIL
    assert "disagree" in r.message


def test_check_is_skipped_where_there_is_no_capital_stack():
    """The assumptions form cannot see a DCF. `skipped` is the honest
    answer there — rendering it as a pass would claim we looked."""
    r = _check(checks.CheckInput(ttm_noi=340_000))
    assert r.status == checks.SKIPPED


def test_a_boundary_penny_is_inside_the_tolerance():
    model = build_returns_model(
        adjusted_ttm_noi=400_000, asking_price=5_000_000, nrsf=50_000,
        expense_ratio=0.42)
    su = dict(model["sources_uses"])
    su["total_uses"] += checks.SOURCES_USES_TOLERANCE_ABS
    su["total_sources"] = su["total_uses"]
    assert _check(checks.CheckInput(scenarios=model["scenarios"],
                                    sources_uses=su)).status == checks.PASS


# ── 9. Value-add engine shares the basis ─────────────────────────────

def test_value_add_basis_includes_the_reserve(mock_cim_data,
                                              base_financial_analysis):
    from model.value_add_model import run_value_add_scenarios

    plain = run_value_add_scenarios(
        cim_data=mock_cim_data, financial_analysis=base_financial_analysis,
        asking_price=5_000_000, capex=0)
    reserved = run_value_add_scenarios(
        cim_data=mock_cim_data, financial_analysis=base_financial_analysis,
        asking_price=5_000_000, capex=0, reserve=100_000)
    base = ScenarioType.BASE
    assert (reserved[base]["total_basis"]
            == pytest.approx(plain[base]["total_basis"] + 100_000))
    assert reserved[base]["irr"] < plain[base]["irr"]


# ── 10. Output surfaces ──────────────────────────────────────────────

def _stack():
    return build_sources_uses(price=5_000_000, capex=250_000,
                              acquisition_cost=50_000, reserve=75_000,
                              gp_coinvest_pct=0.10)


def test_the_capital_stack_reaches_the_excel_workbook(tmp_path,
                                                      mock_cim_data):
    from openpyxl import load_workbook
    from output.excel_writer import generate_excel

    path = generate_excel(
        property_name="SU Test", cim_data=mock_cim_data,
        financial_analysis={}, scenario_results={}, sensitivity={},
        max_offer={}, sources_uses=_stack(), output_dir=str(tmp_path))

    wb = load_workbook(path)
    assert "Sources & Uses" in wb.sheetnames
    ws = wb["Sources & Uses"]
    cells = {str(c.value) for row in ws.iter_rows() for c in row if c.value}
    assert "Purchase Price" in cells
    assert "Operating Reserve" in cells
    assert "Total Uses" in cells
    assert "LP Equity" in cells
    assert "In balance" in cells
    values = [c.value for row in ws.iter_rows() for c in row
              if isinstance(c.value, (int, float))]
    assert 5_375_000 in values          # total uses
    assert 537_500 in values            # GP co-invest at 10%


def test_no_capital_stack_means_no_sheet(tmp_path, mock_cim_data):
    """The CLI path and any older caller pass nothing — the workbook must
    still build, without an empty sheet."""
    from openpyxl import load_workbook
    from output.excel_writer import generate_excel

    path = generate_excel(
        property_name="SU Test", cim_data=mock_cim_data,
        financial_analysis={}, scenario_results={}, sensitivity={},
        max_offer={}, output_dir=str(tmp_path))
    assert "Sources & Uses" not in load_workbook(path).sheetnames


def test_the_capital_stack_reaches_the_memo(tmp_path, mock_cim_data):
    from docx import Document
    from output.memo_writer import generate_memo

    path = generate_memo(
        property_name="SU Test", cim_data=mock_cim_data, gate_results=[],
        market_analysis={}, physical_analysis={}, financial_analysis={},
        rent_analysis={}, scenario_results={"base": {"noi_projection": [1]}},
        value_add={}, risk_analysis={}, max_offer={},
        sources_uses=_stack(), output_dir=str(tmp_path))

    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    text += "\n".join(c.text for t in doc.tables for r in t.rows
                      for c in r.cells)
    assert "Sources & Uses" in text
    assert "Operating Reserve" in text
    assert "$5,375,000" in text
    assert "Total equity required: $5,375,000" in text
    # Zero lines are PRINTED, not dropped: a reader has to be able to see
    # this is an all-equity underwrite rather than infer it.
    assert "Senior Debt" in text
    assert "Financing Costs" in text


def test_capital_context_zips_uses_and_sources_into_one_table():
    from webapp.results import capital_context

    ctx = capital_context({"sources_uses": _stack()})
    assert ctx["has_capital"] is True
    assert len(ctx["capital_rows"]) == 5        # uses is the longer side
    first = ctx["capital_rows"][0]
    assert first["use"]["label"] == "Purchase Price"
    assert first["use"]["amount"] == "$5,000,000"
    assert first["source"]["label"] == "Senior Debt"
    # Sources runs out first; the template renders a blank, not a crash.
    assert ctx["capital_rows"][3]["source"] is None
    assert ctx["total_equity"] == "$5,375,000"
    assert ctx["gp_coinvest_pct"] == "10%"


def test_capital_context_is_absent_for_a_run_that_predates_item_d():
    from webapp.results import capital_context

    assert capital_context({})["has_capital"] is False


def test_zero_lines_are_kept_but_marked_quiet():
    """The default all-equity deal has three $0 Uses rows. They stay —
    "no reserve was underwritten" is a statement — but they are marked so
    the page can render them without shouting."""
    from webapp.results import capital_context

    ctx = capital_context({"sources_uses": build_sources_uses(price=5_000_000)})
    by_label = {r["use"]["label"]: r["use"] for r in ctx["capital_rows"]
                if r["use"]}
    assert by_label["Operating Reserve"]["zero"] is True
    assert by_label["Purchase Price"]["zero"] is False
    assert by_label["Operating Reserve"]["amount"] == "$0"


def test_an_unbalanced_stack_is_reported_even_without_a_dcf():
    """`skipped` must not swallow a finding: no scenarios to reconcile
    against is a reason not to compare with the DCF, not a reason to stop
    reporting that the stack does not balance against itself."""
    su = dict(_stack())
    su["total_sources"] = su["total_uses"] - 500
    r = _check(checks.CheckInput(sources_uses=su))
    assert r.status == checks.FAIL
    assert "does not balance" in r.message


# ── 11. Review repairs ───────────────────────────────────────────────

def test_pct_of_price_capex_scales_across_the_sensitivity_price_axis():
    """The grid's row axis IS price. Holding CapEx at the asking-price
    dollars makes every row but the centre describe a deal whose CapEx did
    not move with its price — the same defect the solvers carry
    capex_pct_of_price to avoid (review finding)."""
    import config as cfg

    pct = 0.05
    price = 5_000_000
    common = dict(adjusted_ttm_noi=400_000, asking_price=price, nrsf=50_000,
                  expense_ratio=0.42)
    scaled = build_returns_model(capex=price * pct,
                                 capex_pct_of_price=pct, **common)
    frozen = build_returns_model(capex=price * pct, **common)

    # Centre cell is the asking price, so both agree there.
    assert (scaled["sensitivity"]["irr_grid"][4][4]
            == pytest.approx(frozen["sensitivity"]["irr_grid"][4][4], abs=1e-12))
    # The -10% price row must not be.
    assert scaled["sensitivity"]["irr_grid"][0][4] != pytest.approx(
        frozen["sensitivity"]["irr_grid"][0][4], abs=1e-9)
    # And the scaled cell must equal a hand-built projection at that price.
    cheap = scaled["sensitivity"]["price_values"][0]
    expected = project_cash_flows(
        ttm_noi=400_000, price=cheap, capex=cheap * pct,
        params=cfg.SCENARIO_DEFAULTS[ScenarioType.BASE],
        expense_ratio=0.42, coerce_exit_cap=False,
        exit_cap_override=scaled["sensitivity"]["cap_values"][4])["irr"]
    assert scaled["sensitivity"]["irr_grid"][0][4] == pytest.approx(
        expected, abs=1e-12)


def test_a_rate_that_resolves_to_zero_becomes_a_run_warning(mock_cim_data,
                                                            tmp_path,
                                                            monkeypatch):
    """A saved basis outlives the value it was validated against: a
    re-extraction that loses NRSF turns a valid $/SF CapEx into $0 on the
    next run. A quiet $0 line in the capital stack is an empty state
    hiding a real failure (review finding)."""
    # data.comp_db binds COMP_DB_PATH at import — redirect it or the run
    # writes into the repo's real comps database.
    monkeypatch.setattr("data.comp_db.COMP_DB_PATH", str(tmp_path / "c.db"))
    from engine import AnalysisResult, run_analysis

    mock_cim_data.nrsf = None
    mock_cim_data.capex_estimate = 0.50
    result = AnalysisResult(pdf_path=str(tmp_path / "none.pdf"))
    result.cim_data = mock_cim_data
    run_analysis(result, output_dir=str(tmp_path),
                 capital_structure={"capex_basis": "per_sf"})

    assert any("resolved to $0" in e for e in result.errors), result.errors
    assert any("CapEx" in e for e in result.errors)


def test_a_resolvable_rate_raises_no_warning(mock_cim_data, tmp_path,
                                             monkeypatch):
    monkeypatch.setattr("data.comp_db.COMP_DB_PATH", str(tmp_path / "c.db"))
    from engine import AnalysisResult, run_analysis

    mock_cim_data.capex_estimate = 0.50          # nrsf is 50,000
    result = AnalysisResult(pdf_path=str(tmp_path / "none.pdf"))
    result.cim_data = mock_cim_data
    run_analysis(result, output_dir=str(tmp_path),
                 capital_structure={"capex_basis": "per_sf"})

    assert not any("resolved to $0" in e for e in result.errors)
    assert result.sources_uses["uses"][2]["amount"] == 25_000.0


def test_unconverged_solver_is_flagged_for_the_returns_tab():
    from webapp.results import returns_context

    ctx = returns_context({"max_offer": {"max_price": 1_500_000.0,
                                         "converged": False}})
    assert ctx["max_offer_unconverged"] is True
    assert returns_context(
        {"max_offer": {"max_price": 3_900_000.0, "converged": True}}
    )["max_offer_unconverged"] is False
    # A run predating the flag has no `converged` key — a missing field is
    # not evidence of a failure, so it must not be flagged retroactively.
    assert returns_context(
        {"max_offer": {"max_price": 3_900_000.0}}
    )["max_offer_unconverged"] is False
    assert returns_context({})["max_offer_unconverged"] is False
