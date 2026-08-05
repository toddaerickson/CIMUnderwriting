"""Item T Category 1 — every moved key has ONE definition, and every
surface that restated it now follows it.

The characterization suite next door proves this PR changed nothing. That
is only half the claim. A literal moved into config and then read back
*wrongly* — the caller binding a stale copy, or reading a different key —
reproduces the snapshot perfectly, because the value is the same value.
Byte-for-byte green is exactly what a dead wire looks like.

So these tests move the config value and assert every surface moves WITH
it. That is the item's actual acceptance criterion ("a ConfigOverride
delta changes every output the audit found divergent"), and it is the
test that fails if someone later re-freezes one of these bindings at
import time.

Patching is done with `monkeypatch.setitem` on the live dict, which is
precisely what `webapp.services._merge_patch` does for a real
ConfigOverride row — same mutation, same in-place semantics, so a module
that survives this survives a settings override. Scalars have no such
mechanism (a module binding one by value cannot see a patch at all), so
`MGMT_FEE_TARGET_PCT` is asserted to be read through `cfg.` at call time.
"""

import math
from unittest import mock

import pytest

import config as cfg
from analysis.financials import analyze_financials
from analysis.market import analyze_market
from analysis.risks import identify_risks
from analysis.value_add import identify_value_add
from engine import AnalysisResult, run_analysis
from registry import ScenarioType
from tests.test_characterization import stabilized_deal


@pytest.fixture(autouse=True)
def isolated_comp_db(tmp_path, monkeypatch):
    """Same reason as the characterization suite: `engine.run_analysis`
    writes every completed run into the comp DB and
    `analysis.financials._get_benchmarks` reads it back, so an unisolated
    run here would change the benchmarks a LATER test resolves. Patch the
    module attribute, not `config.COMP_DB_PATH` — `data.comp_db` binds it
    at import."""
    import data.comp_db as comp_db_module
    monkeypatch.setattr(comp_db_module, "COMP_DB_PATH",
                        str(tmp_path / "comps.db"))


# ── GATES["max_noi_step_up"] — the flag AND the sentence ─────────────

def _risks_for(cim_data):
    return identify_risks(cim_data, gate_results=[], financial_analysis={},
                          scenario_results={})["risks"]


def _risk(risks, name):
    return next((r for r in risks if r["risk"] == name), None)


def test_noi_step_up_flag_follows_the_gate(mock_cim_data, monkeypatch):
    """340k TTM -> 360k Yr1 is a 5.9% step-up: under the 15% default, over
    a 5% gate. The literal was `0.15` in `analysis/risks.py`."""
    assert _risk(_risks_for(mock_cim_data), "Aggressive CIM pro forma") is None

    monkeypatch.setitem(cfg.GATES, "max_noi_step_up", 0.05)
    assert _risk(_risks_for(mock_cim_data), "Aggressive CIM pro forma")


def test_noi_step_up_sentence_quotes_the_gate_it_used(mock_cim_data,
                                                      monkeypatch):
    """The risk text said "exceeds 15% step-up threshold" as a string
    literal, so an overridden gate produced a sentence describing a
    threshold the code had not applied."""
    monkeypatch.setitem(cfg.GATES, "max_noi_step_up", 0.05)
    risk = _risk(_risks_for(mock_cim_data), "Aggressive CIM pro forma")

    assert "exceeds 5% step-up threshold" in risk["description"]
    assert "15%" not in risk["description"]


# ── GATES["population_3mi"] — four surfaces, one number ───────────────

def test_population_gate_drives_every_market_surface(mock_cim_data,
                                                     monkeypatch):
    """`analysis/market.py` restated the 50,000 gate three times (the
    adequacy flag, the narrative sentence, the demand negative) and
    `analysis/risks.py` a fourth (severity).

    60,000 people: above the 50,000 default, below a 65,000 override, and
    below the separate 75,000 "dense trade area" tier that would otherwise
    short-circuit the demand branch (see the test below).
    """
    mock_cim_data.population_3mi = 60_000

    market = analyze_market(mock_cim_data)
    assert market["demographics"]["pop_3mi_adequate"] is True
    assert "Adequate density" in market["demographics"]["pop_narrative"]
    assert not any("Thin trade area" in n
                   for n in market["demand_drivers"]["negatives"])

    monkeypatch.setitem(cfg.GATES, "population_3mi", 65_000)
    market = analyze_market(mock_cim_data)

    assert market["demographics"]["pop_3mi_adequate"] is False
    assert "Thin trade area" in market["demographics"]["pop_narrative"]
    assert any("Thin trade area" in n
               for n in market["demand_drivers"]["negatives"])


def test_the_density_tier_can_no_longer_shadow_the_population_gate(
        mock_cim_data, monkeypatch):
    """This test previously pinned the OPPOSITE, as a known Category 1
    limitation: a gate raised above the 75,000 tier left a deal reporting
    as a POSITIVE demand driver while the adequacy flag beside it said
    False. Making the tier settings-editable turned that from a fixed
    quirk into something a user could produce on purpose, so
    `_assess_demand` now takes `max(preferred_density, population_3mi)`.

    The two settings are independent, which means nothing stops someone
    setting the tier below the gate — the guard is that the OUTPUT stays
    coherent when they do, not that the inputs are policed.
    """
    mock_cim_data.population_3mi = 80_000
    monkeypatch.setitem(cfg.GATES, "population_3mi", 100_000)
    market = analyze_market(mock_cim_data)

    assert market["demographics"]["pop_3mi_adequate"] is False
    assert not any("Dense trade area" in p
                   for p in market["demand_drivers"]["positives"])


def test_a_tier_set_below_the_gate_cannot_contradict_it(mock_cim_data,
                                                        monkeypatch):
    """The same guard from the other direction: an operator lowering the
    preferred-density tier below the gate must not manufacture a deal that
    is simultaneously a demand positive and inadequate.

    BOTH faces are asserted here, and the second one is why: the first
    draft of this PR guarded `market.py` only, and `risks.py` — reading
    the same key — silently stopped raising "Limited trade area
    population" for a deal that fails the gate outright. Asserting the
    market face alone read as full coverage of a "one threshold, two
    faces" key while testing one of them.
    """
    mock_cim_data.population_3mi = 45_000
    monkeypatch.setitem(cfg.POPULATION_TIERS, "preferred_density", 40_000)
    market = analyze_market(mock_cim_data)

    assert market["demographics"]["pop_3mi_adequate"] is False
    assert not any("Dense trade area" in p
                   for p in market["demand_drivers"]["positives"])
    assert any("Thin trade area" in n
               for n in market["demand_drivers"]["negatives"])

    # The risk face: 45,000 sits ABOVE the lowered tier but BELOW the gate,
    # so an unguarded `pop < preferred_density` never fires.
    risk = _risk(_risks_for(mock_cim_data), "Limited trade area population")
    assert risk is not None
    assert risk["severity"] == "High"        # below the gate, not merely thin


# ── POPULATION_TIERS — narrative grading, settings-editable ──────────

def test_preferred_density_is_one_threshold_with_two_faces(mock_cim_data,
                                                           monkeypatch):
    """It was two separate 75,000 literals — `market.py`'s demand positive
    and `risks.py`'s "Limited trade area population" trigger — which could
    drift into a market that is neither a positive nor a risk, or both."""
    mock_cim_data.population_3mi = 80_000
    assert any("Dense trade area" in p for p in
               analyze_market(mock_cim_data)["demand_drivers"]["positives"])
    assert _risk(_risks_for(mock_cim_data),
                 "Limited trade area population") is None

    monkeypatch.setitem(cfg.POPULATION_TIERS, "preferred_density", 90_000)

    assert not any("Dense trade area" in p for p in
                   analyze_market(mock_cim_data)["demand_drivers"]["positives"])
    assert _risk(_risks_for(mock_cim_data), "Limited trade area population")


def test_strong_density_drives_the_top_narrative_tier(mock_cim_data,
                                                      monkeypatch):
    """`market.py`'s 100,000 "strong demand driver" tier."""
    mock_cim_data.population_3mi = 120_000
    assert "strong demand driver" in analyze_market(
        mock_cim_data)["demographics"]["pop_narrative"]

    monkeypatch.setitem(cfg.POPULATION_TIERS, "strong_density", 150_000)
    narrative = analyze_market(mock_cim_data)["demographics"]["pop_narrative"]

    assert "strong demand driver" not in narrative
    assert "Adequate density" in narrative


def test_population_tiers_are_settings_editable_end_to_end():
    """The registry is what the settings page renders and what
    `build_config_patch` validates against, so an unregistered key is
    accepted by the form and then silently skipped at run time — the exact
    "UI claims the override works" failure item T exists to kill."""
    from webapp.forms import (format_override_value, override_key_registry,
                              parse_override_value)
    from webapp.services import _PATCHED_DICTS, build_config_patch

    reg = override_key_registry()
    for key in ("POPULATION_TIERS.preferred_density",
                "POPULATION_TIERS.strong_density"):
        assert reg[key]["int"] is True and reg[key]["pct"] is False
        # Counts, not percentages: 90000 must round-trip as 90000, not 900.
        assert parse_override_value(key, "90,000") == 90_000
        assert format_override_value(key, 90_000) == "90000"

    # In _PATCHED_DICTS, so the patch actually reaches the analysis modules
    # that bound the dict at import.
    assert "POPULATION_TIERS" in _PATCHED_DICTS
    patch, _solver, skipped = build_config_patch(
        {"POPULATION_TIERS.preferred_density": 90_000})
    assert patch == {"POPULATION_TIERS": {"preferred_density": 90_000}}
    assert skipped == []


@pytest.mark.django_db
def test_a_stored_override_row_reaches_the_analysis(mock_cim_data):
    """The whole chain, from a row in the database to a changed narrative:
    ConfigOverride -> resolve -> build_config_patch -> _patched_config ->
    the module-level dict `analysis.market` bound at import.

    Every earlier link is unit-tested above; this is the one assertion
    that fails if any of them stops composing. The others would all still
    pass with a dict that no consumer reads.
    """
    from django.utils import timezone

    from webapp.models import ConfigOverride
    from webapp.services import (_patched_config, build_config_patch,
                                 resolve_config_overrides)

    mock_cim_data.population_3mi = 80_000
    assert any("Dense trade area" in p for p in
               analyze_market(mock_cim_data)["demand_drivers"]["positives"])

    ConfigOverride.objects.create(key="POPULATION_TIERS.preferred_density",
                                  value=90_000,
                                  effective_date=timezone.localdate())
    deltas = resolve_config_overrides("", timezone.localdate())
    patch, _solver, skipped = build_config_patch(deltas)
    assert skipped == []

    with _patched_config(patch):
        positives = analyze_market(mock_cim_data)["demand_drivers"]["positives"]
        risks = _risks_for(mock_cim_data)
    assert not any("Dense trade area" in p for p in positives)
    assert _risk(risks, "Limited trade area population")

    # And the patch is REVERTED on exit — a leaked mutation would silently
    # reprice every later deal in the same worker process.
    assert cfg.POPULATION_TIERS["preferred_density"] == 75_000


def test_population_narrative_quotes_the_gate_it_used(mock_cim_data,
                                                      monkeypatch):
    """"below 50,000 minimum" was hard-coded into the sentence, so an
    overridden gate printed the OLD number beside the new verdict."""
    monkeypatch.setitem(cfg.GATES, "population_3mi", 100_000)
    narrative = analyze_market(mock_cim_data)["demographics"]["pop_narrative"]

    assert "below 100,000 minimum" in narrative
    assert "50,000" not in narrative


def test_population_severity_follows_the_gate(mock_cim_data, monkeypatch):
    """`risks.py` grades a thin trade area Medium above the gate, High
    below it. At 60,000 the fixture is below the "preferred density"
    trigger either way, so only the SEVERITY should move."""
    mock_cim_data.population_3mi = 60_000
    assert _risk(_risks_for(mock_cim_data),
                 "Limited trade area population")["severity"] == "Medium"

    monkeypatch.setitem(cfg.GATES, "population_3mi", 65_000)
    assert _risk(_risks_for(mock_cim_data),
                 "Limited trade area population")["severity"] == "High"


def test_the_hhi_thresholds_are_not_the_population_gate(mock_cim_data,
                                                        monkeypatch):
    """The trap this move had to avoid: `median_hhi_3mi` is screened at
    50,000 too — dollars, not people. A sweep that replaced every `50_000`
    in `market.py` would have wired household income to the population
    gate, and nothing else in the suite would have noticed."""
    mock_cim_data.median_hhi_3mi = 55_000

    monkeypatch.setitem(cfg.GATES, "population_3mi", 100_000)
    demographics = analyze_market(mock_cim_data)["demographics"]

    assert demographics["hhi_adequate"] is True
    assert "Middle-income" in demographics["hhi_narrative"]


# ── MGMT_FEE_TARGET_PCT — one target, two modules, four sites ────────

def test_mgmt_fee_target_drives_the_financials_adjustment(mock_cim_data,
                                                          monkeypatch):
    """A CIM fee below the benchmark floor is adjusted UP to the target.
    `financials.py` restated it twice as a number and twice as text."""
    mock_cim_data.mgmt_fee_pct = 0.01
    monkeypatch.setattr(cfg, "MGMT_FEE_TARGET_PCT", 0.08)

    fin = analyze_financials(mock_cim_data)
    line = next(ln for ln in fin["expense_analysis"]["lines"]
                if ln["category"] == "Management Fee")
    egr = fin["income_summary"]["egr"]

    assert line["adjusted_pct"] == 0.08
    assert line["adjusted_value"] == pytest.approx(egr * 0.08)
    assert any("Adjusted to 8% of EGR" in a
               for a in fin["expense_analysis"]["adjustments"])


def test_mgmt_fee_target_drives_the_value_add_saving(mock_cim_data,
                                                     monkeypatch):
    """`value_add.py` sized the renegotiation saving off its own copy of
    the target, so the two modules could disagree about the same number.
    Both now resolve it through `resolve_mgmt_fee_target`.

    The fee is 8% — above the (3%, 6%) band, which is the case that used
    to crash `_expense_opportunities` outright (KeyError
    'benchmark_range'). Fixed in #41, so an above-band fee is now the
    natural fixture for "there is a real fee to renegotiate".
    """
    mock_cim_data.mgmt_fee_pct = 0.08
    fin = analyze_financials(mock_cim_data)
    egr = fin["income_summary"]["egr"]

    op = next(o for o in identify_value_add(
        mock_cim_data, fin)["expense_opportunities"]
        if o["category"] == "Management Fee Reduction")
    assert op["est_annual_impact"] == pytest.approx(
        egr * (0.08 - cfg.MGMT_FEE_TARGET_PCT))

    monkeypatch.setattr(cfg, "MGMT_FEE_TARGET_PCT", 0.04)
    op = next(o for o in identify_value_add(
        mock_cim_data, fin)["expense_opportunities"]
        if o["category"] == "Management Fee Reduction")

    assert op["est_annual_impact"] == pytest.approx(egr * (0.08 - 0.04))
    assert "to 4% of EGR" in op["description"]


def test_the_two_modules_cannot_resolve_different_targets(mock_cim_data):
    """The double-count this parameter could have introduced. If
    `analyze_financials` underwrote to 6% while `identify_value_add`
    credited a walk down to the config default, the same dollar would be
    counted twice — once as an expense the model already removed, once as
    upside still to come. Both read the ONE resolver, so a per-deal target
    handed to the engine reaches both or neither.
    """
    mock_cim_data.mgmt_fee_pct = 0.08
    fin = analyze_financials(mock_cim_data, mgmt_fee_target_pct=0.04)
    egr = fin["income_summary"]["egr"]

    op = next(o for o in identify_value_add(
        mock_cim_data, fin, mgmt_fee_target_pct=0.04)["expense_opportunities"]
        if o["category"] == "Management Fee Reduction")

    assert op["est_annual_impact"] == pytest.approx(egr * (0.08 - 0.04))
    assert "to 4% of EGR" in op["description"]


def test_a_zero_target_is_honoured_not_swallowed(mock_cim_data):
    """0.0 is a legitimate target — a self-managed property underwritten
    with no third-party fee — so the resolver keys on `is None`. A falsy
    check would silently substitute the 6% config default and quietly add
    an expense the analyst deliberately removed."""
    from analysis.financials import resolve_mgmt_fee_target

    assert resolve_mgmt_fee_target(0.0) == 0.0
    assert resolve_mgmt_fee_target(None) == cfg.MGMT_FEE_TARGET_PCT

    mock_cim_data.mgmt_fee_pct = None
    fin = analyze_financials(mock_cim_data, mgmt_fee_target_pct=0.0)
    line = next(ln for ln in fin["expense_analysis"]["lines"]
                if ln["category"] == "Management Fee")

    assert line["adjusted_value"] == 0.0


def test_an_omitted_fee_is_underwritten_at_the_target(mock_cim_data):
    """The highest-impact path, and the one the characterization net does
    NOT reach: a CIM that omits its management fee entirely.

    All three snapshot fixtures miss it — `stabilized` states 5%,
    `value_add` states 6%, and `thin` omits the fee but also has no EGR,
    so `elif egr:` never fires and no fee is assumed at all. So the
    default moving 5% -> 6% shows up in the snapshots only as a lost
    value-add opportunity, never as the NOI reduction it mainly is. This
    test is that missing coverage, asserted directly.
    """
    mock_cim_data.mgmt_fee_pct = None
    fin = analyze_financials(mock_cim_data)
    egr = fin["income_summary"]["egr"]
    line = next(ln for ln in fin["expense_analysis"]["lines"]
                if ln["category"] == "Management Fee")

    assert line["flag"] == "NOT FOUND"
    assert line["adjusted_value"] == pytest.approx(egr * 0.06)
    assert any("Assumed 6% of EGR" in a
               for a in fin["expense_analysis"]["adjustments"])


@pytest.mark.django_db
def test_the_mgmt_fee_target_round_trips_through_the_assumptions_form():
    """The percent-vs-decimal boundary, which is where a field of this
    shape usually breaks: the form takes a WHOLE number (4 meaning 4%)
    and every consumer wants a decimal. A one-way conversion stores 4.0
    as a 400% fee, or redisplays 0.04 in a box labelled "%".

    Also asserts the delta discipline the section uses: a field left at
    the config default writes NO key, so the deal stays on whatever the
    default later becomes rather than freezing today's value into every
    deal ever saved.
    """
    from django.http import QueryDict

    from webapp.forms import AssumptionsForm, build_initial, build_overrides
    from webapp.models import Deal

    deal = Deal.objects.create(deal_id="mf-rt", property_name="RT")

    form = AssumptionsForm(data={"mgmt_fee_target_pct": "4"})
    assert form.is_valid(), form.errors
    stored = build_overrides(form.cleaned_data, QueryDict(""), deal)
    assert stored["mgmt_fee_target_pct"] == pytest.approx(0.04)

    # Redisplay lands on the whole number the analyst typed, not 0.04.
    deal.assumption_overrides = stored
    deal.save()
    assert build_initial(deal)["mgmt_fee_target_pct"] == pytest.approx(4.0)

    # At the default: no key stored.
    at_default = AssumptionsForm(
        data={"mgmt_fee_target_pct": str(cfg.MGMT_FEE_TARGET_PCT * 100)})
    assert at_default.is_valid(), at_default.errors
    assert "mgmt_fee_target_pct" not in build_overrides(
        at_default.cleaned_data, QueryDict(""), deal)


def test_the_engine_accepts_a_per_deal_mgmt_fee_target():
    """`webapp.services` hands the stored override straight to
    `run_analysis`, so the parameter has to exist by that exact name.
    Pinned by signature rather than by a full run — the arithmetic is
    covered above; what breaks silently is a rename."""
    import inspect

    from engine import run_analysis

    params = inspect.signature(run_analysis).parameters
    assert "mgmt_fee_target_pct" in params
    assert params["mgmt_fee_target_pct"].default is None


@pytest.mark.django_db
def test_a_saved_mgmt_fee_target_actually_reaches_the_run(monkeypatch,
                                                          tmp_path, settings):
    """The link every other test in this file would pass without: the
    worker reading the stored override and handing it to the engine.

    Found by mutation — deleting the `mgmt_fee_target_pct=` line from
    `webapp/services.py`'s `run_analysis` call left the whole suite green.
    The analyst would type 4%, the page would save it, redisplay it, and
    the model would quietly underwrite at 6% — the exact "UI claims the
    override works and the model proves otherwise" failure item T exists
    to kill, reintroduced by the very PR adding the field.

    A 0.0 target is used deliberately: it is both a legitimate value and
    the one an `or`-style fallback would swallow.
    """
    from tests.test_web_runs import _make_extracted_deal, _start_run

    # `deals_dir` is a local fixture in three other test modules; inlined
    # here rather than copied a fourth time.
    deals_dir = tmp_path / "deals"
    deals_dir.mkdir()
    settings.CIM_DEALS_DIR = str(deals_dir)
    seen = {}

    def _fake(result, progress=None, output_dir=None, custom_scenarios=None,
              custom_va_scenarios=None, solver_target_irr=None, enrich=False,
              expense_line_overrides=None, hold_years=None,
              transaction_costs=None, capital_structure=None,
              market_cap_rate=None, market_cap=None,
              debt_terms=None, waterfall_terms=None, am_fee_pct=None,
              mgmt_fee_target_pct=None):
        seen["mgmt_fee_target_pct"] = mgmt_fee_target_pct
        result.gate_results = []
        result.gate_summary = {"passed": 0, "failed": 0, "tbd": 0, "total": 0,
                               "recommendation": "PURSUE",
                               "failed_gates": [], "tbd_gates": []}
        return result

    monkeypatch.setattr("webapp.services.run_analysis", _fake)
    deal = _make_extracted_deal(deals_dir)
    deal.assumption_overrides = {"mgmt_fee_target_pct": 0.0}
    deal.save()
    run = _start_run(deal)

    assert seen["mgmt_fee_target_pct"] == 0.0
    assert run.applied_overrides["assumptions"]["mgmt_fee_target_pct"] == 0.0


@pytest.mark.django_db
def test_a_run_on_the_default_still_stamps_the_target_it_used(monkeypatch,
                                                              tmp_path,
                                                              settings):
    """Stamped RESOLVED, not as a delta — the discipline `hold_years`,
    `transaction_costs` and the debt/waterfall blocks already follow.

    This PR moved the default 5% -> 6%. Without a resolved stamp, a run
    from before and a run from after — neither with a per-deal override —
    carry byte-identical `applied_overrides` while underwriting a
    management fee 100bp apart, which is 1.6% of adjusted NOI on a deal
    whose CIM omits the fee. A past run has to say what it ran under.
    """
    from tests.test_web_runs import _make_extracted_deal, _start_run

    deals_dir = tmp_path / "deals"
    deals_dir.mkdir()
    settings.CIM_DEALS_DIR = str(deals_dir)

    def _fake(result, *a, **kw):
        result.gate_results = []
        result.gate_summary = {"passed": 0, "failed": 0, "tbd": 0, "total": 0,
                               "recommendation": "PURSUE",
                               "failed_gates": [], "tbd_gates": []}
        return result

    monkeypatch.setattr("webapp.services.run_analysis", _fake)
    deal = _make_extracted_deal(deals_dir)
    deal.assumption_overrides = {}          # nothing set — pure default
    deal.save()
    run = _start_run(deal)

    assert (run.applied_overrides["assumptions"]["mgmt_fee_target_pct"]
            == cfg.MGMT_FEE_TARGET_PCT)


def test_the_target_is_the_top_of_the_benchmark_band():
    """Replaces this file's earlier
    `test_mgmt_fee_target_is_not_the_benchmark_midpoint`, which pinned the
    old 5%.

    6% is the band's HIGH end, and that is the whole underwriting
    argument: a CIM omitting its management fee is the common case, and
    the conservative read of an omission is the most expensive credible
    number, not a comfortable middle one. Operator's call 2026-08-04.

    The midpoint assertion survives for the same reason it was written:
    a derived `(low + high) / 2` looks like a tidy-up and silently
    re-underwrites every deal with an understated or missing fee.
    """
    low, high = cfg.EXPENSE_BENCHMARKS["mgmt_fee_pct"]

    assert cfg.MGMT_FEE_TARGET_PCT == high
    assert cfg.MGMT_FEE_TARGET_PCT != (low + high) / 2


# ── SOLVER_TARGET_IRR — the setting, resolved at call time ──────────

def test_the_solver_target_is_read_at_call_time_not_frozen_at_import():
    """Both unlevered solvers took `target_irr: float = SOLVER_TARGET_IRR`
    — a default ARGUMENT, which Python evaluates once at import. The value
    froze at whatever config held on first import of `model.solver`, so
    the only way a settings change reached the solver was by being
    threaded in as a parameter (which the web app does, and the CLI does
    not).

    Patching `cfg` must now move the answer with no parameter passed at
    all. That is what "read at call time" means, and a default-argument
    binding cannot do it.
    """
    from model.solver import resolve_target_irr, solve_max_price

    assert resolve_target_irr(None) == cfg.SOLVER_TARGET_IRR

    base = solve_max_price(adjusted_ttm_noi=400_000.0, capex=0,
                           expense_ratio=0.40)
    assert base["target_irr"] == pytest.approx(0.10)

    with mock.patch.object(cfg, "SOLVER_TARGET_IRR", 0.14):
        raised = solve_max_price(adjusted_ttm_noi=400_000.0, capex=0,
                                 expense_ratio=0.40)

    assert raised["target_irr"] == pytest.approx(0.14)
    # A higher bar buys less: the price has to FALL.
    assert raised["max_price"] < base["max_price"]


def test_a_zero_solver_target_is_a_question_not_a_blank():
    """`engine.py` guarded with `if solver_target_irr:` and
    `webapp/services.py` resolved with `or`, so a 0% target — the price at
    which the deal merely breaks even — was silently answered at 10%
    instead. The form accepts it (`min_value=0`), so it was reachable."""
    from model.solver import resolve_target_irr, solve_max_price

    assert resolve_target_irr(0.0) == 0.0

    out = solve_max_price(adjusted_ttm_noi=400_000.0, capex=0,
                          expense_ratio=0.40, target_irr=0.0)

    assert out["target_irr"] == 0.0
    assert out["max_price"] > solve_max_price(
        adjusted_ttm_noi=400_000.0, capex=0,
        expense_ratio=0.40)["max_price"]


def test_a_zero_target_survives_the_engine_too(tmp_path, monkeypatch):
    """The resolver alone is not enough — `engine.py` decides whether to
    FORWARD the target at all, and its `if solver_target_irr:` dropped a
    0.0 before the solver ever saw it. Found by mutation: reverting that
    one guard left the suite green, because every other test in this file
    exercises the solver directly or passes a non-zero target.
    """
    import data.comp_db as comp_db_module
    monkeypatch.setattr(comp_db_module, "COMP_DB_PATH",
                        str(tmp_path / "comps.db"))

    result = AnalysisResult(pdf_path="z.pdf", cim_data=stabilized_deal())
    run_analysis(result, output_dir=str(tmp_path), solver_target_irr=0.0)

    assert result.max_offer["target_irr"] == 0.0


@pytest.mark.django_db
def test_a_settings_row_moves_a_real_max_offer(tmp_path, monkeypatch):
    """The assertion the existing coverage stops short of: every current
    test of this path either inspects a signature or watches a
    MONKEYPATCHED `run_analysis` record the kwarg it was handed. None runs
    the real pipeline and checks the number moved.

    ConfigOverride row -> resolve -> build_config_patch -> the worker's
    `solver_irr` -> `run_analysis` -> `solve_max_price`.
    """
    from django.utils import timezone

    from engine import AnalysisResult, run_analysis
    from tests.test_characterization import stabilized_deal
    from webapp.models import ConfigOverride
    from webapp.services import build_config_patch, resolve_config_overrides

    import data.comp_db as comp_db_module
    monkeypatch.setattr(comp_db_module, "COMP_DB_PATH",
                        str(tmp_path / "comps.db"))

    def _run(**kw):
        out = tmp_path / f"o{len(list(tmp_path.iterdir()))}"
        out.mkdir()
        r = AnalysisResult(pdf_path="s.pdf", cim_data=stabilized_deal())
        run_analysis(r, output_dir=str(out), **kw)
        return r.max_offer

    base = _run()

    ConfigOverride.objects.create(key="SOLVER_TARGET_IRR", value=0.14,
                                  effective_date=timezone.localdate())
    deltas = resolve_config_overrides("", timezone.localdate())
    _patch, solver_irr, skipped = build_config_patch(deltas)
    assert skipped == [] and solver_irr == 0.14

    overridden = _run(solver_target_irr=solver_irr)

    assert base["target_irr"] == pytest.approx(0.10)
    assert overridden["target_irr"] == pytest.approx(0.14)
    assert overridden["max_price"] < base["max_price"]


# ── The documents — the surfaces the result object cannot prove ──────

def _documents_for(tmp_path, **kwargs):
    """One real end-to-end run, read back off disk.

    The workbook's grid colors and the max-offer captions are Category 1
    targets that leave `AnalysisResult`'s own fields untouched, so nothing
    short of opening the rendered files can show them moving.
    """
    from tests.test_characterization import _docx_content, _xlsx_content

    result = AnalysisResult(pdf_path="rt.pdf", cim_data=stabilized_deal())
    run_analysis(result, output_dir=str(tmp_path), **kwargs)
    return (_docx_content(result.memo_path),
            _xlsx_content(result.excel_path))


def _recommendation(base_irr: float) -> str:
    """Section 10 rendered on its own, with every gate passing.

    Called directly rather than through a pipeline run because the
    threshold sentence is only REACHABLE when no gate failed and none is
    TBD — a failed gate short-circuits to DECLINE and never reads the IRR
    at all. Both characterization fixtures that reach the memo have a
    failing gate, so an end-to-end run cannot exercise this branch.
    """
    from docx import Document

    from output.memo_writer import _add_section_10

    doc = Document()
    _add_section_10(doc, gate_results=[{"gate": 1, "name": "g",
                                        "result": "PASS"}],
                    scenario_results={"base": {"irr": base_irr}},
                    max_offer={}, risk_analysis={}, cim_data=None)
    return "\n".join(p.text for p in doc.paragraphs)


def test_memo_recommendation_threshold_follows_the_gate(monkeypatch):
    """"meets the 10% IRR target" was a string literal sitting beside a
    `>= 0.10` comparison — two copies of the gate, neither reading it. A
    12% deal clears the default and misses a 15% gate, so BOTH the
    verdict and the number in the sentence have to move."""
    text = _recommendation(0.12)
    assert "RECOMMENDATION: PURSUE" in text
    assert "meet the 10% IRR target" in text

    monkeypatch.setitem(cfg.GATES, "min_irr_5yr", 0.15)
    text = _recommendation(0.12)

    assert "RECOMMENDATION: PURSUE CONTINGENT ON" in text
    assert "base case IRR is below 15% target" in text
    assert "10%" not in text


def test_sensitivity_legend_follows_the_gate_and_the_strong_band(
        tmp_path, monkeypatch):
    """The workbook legend spelled out "≥ 12% IRR" / "10-12% IRR" /
    "< 10% IRR" as three literals describing thresholds held as two
    others. All five are now one pair."""
    monkeypatch.setitem(cfg.GATES, "min_irr_5yr", 0.08)
    monkeypatch.setattr(cfg, "IRR_STRONG_THRESHOLD", 0.20)
    _, workbook = _documents_for(tmp_path)

    legend = [c["v"] for c in workbook["Sensitivity"].values()
              if isinstance(c["v"], str) and "IRR" in c["v"]]

    assert "  ≥ 20% IRR" in legend
    assert "  8-20% IRR" in legend
    assert "  < 8% IRR" in legend


def test_sensitivity_colors_follow_the_gate(tmp_path, monkeypatch):
    """The fills, not just the caption. A gate above every IRR in the
    grid must leave no green and no yellow cell in it — proof the
    comparison reads config and not the old `0.12` / `0.10`."""
    monkeypatch.setitem(cfg.GATES, "min_irr_5yr", 0.90)
    monkeypatch.setattr(cfg, "IRR_STRONG_THRESHOLD", 0.95)
    _, workbook = _documents_for(tmp_path)

    fills = {c.get("fill") for c in workbook["Sensitivity"].values()
             if isinstance(c["v"], float)}

    assert "00C6EFCE" not in fills          # green — none can clear 95%
    assert "00FFEB9C" not in fills          # yellow — none can clear 90%
    assert "00FFC7CE" in fills              # red — every cell


def test_max_offer_captions_quote_the_target_actually_solved_for(tmp_path):
    """The captions must name the target the SOLVER used, not a constant
    that happens to match it today.

    Driven through the per-deal `solver_target_irr` — the route an
    analyst actually uses, and the one that was the ONLY route back when
    `solve_max_price` bound `SOLVER_TARGET_IRR` as a default argument
    frozen at import. That binding is gone (see
    `test_the_solver_target_is_read_at_call_time_not_frozen_at_import`
    above, which covers the config route), and this test kept passing
    across the change unaltered, which is what it was written to do.
    """
    memo, workbook = _documents_for(tmp_path, solver_target_irr=0.14)
    text = "\n".join(memo["paragraphs"])

    assert "Maximum Offer Price (for 14% Base Case IRR)" in text
    assert "At a target 14% base case unlevered IRR" in text
    assert not any("10% Base Case IRR" in p for p in memo["paragraphs"])

    captions = [c["v"] for c in workbook["Max Offer"].values()
                if isinstance(c["v"], str) and "Max Price" in c["v"]]
    assert captions and all("10%" not in c for c in captions)


def test_max_offer_caption_falls_back_to_the_config_target(mock_cim_data,
                                                           monkeypatch):
    """The `.get()` defaults behind those captions were the literal 0.10.
    A max-offer dict with no `target_irr` — the only case the default
    fires — must read config, or the fallback quietly re-hard-codes the
    number the rest of this move just removed."""
    from docx import Document

    from output.memo_writer import _add_section_1

    monkeypatch.setattr(cfg, "SOLVER_TARGET_IRR", 0.14)
    doc = Document()
    _add_section_1(doc, mock_cim_data, gate_results=[], scenario_results={},
                   max_offer={"max_price": 1_000_000})
    text = "\n".join(p.text for p in doc.paragraphs)

    assert "Maximum Offer Price (for 14% Base Case IRR)" in text


# ── Item T Category 2: config.VALUE_ADD_ASSUMPTIONS ──────────────────
#
# The whole opportunity-sizing layer of `analysis/value_add.py` was
# literals. The characterization net could not have caught a bad move
# here — worse, it cannot even SEE part of it: the economic-occupancy
# description reaches no snapshot at all (`grep "assumes" tests/snapshots`
# returns nothing), so byte-for-byte green over there proves nothing
# about this. Each key below is patched on the live dict, exactly as
# `webapp.services._merge_patch` patches a real ConfigOverride row, and
# both faces are asserted: the DOLLARS booked and the SENTENCE printed.

def _va_cim(mock_cim_data, **overrides):
    for k, v in overrides.items():
        setattr(mock_cim_data, k, v)
    return mock_cim_data


def _fin(gpr=600_000.0, egr=550_000.0, revenue=560_000.0):
    return {"income_summary": {"gpr": gpr, "egr": egr,
                               "total_revenue": revenue}}


def _op(result, category):
    return next((o for o in result["revenue_opportunities"]
                 if o["category"] == category), None)


def test_the_occupancy_target_drives_the_trigger_the_dollars_and_the_prose(
        mock_cim_data, monkeypatch):
    cim = _va_cim(mock_cim_data, physical_occupancy=0.90,
                  economic_occupancy=None, other_income=0.0)

    base = _op(identify_value_add(cim, _fin()), "Occupancy Improvement")
    assert base is not None and "to 93%" in base["description"]

    monkeypatch.setitem(cfg.VALUE_ADD_ASSUMPTIONS, "occupancy_target", 0.98)
    moved = _op(identify_value_add(cim, _fin()), "Occupancy Improvement")
    assert "to 98%" in moved["description"]
    assert moved["est_annual_impact"] > base["est_annual_impact"]

    # and the target is a TRIGGER too: below it, no opportunity at all
    monkeypatch.setitem(cfg.VALUE_ADD_ASSUMPTIONS, "occupancy_target", 0.85)
    assert _op(identify_value_add(cim, _fin()), "Occupancy Improvement") is None


def test_the_spread_recovery_share_drives_the_dollars_and_the_sentence(
        mock_cim_data, monkeypatch):
    """The path NO characterization snapshot renders. It was the literal
    0.5 with the word "half" beside it in prose — two copies of one
    number, in the same expression, which is the exact shape item T
    exists to kill."""
    cim = _va_cim(mock_cim_data, physical_occupancy=0.90,
                  economic_occupancy=0.75, other_income=0.0)

    base = _op(identify_value_add(cim, _fin()), "Economic Occupancy Recovery")
    assert base is not None
    # 600,000 * (0.90 - 0.75) * 0.50
    assert base["est_annual_impact"] == pytest.approx(45_000.0)
    assert "50% of the spread is recoverable" in base["description"]

    monkeypatch.setitem(cfg.VALUE_ADD_ASSUMPTIONS,
                        "spread_recovery_share", 0.80)
    moved = _op(identify_value_add(cim, _fin()), "Economic Occupancy Recovery")
    assert moved["est_annual_impact"] == pytest.approx(72_000.0)
    assert "80% of the spread is recoverable" in moved["description"]
    assert "50%" not in moved["description"]


def test_the_ecri_keys_drive_the_trigger_the_uplift_and_the_band(
        mock_cim_data, monkeypatch):
    cim = _va_cim(mock_cim_data, physical_occupancy=0.90,
                  economic_occupancy=None, other_income=0.0)

    base = _op(identify_value_add(cim, _fin()), "Revenue Management / ECRI")
    assert base is not None
    assert base["est_annual_impact"] == pytest.approx(550_000.0 * 0.03)
    assert "targeting 8-10% annual increases" in base["description"]
    assert "tenants > 6 months" in base["description"]

    # the floor is a real gate: raise it above the deal and the op goes
    monkeypatch.setitem(cfg.VALUE_ADD_ASSUMPTIONS, "ecri_min_occupancy", 0.95)
    assert _op(identify_value_add(cim, _fin()),
               "Revenue Management / ECRI") is None

    monkeypatch.setitem(cfg.VALUE_ADD_ASSUMPTIONS, "ecri_min_occupancy", 0.88)
    monkeypatch.setitem(cfg.VALUE_ADD_ASSUMPTIONS, "ecri_egr_uplift", 0.05)
    monkeypatch.setitem(cfg.VALUE_ADD_ASSUMPTIONS, "ecri_increase_range",
                        (0.12, 0.15))
    monkeypatch.setitem(cfg.VALUE_ADD_ASSUMPTIONS,
                        "ecri_tenant_tenure_months", 9)
    moved = _op(identify_value_add(cim, _fin()), "Revenue Management / ECRI")
    assert moved["est_annual_impact"] == pytest.approx(550_000.0 * 0.05)
    assert "targeting 12-15% annual increases" in moved["description"]
    assert "tenants > 9 months" in moved["description"]


def test_the_ancillary_keys_drive_the_trigger_the_uplift_and_the_band(
        mock_cim_data, monkeypatch):
    cim = _va_cim(mock_cim_data, physical_occupancy=0.90,
                  economic_occupancy=None, other_income=10_000.0)

    base = _op(identify_value_add(cim, _fin()), "Ancillary Revenue")
    assert base is not None
    assert base["est_annual_impact"] == pytest.approx(560_000.0 * 0.03)
    assert "target 5-8% of revenue" in base["description"]

    # other income is 10k/560k = 1.8%; drop the trigger below that and
    # the line stops reading as under-exploited
    monkeypatch.setitem(cfg.VALUE_ADD_ASSUMPTIONS, "ancillary_min_share", 0.01)
    assert _op(identify_value_add(cim, _fin()), "Ancillary Revenue") is None

    monkeypatch.setitem(cfg.VALUE_ADD_ASSUMPTIONS, "ancillary_min_share", 0.05)
    monkeypatch.setitem(cfg.VALUE_ADD_ASSUMPTIONS,
                        "ancillary_revenue_uplift", 0.06)
    monkeypatch.setitem(cfg.VALUE_ADD_ASSUMPTIONS,
                        "ancillary_target_share", (0.06, 0.11))
    moved = _op(identify_value_add(cim, _fin()), "Ancillary Revenue")
    assert moved["est_annual_impact"] == pytest.approx(560_000.0 * 0.06)
    assert "target 6-11% of revenue" in moved["description"]


def test_the_uplift_total_follows_every_assumption_it_sums(mock_cim_data,
                                                           monkeypatch):
    """`estimated_noi_uplift` is the number that reaches the memo and the
    investor summary. A key moved into config but read back wrongly would
    leave the individual op right and this total stale, or vice versa."""
    cim = _va_cim(mock_cim_data, physical_occupancy=0.90,
                  economic_occupancy=0.75, other_income=0.0)
    before = identify_value_add(cim, _fin())["estimated_noi_uplift"]

    monkeypatch.setitem(cfg.VALUE_ADD_ASSUMPTIONS,
                        "spread_recovery_share", 0.80)
    after = identify_value_add(cim, _fin())["estimated_noi_uplift"]
    assert after == pytest.approx(before + 27_000.0)     # 600k * 0.15 * 0.30


# ── Item T Category 2: config.RENOVATION_COST ────────────────────────

def _capex(cim):
    return {i["item"]: i for i in identify_value_add(cim, _fin())["capex_items"]}


def test_the_renovation_age_triggers_come_from_config(mock_cim_data,
                                                      monkeypatch):
    from registry import asset_age

    cim = _va_cim(mock_cim_data, physical_occupancy=0.90,
                  economic_occupancy=None, other_income=0.0,
                  year_built=2014, nrsf=50_000)
    age = asset_age(2014)
    assert age is not None and 10 < age <= 15, (
        "fixture vintage must sit between the security and LED triggers "
        f"for this test to discriminate; age is {age}")

    items = _capex(cim)
    assert "Security System Upgrade" in items          # min_age 10, age > 10
    assert "LED Lighting Upgrade" not in items         # min_age 15, age <= 15

    monkeypatch.setitem(cfg.RENOVATION_COST["led_lighting"], "min_age", 5)
    assert "LED Lighting Upgrade" in _capex(cim)

    monkeypatch.setitem(cfg.RENOVATION_COST["security"], "min_age", 99)
    assert "Security System Upgrade" not in _capex(cim)


def test_the_roof_high_priority_age_comes_from_config(mock_cim_data,
                                                      monkeypatch):
    cim = _va_cim(mock_cim_data, physical_occupancy=0.90,
                  economic_occupancy=None, other_income=0.0,
                  year_built=1999, nrsf=50_000)
    assert _capex(cim)["Roof Replacement / Repair"]["priority"] == "Medium"

    monkeypatch.setitem(cfg.RENOVATION_COST["roof"], "high_priority_age", 20)
    assert _capex(cim)["Roof Replacement / Repair"]["priority"] == "High"


def test_the_renovation_costs_come_from_config(mock_cim_data, monkeypatch):
    cim = _va_cim(mock_cim_data, physical_occupancy=0.90,
                  economic_occupancy=None, other_income=0.0,
                  year_built=1990, nrsf=50_000)
    items = _capex(cim)
    # per_sf costs multiply NRSF; flat amounts do not
    assert items["Roof Replacement / Repair"]["est_cost_range"] == \
        "$75,000 - $150,000"
    assert items["Security System Upgrade"]["est_cost_range"] == \
        "$15,000 - $50,000"
    assert items["Signage & Curb Appeal"]["est_cost_range"] == \
        "$5,000 - $25,000"

    monkeypatch.setitem(cfg.RENOVATION_COST["roof"], "per_sf", (2.0, 4.0))
    monkeypatch.setitem(cfg.RENOVATION_COST["signage"], "amount",
                        (7_000, 30_000))
    items = _capex(cim)
    assert items["Roof Replacement / Repair"]["est_cost_range"] == \
        "$100,000 - $200,000"
    assert items["Signage & Curb Appeal"]["est_cost_range"] == \
        "$7,000 - $30,000"


def test_the_capex_list_keeps_config_declaration_order(mock_cim_data):
    """Render order is the config dict's order, and the age-gated items
    still precede the always-on ones — the shape the five hand-written
    branches produced."""
    cim = _va_cim(mock_cim_data, physical_occupancy=0.90,
                  economic_occupancy=None, other_income=0.0,
                  year_built=1990, nrsf=50_000)
    ordered = [i["item"] for i in identify_value_add(cim, _fin())["capex_items"]]
    assert ordered == [cfg.RENOVATION_COST[k]["item"]
                       for k in ("roof", "led_lighting", "security",
                                 "signage", "website")]


def test_a_deal_with_no_vintage_still_gets_the_ungated_items(mock_cim_data):
    """`year_built=None` skips every `min_age` spec and keeps the rest —
    the pre-existing behaviour, now a property of the data rather than of
    where the `if year_built:` block happened to end."""
    cim = _va_cim(mock_cim_data, physical_occupancy=0.90,
                  economic_occupancy=None, other_income=0.0,
                  year_built=None, nrsf=50_000)
    assert [i["item"] for i in identify_value_add(cim, _fin())["capex_items"]] \
        == ["Signage & Curb Appeal", "Website & Digital Presence"]


def test_a_per_sf_item_without_nrsf_stays_tbd(mock_cim_data):
    """It is a diligence prompt, not an underwriting input — so it says
    TBD rather than inventing a square footage. (`nrsf or 1` elsewhere is
    item T Category 4's target; this is not that.)"""
    cim = _va_cim(mock_cim_data, physical_occupancy=0.90,
                  economic_occupancy=None, other_income=0.0,
                  year_built=1990, nrsf=None)
    items = _capex(cim)
    assert items["Roof Replacement / Repair"]["est_cost_range"] == "TBD"
    assert items["Security System Upgrade"]["est_cost_range"] == \
        "$15,000 - $50,000"          # flat amounts do not need NRSF


def test_value_add_no_longer_imports_the_expense_benchmarks(mock_cim_data):
    """It imported `EXPENSE_BENCHMARKS` and never used it. An unused
    import of a settings-editable dict reads like this module benchmarks
    against it — and a reader who believes that will look for a bug that
    does not exist, or add one that does."""
    import analysis.value_add as va_mod

    assert not hasattr(va_mod, "EXPENSE_BENCHMARKS")


@pytest.mark.django_db
def test_a_stored_row_reaches_the_value_add_assumptions(mock_cim_data):
    """The Category 2 twin of the market test above: a real row in the
    database changing a real dollar figure and a real sentence. Every
    link is unit-tested separately; only this one fails if they stop
    composing — which is what "the UI claims the override works and the
    model proves otherwise" looks like, the defect item T names in its
    own rationale.

    Also pins that `VALUE_ADD_ASSUMPTIONS` is a `_PATCHED_DICTS` entry AND
    an `override_key_registry` key. Being one without the other is silent:
    a registry key with no patch lane saves a row nothing reads; a patch
    lane with no registry key is unreachable from the settings page and
    `build_config_patch` reports it `skipped`."""
    from django.utils import timezone

    from webapp.models import ConfigOverride
    from webapp.services import (_patched_config, build_config_patch,
                                 resolve_config_overrides)

    mock_cim_data.physical_occupancy = 0.90
    mock_cim_data.economic_occupancy = 0.75
    mock_cim_data.other_income = 0.0

    before = _op(identify_value_add(mock_cim_data, _fin()),
                 "Economic Occupancy Recovery")
    assert before["est_annual_impact"] == pytest.approx(45_000.0)

    ConfigOverride.objects.create(
        key="VALUE_ADD_ASSUMPTIONS.spread_recovery_share", value=0.80,
        effective_date=timezone.localdate())
    deltas = resolve_config_overrides("", timezone.localdate())
    patch, _solver, skipped = build_config_patch(deltas)
    assert skipped == [], "the key is not reachable from the settings page"

    with _patched_config(patch):
        after = _op(identify_value_add(mock_cim_data, _fin()),
                    "Economic Occupancy Recovery")
    assert after["est_annual_impact"] == pytest.approx(72_000.0)
    assert "80% of the spread is recoverable" in after["description"]

    # reverted on exit — a leaked mutation reprices every later deal in
    # the same worker process
    assert cfg.VALUE_ADD_ASSUMPTIONS["spread_recovery_share"] == 0.50


@pytest.mark.django_db
def test_the_settings_page_offers_every_value_add_assumption(client,
                                                             django_user_model):
    """A config key with no registry entry is invisible: it exists, it
    matters, and no operator can reach it. All nine are offered, and each
    inherits a bound from its shape (PR #45) rather than needing one
    written by hand."""
    from webapp.forms import override_key_registry

    user = django_user_model.objects.create_user(username="op2", password="x")
    client.force_login(user)
    reg = override_key_registry()

    for key in cfg.VALUE_ADD_ASSUMPTIONS:
        dotted = f"VALUE_ADD_ASSUMPTIONS.{key}"
        assert dotted in reg, f"{dotted} is not editable from settings"
        lo, hi = reg[dotted]["bounds"]
        assert (lo, hi) != (None, None)

    content = client.get("/settings/").content.decode()
    assert "Value-Add Opportunity Assumptions" in content

    # RENOVATION_COST is deliberately NOT offered — presentation only,
    # see the note in webapp.services._PATCHED_DICTS
    assert not any(k.startswith("RENOVATION_COST") for k in reg)


# ── Item T Category 2: the asset-age register ────────────────────────

def test_no_age_threshold_survives_as_a_bare_literal():
    """The register in `config.ASSET_AGE_LADDERS` is only worth having if
    it is complete. Three ladders that disagree is a decision the
    operator made on 2026-08-05 with the deltas measured; a FOURTH ladder
    appearing quietly in some module is not a decision at all, it is the
    drift item T exists to stop.

    So: walk `analysis/`, `model/` and `output/` and fail on any
    comparison of an asset age against a numeric literal. `registry.py`
    is excluded — it DEFINES `AGE_BANDS`, and `config.py` holds the
    values, which is the whole point.
    """
    import ast
    from pathlib import Path

    root = Path(__file__).resolve().parent.parent
    offenders = []

    def is_age(node):
        return ((isinstance(node, ast.Name) and node.id == "age")
                or (isinstance(node, ast.Call)
                    and isinstance(node.func, ast.Name)
                    and node.func.id == "asset_age"))

    def is_num(node):
        return (isinstance(node, ast.Constant)
                and isinstance(node.value, (int, float))
                and not isinstance(node.value, bool))

    for pkg in ("analysis", "model", "output"):
        for path in sorted((root / pkg).glob("*.py")):
            tree = ast.parse(path.read_text(), filename=str(path))
            for node in ast.walk(tree):
                if not isinstance(node, ast.Compare):
                    continue
                left_age = is_age(node.left)
                right_age = any(is_age(c) for c in node.comparators)
                hit = ((left_age and any(is_num(c) for c in node.comparators))
                       or (right_age and is_num(node.left)))
                if hit:
                    offenders.append(
                        f"{path.relative_to(root)}:{node.lineno} — "
                        f"{ast.unparse(node)}")

    assert not offenders, (
        "asset age compared to a bare literal; declare the threshold in "
        "config.py and list it in ASSET_AGE_LADDERS:\n  "
        + "\n  ".join(offenders))


def test_the_age_register_names_every_ladder_that_exists():
    """`ASSET_AGE_LADDERS` is prose unless something checks that each
    name it lists actually resolves — a register pointing at a ladder
    someone renamed is worse than no register."""
    import registry

    assert cfg.ASSET_AGE_LADDERS == (
        "registry.AGE_BANDS", "RENOVATION_COST", "RISK_TRIGGERS")
    assert registry.AGE_BANDS
    assert any("min_age" in spec for spec in cfg.RENOVATION_COST.values())
    assert cfg.RISK_TRIGGERS["aging_plant_age"]


def test_the_aging_plant_risk_reads_its_trigger_from_config(mock_cim_data,
                                                            monkeypatch):
    """It was `if age > 25:`, a bare literal — the second of the three
    ladders, and the one no other module could see."""
    import datetime as dt

    this_year = dt.date.today().year
    mock_cim_data.year_built = this_year - 20        # 20 < 25, no risk today
    assert _risk(_risks_for(mock_cim_data), "Aging physical plant") is None

    monkeypatch.setitem(cfg.RISK_TRIGGERS, "aging_plant_age", 15)
    risk = _risk(_risks_for(mock_cim_data), "Aging physical plant")
    assert risk is not None
    assert "20 years old" in risk["description"]


# ── Item T Category 3: the model layer's own hard-codes ──────────────
#
# Three moves, and the characterization net can only see one of them.
# `SENSITIVITY_GRID` and `EXPENSE_RATIO` are pure literal→config moves,
# so byte-for-byte green over there is exactly what a dead wire looks
# like (Category 1's lesson). `SOLVER_BOUNDS` is the opposite problem:
# it MOVED the snapshots deliberately, and a moved snapshot proves the
# wire is live but nothing about whether it is right. Both halves get
# asserted here.


# ── SOLVER_BOUNDS — one bracket, three solvers ───────────────────────

def test_the_bracket_is_read_at_call_time_not_frozen_at_import(monkeypatch):
    """The whole point of a config key. `model.solver` imports four other
    config names by value at the top of the file; the bracket must not
    join them.

    The REBIND at the end is the half that discriminates. `setitem` on the
    live dict is visible even to a module that did `from config import
    SOLVER_BOUNDS`, because both names point at the same object — so a
    setitem-only test passes on a frozen import and proves nothing.
    Replacing the attribute is the mutation a frozen binding cannot see.
    """
    from model.solver import solver_price_bracket

    assert solver_price_bracket(300_000) == (1_500_000.0, 15_000_000.0)

    monkeypatch.setitem(cfg.SOLVER_BOUNDS, "dear_entry_cap", 0.03)
    assert solver_price_bracket(300_000) == (1_500_000.0, 10_000_000.0)

    monkeypatch.setitem(cfg.SOLVER_BOUNDS, "cheap_entry_cap", 0.25)
    assert solver_price_bracket(300_000) == (1_200_000.0, 10_000_000.0)

    monkeypatch.setattr(cfg, "SOLVER_BOUNDS",
                        {"cheap_entry_cap": 0.10, "dear_entry_cap": 0.05,
                         "zero_noi_low_price": 1, "zero_noi_high_price": 2})
    assert solver_price_bracket(300_000) == (3_000_000.0, 6_000_000.0)


def test_a_non_positive_noi_falls_back_to_the_dollar_window(monkeypatch):
    """An implied cap rate on zero or negative NOI is meaningless — the
    division would return 0 or flip the bracket's ends. All three solvers
    carried the same pair of raw dollar bounds for it."""
    from model.solver import solver_price_bracket

    for noi in (0, None, -50_000):
        assert solver_price_bracket(noi) == (100_000.0, 50_000_000.0)

    monkeypatch.setitem(cfg.SOLVER_BOUNDS, "zero_noi_high_price", 9_000_000)
    assert solver_price_bracket(0) == (100_000.0, 9_000_000.0)


def test_all_three_solvers_bisect_the_same_bracket():
    """The defect this key closes was a DISAGREEMENT, not an absence: the
    static and levered solvers stopped at a 3% implied entry cap and the
    value-add solver went to 2%, with nothing recording that they
    differed. An AST walk is the only check that stays true — a test that
    calls the three solvers and compares answers cannot tell a shared
    bracket from two brackets that happen to bracket the same root.
    """
    import ast
    from pathlib import Path

    source = (Path(__file__).resolve().parent.parent
              / "model" / "solver.py").read_text()
    tree = ast.parse(source)

    users = set()
    for node in ast.walk(tree):
        if not isinstance(node, ast.FunctionDef):
            continue
        if not node.name.startswith("solve_max_price"):
            continue
        calls = {n.func.id for n in ast.walk(node)
                 if isinstance(n, ast.Call) and isinstance(n.func, ast.Name)}
        assert "solver_price_bracket" in calls, (
            f"{node.name} does not call solver_price_bracket — it has a "
            "bracket of its own again")
        users.add(node.name)

    assert users == {"solve_max_price", "solve_max_price_value_add",
                     "solve_max_price_levered"}


def test_a_truncated_answer_is_reported_not_returned_silently(caplog):
    """The measurement that chose 2% over 3%, asserted as behaviour.

    Bisection cannot find a root outside its bracket: every iteration
    pushes `low` up and the loop ends holding `high`, which is a price, in
    the shape of an answer, at an IRR nowhere near the target. Nothing on
    any surface reads `converged`, so before this the only signal was a
    number that happened to be round.

    Driven by squeezing the bracket rather than by contriving a deal —
    same arithmetic, and it does not depend on any fixture continuing to
    be extreme enough.
    """
    import logging

    from model.solver import solve_max_price

    with mock.patch.dict(cfg.SOLVER_BOUNDS, {"dear_entry_cap": 0.19}):
        with caplog.at_level(logging.WARNING, logger="cim_analyst"):
            out = solve_max_price(adjusted_ttm_noi=300_000, capex=0,
                                  expense_ratio=0.40)

    assert out["converged"] is False
    # The answer IS the ceiling, to the dollar — 300,000 / 0.19.
    assert out["max_price"] == pytest.approx(300_000 / 0.19)
    assert "the answer is the search ceiling" in caplog.text
    assert "SOLVER_BOUNDS" in caplog.text


def test_a_converged_answer_says_nothing(caplog):
    """The other half, and the one that makes the warning worth having: a
    warning that fires on every deal is noise nobody reads. The default
    bracket converges on the default fixture."""
    import logging

    from model.solver import solve_max_price

    with caplog.at_level(logging.WARNING, logger="cim_analyst"):
        out = solve_max_price(adjusted_ttm_noi=300_000, capex=0,
                              expense_ratio=0.40)

    assert out["converged"] is True
    assert "the answer is the search" not in caplog.text


def test_the_bracket_moves_a_real_max_offer(tmp_path, monkeypatch):
    """End to end, through the pipeline rather than the solver, because
    what a squeezed bracket does to a REPORTED max offer is the thing
    worth pinning. `run_analysis` reaches the static solver; a bracket
    frozen anywhere along that path leaves this unchanged."""
    import data.comp_db as comp_db_module
    monkeypatch.setattr(comp_db_module, "COMP_DB_PATH",
                        str(tmp_path / "comps.db"))

    def _max_offer(out_name):
        out = tmp_path / out_name
        out.mkdir()
        r = AnalysisResult(pdf_path="b.pdf", cim_data=stabilized_deal())
        run_analysis(r, output_dir=str(out))
        return r.max_offer

    base = _max_offer("base")
    assert base["converged"] is True

    monkeypatch.setitem(cfg.SOLVER_BOUNDS, "dear_entry_cap", 0.08)
    squeezed = _max_offer("squeezed")

    assert squeezed["max_price"] < base["max_price"]
    assert squeezed["converged"] is False


# ── SENSITIVITY_GRID — the axes, not the nine offsets ────────────────

def test_the_grid_reproduces_the_literal_lists_it_replaced():
    """The exact nine-and-nine `_build_sensitivity` carried inline. Not
    approx: these multiply a price and add to a cap rate, and the labels
    are formatted from them, so a value one float-ulp off changes a
    rendered column header."""
    from model.returns_model import _axis_offsets

    assert _axis_offsets(0.10, 0.025, "price") == [
        -0.10, -0.075, -0.05, -0.025, 0.0, 0.025, 0.05, 0.075, 0.10]
    assert _axis_offsets(0.0100, 0.0025, "exit cap") == [
        -0.0100, -0.0075, -0.0050, -0.0025, 0.0, 0.0025, 0.0050, 0.0075,
        0.0100]


@pytest.mark.parametrize("span,step", [(0.10, 0.025),     # the shipped price axis
                                       (0.0100, 0.0025),  # the shipped cap axis
                                       (0.0030, 0.0006),  # ±30bps in 6bp steps
                                       (0.0015, 0.0003)])
def test_the_centre_offset_is_a_positive_zero(span, step):
    """Accumulating from −span can leave the centre as a residue that is
    NEGATIVE zero, and `f"{-0.0:+.1%}"` renders "-0.0%" where the column
    has always read "+0.0%". Building outward from the centre makes it
    exactly `0 * step`.

    The last two pairs are why this test is parametrized rather than
    written against the shipped axes. On 0.10/0.025 the two constructions
    agree exactly — `-0.10 + 4 × 0.025` is a clean 0.0 — so a test using
    only the shipped values passes on BOTH implementations and guards
    nothing. It was written that way first and a deliberate mutation
    walked straight through it. A sweep of valid (span, step) pairs found
    553 that do misplace the centre; ±30bps in 6bp steps is one of them
    and is a plausible cap axis for a tight market.
    """
    from model.returns_model import _axis_offsets

    offsets = _axis_offsets(span, step, "price")
    centre = offsets[len(offsets) // 2]

    assert f"{centre:+.1%}" == "+0.0%"
    assert math.copysign(1, centre) == 1.0


def test_the_grid_axes_come_from_config(monkeypatch):
    """Both axes, both dimensions: a wider span adds cells, a coarser
    step removes them, and the labels follow."""
    from model.returns_model import build_returns_model

    def _grid():
        return build_returns_model(adjusted_ttm_noi=300_000,
                                   asking_price=4_000_000, nrsf=50_000,
                                   capex=0)["sensitivity"]

    base = _grid()
    assert len(base["price_labels"]) == 9 and len(base["cap_labels"]) == 9
    assert base["price_labels"][0] == "-10.0%"

    monkeypatch.setitem(cfg.SENSITIVITY_GRID, "price_span", 0.20)
    monkeypatch.setitem(cfg.SENSITIVITY_GRID, "exit_cap_step", 0.0050)
    moved = _grid()

    assert len(moved["price_labels"]) == 17
    assert moved["price_labels"][0] == "-20.0%"
    assert len(moved["cap_labels"]) == 5
    assert len(moved["irr_grid"]) == 17 and len(moved["irr_grid"][0]) == 5
    # The centre cell is still the base case, whatever the axes do.
    assert moved["irr_grid"][8][2] == pytest.approx(base["irr_grid"][4][4])


def test_the_grid_axes_divide_evenly():
    """config.py's own values, checked live — the same discipline
    `test_every_default_sits_inside_its_own_bounds` applies to bounds. A
    span that is not a whole multiple of its step silently stops the axis
    short, which reads as a layout quirk rather than as lost downside."""
    from model.returns_model import _axis_offsets

    grid = cfg.SENSITIVITY_GRID
    for span, step, axis in (("price_span", "price_step", "price"),
                             ("exit_cap_span", "exit_cap_step", "exit cap")):
        offsets = _axis_offsets(grid[span], grid[step], axis)
        assert offsets[0] == pytest.approx(-grid[span])
        assert offsets[-1] == pytest.approx(grid[span])


@pytest.mark.parametrize("span,step", [(0.10, 0.03), (0.10, 0.0), (0.10, -0.01)])
def test_a_bad_axis_raises_instead_of_truncating(span, step):
    from model.returns_model import _axis_offsets

    with pytest.raises(ValueError):
        _axis_offsets(span, step, "price")


# ── EXPENSE_RATIO — the default, the clamp, and their relation ───────

def test_the_expense_ratio_default_sits_inside_the_benchmark_band():
    """The relation the scope asked to be stated once. A band edited past
    the default would leave `clamp_expense_ratio(None)` returning a ratio
    the benchmarks themselves call implausible."""
    low, high = cfg.EXPENSE_BENCHMARKS["opex_revenue_ratio"]

    assert low <= cfg.EXPENSE_RATIO["default"] <= high
    # And deliberately NOT the midpoint — same argument as the management
    # fee target. A derived `(low + high) / 2` looks like a tidy-up and
    # silently re-underwrites every deal whose financials yield no ratio.
    assert cfg.EXPENSE_RATIO["default"] != (low + high) / 2


def test_the_clamp_is_the_band_widened_not_a_second_pair(monkeypatch):
    """It was `EXPENSE_RATIO_CLAMP = (0.25, 0.65)` in registry.py, beside
    a band of (0.35, 0.55) it had no stated relation to. Editing the band
    now moves the clamp with it; before, an operator who widened the band
    kept the old clamp and had ratios clipped that they had just declared
    credible."""
    from registry import clamp_expense_ratio, expense_ratio_clamp

    assert expense_ratio_clamp() == (0.25, 0.65)

    monkeypatch.setitem(cfg.EXPENSE_BENCHMARKS, "opex_revenue_ratio",
                        (0.30, 0.70))
    assert expense_ratio_clamp() == (0.20, 0.80)
    assert clamp_expense_ratio(0.75) == 0.75          # was clipped to 0.65

    monkeypatch.setitem(cfg.EXPENSE_RATIO, "clamp_tolerance", 0.0)
    assert expense_ratio_clamp() == (0.30, 0.70)
    assert clamp_expense_ratio(0.75) == 0.70


def test_the_default_and_the_clamp_are_two_different_readings(monkeypatch):
    """The coincident-values trap: today `clamp_expense_ratio(None)`
    returns 0.40 because the default happens to sit inside the clamp, so
    a test asserting 0.40 passes whichever of the two the code reads.

    Moving the band until the clamp EXCLUDES the default separates them —
    the default is applied first and then clamped, so the answer is the
    clamp floor, not the default.
    """
    from registry import clamp_expense_ratio

    assert clamp_expense_ratio(None) == 0.40

    monkeypatch.setitem(cfg.EXPENSE_RATIO, "default", 0.30)
    assert clamp_expense_ratio(None) == 0.30

    monkeypatch.setitem(cfg.EXPENSE_BENCHMARKS, "opex_revenue_ratio",
                        (0.55, 0.60))
    assert clamp_expense_ratio(None) == 0.45         # clamp floor wins
    assert clamp_expense_ratio(0.30) == 0.45         # and it is not the default


def test_the_expense_ratio_default_reaches_the_projection(monkeypatch):
    """`project_cash_flows` is the ONE projection, so this single call
    site carries the assumed expense load into every scenario, every
    sensitivity cell and every solver iteration. `expense_ratio=None` is
    the path a deal takes when the financials yield no ratio at all.

    What the ratio actually does is worth stating, because the obvious
    assertion is wrong: Year 1 NOI is `ttm_noi × (1 + bump)` and does not
    depend on it at all. The ratio SPLITS that NOI into revenue and
    expenses so the two can grow at different rates — so it shows up from
    Year 2 onward, and a heavier assumed load means a bigger expense base
    compounding at `exp_growth`, which drags the terminal NOI and the IRR
    with it. Asserting `noi[0]` would have passed on a dead wire.
    """
    from analysis.valuation import project_cash_flows

    def _run():
        return project_cash_flows(ttm_noi=300_000, price=4_000_000, capex=0,
                                  params=cfg.SCENARIO_DEFAULTS[
                                      ScenarioType.BASE],
                                  expense_ratio=None,
                                  exit_cap=0.0625)

    base = _run()
    assert base["revenue"][0] == pytest.approx(315_000 / (1 - 0.40))

    monkeypatch.setitem(cfg.EXPENSE_RATIO, "default", 0.50)
    heavier = _run()

    assert heavier["revenue"][0] == pytest.approx(315_000 / (1 - 0.50))
    assert heavier["noi"][0] == base["noi"][0]        # Year 1 is unaffected
    assert heavier["noi"][-1] < base["noi"][-1]
    assert heavier["irr"] < base["irr"]


@pytest.mark.django_db
def test_a_stored_row_reaches_the_expense_ratio(monkeypatch):
    """The Category 3 twin of the market and value-add tests above: a real
    row in the database moving a real IRR. Pins that `EXPENSE_RATIO` is
    BOTH a `_PATCHED_DICTS` entry and an `override_key_registry` key —
    one without the other is silent in the way item T exists to kill."""
    from django.utils import timezone

    from analysis.valuation import project_cash_flows
    from webapp.models import ConfigOverride
    from webapp.services import (_patched_config, build_config_patch,
                                 resolve_config_overrides)

    def _irr():
        return project_cash_flows(ttm_noi=300_000, price=4_000_000, capex=0,
                                  params=cfg.SCENARIO_DEFAULTS[
                                      ScenarioType.BASE],
                                  expense_ratio=None,
                                  exit_cap=0.0625)["irr"]

    before = _irr()

    ConfigOverride.objects.create(key="EXPENSE_RATIO.default", value=0.50,
                                  effective_date=timezone.localdate())
    deltas = resolve_config_overrides("", timezone.localdate())
    patch, _solver, skipped = build_config_patch(deltas)
    assert skipped == [], "the key is not reachable from the settings page"

    with _patched_config(patch):
        after = _irr()
    assert after < before

    # reverted on exit — a leaked mutation reprices every later deal in
    # the same worker process
    assert cfg.EXPENSE_RATIO["default"] == 0.40


def test_the_registry_constants_are_gone():
    """`registry.DEFAULT_EXPENSE_RATIO` and `EXPENSE_RATIO_CLAMP` were
    module scalars: a module binding one by value at import can never see
    a settings patch, so leaving them behind as aliases would leave a
    second, frozen source of the same numbers — the exact shape item T
    is removing."""
    import registry

    assert not hasattr(registry, "DEFAULT_EXPENSE_RATIO")
    assert not hasattr(registry, "EXPENSE_RATIO_CLAMP")


def test_the_composed_clamp_cannot_reach_a_division_by_zero(caplog):
    """The defect the derivation INTRODUCED, found by the pre-push audit.

    `_bounds_for` bounds each editable field on its own shape, and both
    inputs here are shares in (0, 1) — individually legal at every value.
    Their sum is not: an `opex_revenue_ratio` high of 0.90 and a
    `clamp_tolerance` of 0.10 compose to a clamp ceiling of exactly 1.0,
    and `analysis/valuation.py` then evaluates `yr1_noi / (1 - ratio)` on
    a deal whose expenses reach revenue. That is a ZeroDivisionError and
    an unhandled 500, reached by two settings edits that each passed
    validation.

    Unreachable before this PR — the clamp was the constant (0.25, 0.65),
    which no operator could touch. Making it follow the band is what put
    a composed value in reach, so the bound belongs with the derivation.
    """
    import logging

    import registry
    from analysis.valuation import project_cash_flows
    from registry import EXPENSE_RATIO_LIMITS, expense_ratio_clamp

    registry._WARNED.clear()          # the warning is once-per-process
    with mock.patch.dict(cfg.EXPENSE_BENCHMARKS,
                         {"opex_revenue_ratio": (0.35, 0.90)}), \
            mock.patch.dict(cfg.EXPENSE_RATIO, {"clamp_tolerance": 0.10}):
        with caplog.at_level(logging.WARNING, logger="cim_analyst"):
            low, high = expense_ratio_clamp()

        assert high == EXPENSE_RATIO_LIMITS[1]
        assert high < 1.0
        assert "outside the limits" in caplog.text

        # and the projection it protects survives a deal whose stated
        # expenses exceed its revenue
        out = project_cash_flows(
            ttm_noi=300_000, price=4_000_000, capex=0,
            params=cfg.SCENARIO_DEFAULTS[ScenarioType.BASE],
            expense_ratio=1.20, exit_cap=0.0625)
        assert out["revenue"][0] > 0


def test_a_band_low_under_the_tolerance_cannot_derive_a_negative_floor():
    """The same composition from the other side: a band low of 0.05 minus
    a 0.10 tolerance derives a clamp FLOOR of −0.05, which would let a
    deal be underwritten on expenses below zero — free money, and no
    exception to notice it by."""
    from registry import EXPENSE_RATIO_LIMITS, clamp_expense_ratio, \
        expense_ratio_clamp

    with mock.patch.dict(cfg.EXPENSE_BENCHMARKS,
                         {"opex_revenue_ratio": (0.05, 0.55)}):
        low, _high = expense_ratio_clamp()

        assert low == EXPENSE_RATIO_LIMITS[0] == 0.0
        assert clamp_expense_ratio(-0.30) == 0.0


def test_the_limits_are_not_reached_on_the_shipped_defaults():
    """The bound must be a backstop, not a live constraint — if it bound
    today it would be silently re-underwriting every deal, and the
    warning above would fire on every run."""
    from registry import EXPENSE_RATIO_LIMITS, expense_ratio_clamp

    floor, ceiling = EXPENSE_RATIO_LIMITS
    low, high = expense_ratio_clamp()

    assert (low, high) == (0.25, 0.65)
    assert floor < low and high < ceiling


@pytest.mark.parametrize("band,tolerance", [((1.00, 1.00), 0.0),
                                            ((0.97, 0.99), 0.0),
                                            ((0.99, 1.00), 0.05)])
def test_the_bounded_clamp_can_never_come_back_inverted(band, tolerance):
    """The hole the FIRST version of the bound left, found by the second
    audit pass — and it is the same shape as the bug it was fixing.

    Bounding only the high end leaves the low end free to climb past the
    ceiling: a band of (1.0, 1.0) with a zero tolerance derived
    (1.0, 1.0) and bounded it to (1.0, 0.95). `max(lo, min(hi, r))` on an
    inverted pair returns `lo` for EVERY input — the clamp stops reading
    its argument and hands back 1.0, which is the exact
    ZeroDivisionError the limits exist to prevent, now reached THROUGH
    the guard.

    The first fix's own test could not see this: it used a band of
    (0.35, 0.90), where only the high end needed clamping. Coincidence,
    not coverage — the third time that pattern has bitten in this item.
    """
    from analysis.valuation import project_cash_flows
    from registry import EXPENSE_RATIO_LIMITS, clamp_expense_ratio, \
        expense_ratio_clamp

    floor, ceiling = EXPENSE_RATIO_LIMITS
    with mock.patch.dict(cfg.EXPENSE_BENCHMARKS,
                         {"opex_revenue_ratio": band}), \
            mock.patch.dict(cfg.EXPENSE_RATIO,
                            {"clamp_tolerance": tolerance}):
        low, high = expense_ratio_clamp()

        assert low <= high, "inverted clamp — max(lo, min(hi, r)) is now lo"
        assert floor <= low and high <= ceiling

        # the clamp still READS its argument rather than returning a
        # constant, which is what an inverted pair silently does
        assert clamp_expense_ratio(0.40) <= ceiling
        assert clamp_expense_ratio(0.99) <= ceiling

        # and the projection the whole guard protects survives
        out = project_cash_flows(
            ttm_noi=300_000, price=4_000_000, capex=0,
            params=cfg.SCENARIO_DEFAULTS[ScenarioType.BASE],
            expense_ratio=None, exit_cap=0.0625)
        assert out["revenue"][0] > 0


def test_the_limit_warning_is_reported_once_not_once_per_projection(caplog):
    """`expense_ratio_clamp` runs inside `project_cash_flows`, which the
    three solvers call up to 50 times each and the sensitivity grid 81
    more — so an unthrottled warning is ~200 identical lines for one
    deal, which buries the one line anybody needed to read."""
    import logging

    import registry
    from analysis.valuation import project_cash_flows

    registry._WARNED.clear()
    with mock.patch.dict(cfg.EXPENSE_BENCHMARKS,
                         {"opex_revenue_ratio": (0.35, 0.90)}):
        with caplog.at_level(logging.WARNING, logger="cim_analyst"):
            for _ in range(25):
                project_cash_flows(
                    ttm_noi=300_000, price=4_000_000, capex=0,
                    params=cfg.SCENARIO_DEFAULTS[ScenarioType.BASE],
                    expense_ratio=None, exit_cap=0.0625)

    hits = [r for r in caplog.records if "outside the limits" in r.getMessage()]
    assert len(hits) == 1, f"{len(hits)} lines for one misconfiguration"


def test_a_different_misconfiguration_still_speaks_up():
    """The throttle keys on the MESSAGE, not on a fired-once flag, so a
    second, different bad combination is not swallowed by the first."""
    import registry
    from registry import expense_ratio_clamp

    registry._WARNED.clear()
    for band in ((0.35, 0.90), (0.35, 1.00)):
        with mock.patch.dict(cfg.EXPENSE_BENCHMARKS,
                             {"opex_revenue_ratio": band}):
            expense_ratio_clamp()

    assert len(registry._WARNED) == 2
