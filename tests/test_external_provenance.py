"""A measured number is not a stated one.

The defect: Census enrichment fills `population_3mi` and `median_hhi_3mi`
BEFORE the pristine snapshot is saved, so a measured demographic sits in
`Deal.cim_json` looking exactly like an extracted one — and every surface
called it "stated in the CIM". On the number Gate 1 is decided by, for a
document that never mentioned it.

Three separate things had to be true for that to happen, and each gets
tests here, because fixing any two still leaves the memo lying:

1. the register had no VOCABULARY for a measurement (design decision 11's
   provenance set was closed at five);
2. the extract-time source log was never PERSISTED, so no later reader
   could have told the difference even with the word for it;
3. the analysis-time enrichment pass OVERWROTE what the extract-time pass
   knew — it finds the field already filled, resolves it at tier 1 and
   reports "CIM/override", which is true of what it was handed and false
   about where the number came from.

The surface assertions are the point (this file's siblings learned that
the hard way: a register that records perfectly and renders nowhere
delivers nothing). Each names the deliberate break that turns it red.
"""

import pytest

from analysis import assumptions as A
from extract.enrichment import (MEASURED_TIER, merge_source_logs,
                                origin_for)
from tests.test_characterization import stabilized_deal
from tests.test_config_single_source import _memo_text, _run

#: A log shaped exactly as `extract.enrichment` writes one: the two
#: demographics it measured, plus the ring centre it measured them around.
MEASURED_LOG = {
    "lat": {"tier": 2, "source": "ZCTA centroid", "value": 32.221},
    "lon": {"tier": 2, "source": "ZCTA centroid", "value": -110.926},
    "population_3mi": {"tier": 2, "source": "Census API", "value": 87_450},
    "median_hhi_3mi": {"tier": 2, "source": "Census API", "value": 64_300.0},
}


def _measured_deal():
    """The stabilized fixture with its demographics MEASURED, not stated —
    value and snapshot agreeing, which is what enrichment-before-save
    produces and what made the mislabel invisible."""
    cim = stabilized_deal()
    cim.population_3mi = 87_450
    cim.median_hhi_3mi = 64_300.0
    snapshot = {"population_3mi": 87_450, "median_hhi_3mi": 64_300.0,
                "nrsf": cim.nrsf}
    return cim, snapshot


def _row(rows, key):
    return next(r for r in rows if r.key == key)


# ── 1. The vocabulary ───────────────────────────────────────────────

def test_a_measured_population_is_not_reported_as_stated_in_the_cim():
    """The defect itself, at the register.

    MUTATION: drop the `enrichment_log` argument at either
    `assumptions.collect`'s call site, or delete the `EXTERNAL` branch in
    `_add_cim_rows` — the row goes back to "stated in the CIM" and every
    number in the memo stays identical.
    """
    cim, snapshot = _measured_deal()
    rows = A.collect(cim_data=cim, cim_snapshot=snapshot,
                     enrichment_log=MEASURED_LOG)

    pop = _row(rows, "cim.population_3mi")
    assert pop.provenance == A.EXTERNAL
    assert pop.provenance_label == "measured from public data"
    assert pop.provenance_label != A.PROVENANCE_LABELS[A.CIM]
    # A field the document really did state is untouched by any of this.
    assert _row(rows, "cim.nrsf").provenance == A.CIM


def test_the_measurement_row_discloses_the_ring_it_measured():
    """A 3-mile ring around a matched building and one around a ZIP's
    centroid are different claims about the same property — which is why
    `enrichment` stamps the centre separately in the first place. Naming
    the API and staying silent about the centre discloses the easier half.

    MUTATION: return `source` alone from `_measurement_detail`.
    """
    cim, snapshot = _measured_deal()
    rows = A.collect(cim_data=cim, cim_snapshot=snapshot,
                     enrichment_log=MEASURED_LOG)
    assert _row(rows, "cim.population_3mi").detail == (
        "Census API — ring centred on the ZCTA centroid")


def test_a_measurement_is_not_counted_among_the_rows_a_human_chose():
    """Operator's call, 2026-08-18: `CHOSEN` answers "did a human or a
    fallback produce this?", and a measurement is neither. Widening it
    would change what memo table B.1 promises a reader.

    MUTATION: add `EXTERNAL` to `CHOSEN` — B.1 grows rows whose intro
    sentence still enumerates only three provenances.
    """
    cim, snapshot = _measured_deal()
    rows = A.collect(cim_data=cim, cim_snapshot=snapshot,
                     enrichment_log=MEASURED_LOG)
    counts = A.summarize(rows)

    assert A.EXTERNAL not in A.CHOSEN
    assert counts["external"] == 2
    assert not _row(rows, "cim.population_3mi").chosen
    # The count identity the memo's lead sentence rests on still holds
    # with a sixth provenance in the vocabulary.
    assert counts["total"] == sum(counts[k] for k in A.PROVENANCE_KEYS)


def test_the_vocabulary_stays_closed_and_derived():
    """Decision 11's set is closed and its labels are the one source for
    the keys. A sixth member must not have been added by widening one and
    forgetting the other."""
    assert A.PROVENANCE_KEYS == tuple(A.PROVENANCE_LABELS)
    assert A.EXTERNAL in A.PROVENANCE_KEYS
    assert A.PROVENANCE_LABELS[A.EXTERNAL].strip()
    # Precedence is declaration order: the resolver takes a CIM-stated
    # population over a measured one (tier 1 beats tier 2), and a
    # measurement beats a shipped default.
    order = list(A.PROVENANCE_KEYS)
    assert order.index(A.CIM) < order.index(A.EXTERNAL) < order.index(A.CONFIG)


def test_without_a_log_every_stated_field_is_still_cim():
    """Not a degraded mode — it is what a run that never enriched means,
    and what every deal extracted before the log existed will report
    forever. It must not guess."""
    cim, snapshot = _measured_deal()
    rows = A.collect(cim_data=cim, cim_snapshot=snapshot)
    assert _row(rows, "cim.population_3mi").provenance == A.CIM
    assert A.summarize(rows)["external"] == 0


# ── 2. The log describes ONE value ──────────────────────────────────

def test_an_analyst_correction_is_not_credited_to_the_census():
    """The log entry records where one particular number came from. Type a
    population over the measured one and the entry stops describing it —
    reporting it anyway credits the Census with an analyst's figure.

    MUTATION: compare only the field NAME in `origin_for`, ignoring
    `value`.
    """
    cim, snapshot = _measured_deal()
    cim.population_3mi = 61_000          # analyst overtyped it

    rows = A.collect(cim_data=cim, cim_snapshot=snapshot,
                     enrichment_log=MEASURED_LOG)
    pop = _row(rows, "cim.population_3mi")
    assert pop.provenance == A.DEAL
    assert pop.was == 87_450, "the correction must still say what it displaced"


def test_the_gate_reports_no_source_for_a_number_nobody_measured():
    """Gate 1 prints the origin beside the population it decided on. Same
    entry, same staleness, same wrong credit — so it asks the same
    question through the same function.

    MUTATION: restore the direct `source_log[...]["source"]` lookup in
    `analysis.filters` — the gate credits "Census API" for a figure the
    analyst typed.
    """
    from analysis.filters import evaluate_gates

    cim, _ = _measured_deal()
    measured = evaluate_gates(cim, {}, None, source_log=MEASURED_LOG)[0]
    assert measured["source"] == "Census API"

    cim.population_3mi = 61_000
    overtyped = evaluate_gates(cim, {}, None, source_log=MEASURED_LOG)[0]
    assert overtyped["actual"] == "61,000"
    assert not overtyped["source"]


def test_origin_for_survives_a_json_round_trip():
    """The log reaches the register through a database column, so the
    comparison meets an int that left as an int and a float that left as a
    float — and `median_hhi_3mi` is the float."""
    import json

    stored = json.loads(json.dumps(MEASURED_LOG))
    assert origin_for(stored, "median_hhi_3mi", 64_300.0)["tier"] == MEASURED_TIER
    assert origin_for(stored, "population_3mi", 87_450)["tier"] == MEASURED_TIER
    assert origin_for(stored, "population_3mi", 87_451) is None
    assert origin_for(stored, "nrsf", 60_000) is None
    assert origin_for({}, "population_3mi", 87_450) is None


# ── 3. The second pass must not erase the first ─────────────────────

def test_a_second_enrichment_pass_does_not_overwrite_a_measurement():
    """The web path enriches twice, and the second pass sees the first
    pass's output already on `cim_data` — resolving it at tier 1,
    "CIM/override". Left alone that is the mislabel arriving by a second
    route, immune to every fix above.

    MUTATION: make `merge_source_logs` a plain `{**prior, **later}`.
    """
    from extract.enrichment import DataResolver

    first, second = {}, {}
    DataResolver(first).resolve("population_3mi", tier1_value=None,
                                tier2_fn=lambda: 87_450)
    DataResolver(second).resolve("population_3mi", tier1_value=87_450,
                                 tier2_fn=lambda: 87_450)
    assert second["population_3mi"]["tier"] == 1, "the shadowing this guards"

    merged = merge_source_logs(first, second)
    assert merged["population_3mi"]["tier"] == MEASURED_TIER
    assert merged["population_3mi"]["source"] == "Census API"


def test_a_later_pass_that_really_measured_wins():
    """The mirror image, and the reason the rule is "a later TIER-1 entry
    never displaces a measurement" rather than "the stored log wins": an
    analyst who corrects the address gets a genuinely new measurement, and
    that one IS the newer fact."""
    prior = {"population_3mi": {"tier": 2, "source": "Census API",
                                "value": 87_450}}
    later = {"population_3mi": {"tier": 2, "source": "Census API",
                                "value": 91_000}}
    assert merge_source_logs(prior, later)["population_3mi"]["value"] == 91_000
    # And a field only the later pass saw is simply added.
    assert "lat" in merge_source_logs(prior, {"lat": {"tier": 2}})


# ── 4. The wires ────────────────────────────────────────────────────

def test_the_engine_carries_a_stored_log_through_to_the_register(tmp_path):
    """The web path's whole wire in one call: `run_analysis` is handed the
    log the extract-time pass produced, and the register it assembles must
    use it. This run does not re-enrich (the fields are already filled),
    which is exactly the case that had no log at all.

    MUTATION: drop `enrichment_log=source_log` from the `collect` call in
    `engine.run_analysis`.
    """
    cim, snapshot = _measured_deal()
    result = _run(cim, tmp_path, cim_snapshot=snapshot,
                  enrichment_log=MEASURED_LOG)

    rows = A.from_dicts(result.assumption_register)
    assert _row(rows, "cim.population_3mi").provenance == A.EXTERNAL
    gate = next(g for g in result.gate_results if g["gate"] == 1)
    assert gate["source"] == "Census API"


def test_the_memo_says_measured_where_it_used_to_say_stated(tmp_path):
    """The surface the whole item exists for. An IC reader looks up the
    population in Appendix B and must not be told the seller stated it.

    MUTATION: any of the above — this is the assertion they all serve.
    """
    cim, snapshot = _measured_deal()
    result = _run(cim, tmp_path, cim_snapshot=snapshot,
                  enrichment_log=MEASURED_LOG)
    text = _memo_text(result.memo_path)

    assert "3-Mile Population" in text
    assert A.PROVENANCE_LABELS[A.EXTERNAL] in text


def test_the_cli_register_can_report_a_measurement_it_took(tmp_path):
    """The CLI enriches exactly once and needs no storage — its log still
    describes the values on `ctx.cim_data`. The sibling test
    `test_a_cli_register_never_claims_a_settings_or_deal_override` pins
    what the CLI cannot have; this pins what it can.

    MUTATION: drop `enrichment_log=source_log` from `run.py`'s `collect`
    call — the CLI's memo silently re-acquires the defect the web app's
    just lost, which is the exact split the fill log shipped with once.
    """
    rows = A.collect(cim_data=_measured_deal()[0],
                     enrichment_log=MEASURED_LOG)
    provenances = {r.provenance for r in rows}
    assert A.EXTERNAL in provenances
    assert A.SETTINGS not in provenances and A.DEAL not in provenances


# ── 5. Persistence — the half nothing could reconstruct ─────────────

@pytest.mark.django_db
def test_the_extract_worker_stores_the_log_beside_the_snapshot(tmp_path,
                                                               monkeypatch):
    """It has to be saved by the pass that measured. Nothing downstream
    can rebuild it: the analysis-time pass finds the field filled and
    reports tier 1.

    MUTATION: remove the `enrichment_log` key from `_extract_worker`'s
    `updates` — extraction still succeeds, every stored number is
    identical, and the deal reports its measured demographics as
    CIM-stated for the rest of its life.
    """
    from engine import AnalysisResult
    from webapp import services
    from webapp.models import Deal

    cim, _ = _measured_deal()

    class _Enrichment:
        source_log = MEASURED_LOG

    def _fake_extract(path):
        result = AnalysisResult(pdf_path=path)
        result.cim_data = cim
        result.extraction_report = {}
        result.enrichment = _Enrichment()
        return result

    monkeypatch.setattr(services, "extract_pdf_data", _fake_extract)
    deal = Deal.objects.create(deal_id="probe", property_name="Probe",
                               extract_requested_at=None)
    services._extract_worker(deal.pk, str(tmp_path / "deal.pdf"), None)

    deal.refresh_from_db()
    assert deal.enrichment_log, "the measuring pass saved no log"
    assert deal.enrichment_log["population_3mi"]["tier"] == MEASURED_TIER
    assert deal.cim_json["population_3mi"] == 87_450, "the snapshot too"


@pytest.mark.django_db
def test_the_assumptions_page_labels_a_measured_demographic(tmp_path):
    """The page had the label and could never reach it: it read the log
    off the LATEST RUN, and a deal with no run yet has none while a run
    that finds the demographics already filled stores no enrichment block
    at all. Both showed "CIM".

    MUTATION: read `source_log` from the run alone, as before.
    """
    from webapp import forms as f
    from webapp.models import Deal

    deal = Deal.objects.create(
        deal_id="probe2", property_name="Probe",
        cim_json={"population_3mi": 87_450, "nrsf": 60_000},
        enrichment_log=MEASURED_LOG)

    form = f.AssumptionsForm(initial={"population_3mi": 87_450,
                                      "nrsf": 60_000})
    rows = f.model_rows(form, [("population_3mi", "3-Mile Population"),
                               ("nrsf", "Rentable SF")],
                        deal.cim_json, deal.enrichment_log)
    by_label = {r["label"]: r["source"] for r in rows}
    assert by_label["3-Mile Population"] == "Census"
    assert by_label["Rentable SF"] == "CIM"


# ── 6. The residuals #101 left open ─────────────────────────────────
#
# Four defects that survived the merge. The first two are the ones the
# #101 session verified against merged `main` rather than trusting a PR
# body; the last two are the catch-all's other exits, which no test had
# ever walked. All four reproduce on 828005c.


def test_an_analyst_value_in_a_field_the_cim_never_filled_is_not_cim():
    """The third side of the same lie. `edited` needs a `prior` to differ
    FROM, so a number typed into an empty box was filed as "stated in the
    CIM" — a claim about a document that does not contain the figure at
    all, on a value a human chose.

    MUTATION: drop the `elif field in entered and prior is None` arm — the
    row reverts to `cim` and leaves memo table B.1 entirely, which is the
    table an IC reader checks to see what the analyst touched.
    """
    cim, snapshot = _measured_deal()
    cim.market_rent_psf = 1.35          # extraction never found one
    assert "market_rent_psf" not in snapshot

    rows = A.collect(cim_data=cim, cim_snapshot=snapshot,
                     enrichment_log=MEASURED_LOG,
                     deal_overrides={"cim_overrides": {
                         "market_rent_psf": 1.35}})

    row = _row(rows, "cim.market_rent_psf")
    assert row.provenance == A.DEAL
    assert row.chosen                       # so it reaches B.1
    assert row.was is None                  # it replaced nothing
    assert "entered" in row.detail


def test_a_field_the_analyst_never_touched_is_still_cim():
    """The other direction, because the branch above keys on membership of
    a delta dict and an over-broad witness would credit the analyst with
    the whole document.

    MUTATION: drop the `field in entered` half and test only `prior is
    None` — every field extraction missed becomes an analyst entry.
    """
    cim, snapshot = _measured_deal()
    cim.market_rent_psf = 1.35
    rows = A.collect(cim_data=cim, cim_snapshot=snapshot,
                     deal_overrides={"cim_overrides": {"nrsf": cim.nrsf}})
    assert _row(rows, "cim.market_rent_psf").provenance == A.CIM


def test_an_analyst_correction_still_reports_what_it_replaced():
    """The pre-existing `edited` arm is checked FIRST and must keep its
    `was`. A correction and a first entry are different facts and the new
    arm must not swallow the one that carries the displaced value.

    MUTATION: reorder the two `DEAL` arms — corrections lose `was`, and
    the memo's "Replaced" column silently empties.
    """
    cim, snapshot = _measured_deal()
    cim.nrsf = 62_000
    rows = A.collect(cim_data=cim, cim_snapshot={**snapshot, "nrsf": 50_000},
                     deal_overrides={"cim_overrides": {"nrsf": 62_000}})
    row = _row(rows, "cim.nrsf")
    assert row.provenance == A.DEAL
    assert row.was == 50_000
    assert "corrected" in row.detail


@pytest.mark.parametrize("entry,expected_source", [
    ({"tier": 3, "source": "comp_db", "value": 61_000}, "comp_db"),
    ({"tier": 4, "source": "default", "value": 61_000}, "default"),
])
def test_a_borrowed_or_invented_value_is_not_reported_as_stated(
        entry, expected_source):
    """The catch-all's unwalked exits. A comp-DB average is about OTHER
    properties and a tier-4 default is invention; both rendered "stated in
    the CIM", which is a worse claim than the tier-2 mislabel #101 fixed —
    a measurement is at least evidence about this property.

    MUTATION: drop the `tier is not None and tier != 1` arm — both fall
    through to `cim`, exactly as they did before this commit.
    """
    cim, snapshot = _measured_deal()
    # Value and snapshot agree — a tier-3/4 fill lands on `cim_data`
    # before `cim_json` is saved, exactly as a measurement does, which is
    # why neither is distinguishable without the log.
    cim.median_hhi_3mi = 61_000
    snapshot = {**snapshot, "median_hhi_3mi": 61_000}
    rows = A.collect(cim_data=cim, cim_snapshot=snapshot,
                     enrichment_log={"median_hhi_3mi": entry})

    row = _row(rows, "cim.median_hhi_3mi")
    assert row.provenance == A.FALLBACK
    assert row.chosen                       # surfaces in B.1, not buried
    assert expected_source in row.detail


def test_an_unknown_tier_degrades_rather_than_raising():
    """`enrichment_log` is a PERSISTED column. A tier a future version
    writes must not 500 the results page of a deal already in the
    database — the same contract `from_dicts` keeps for unknown keys.
    """
    cim, snapshot = _measured_deal()
    rows = A.collect(cim_data=cim, cim_snapshot=snapshot,
                     enrichment_log={"median_hhi_3mi": {
                         "tier": 9, "source": "some future tier",
                         "value": 64_300.0}})
    assert _row(rows, "cim.median_hhi_3mi").provenance == A.FALLBACK


def test_a_tier_none_entry_is_not_mistaken_for_a_borrowed_value():
    """The resolver stamps `{"tier": None}` for "not available". Its value
    is None so it never reaches a row — but the arm above tests `tier is
    not None` precisely so a hand-built or future log carrying a value
    beside a null tier is not read as tier 3.
    """
    cim, snapshot = _measured_deal()
    rows = A.collect(cim_data=cim, cim_snapshot=snapshot,
                     enrichment_log={"nrsf": {
                         "tier": None, "source": "not available",
                         "value": cim.nrsf}})
    assert _row(rows, "cim.nrsf").provenance == A.CIM


def test_no_call_site_supplies_tier_3_or_tier_4(monkeypatch):
    """Why the arm above is unreachable in production TODAY, pinned so
    that stops being true loudly. `enrich_cim_data` passes `tier2_fn` and
    nothing else; the day someone wires the comp DB, this test fails and
    whoever wired it has to decide what the register should say — which is
    the decision that was silently made wrong before.
    """
    from extract import enrichment

    seen = []
    real = enrichment.DataResolver.resolve

    def spy(self, field_name, tier1_value, **kw):
        seen.append(set(k for k, v in kw.items() if v is not None))
        return real(self, field_name, tier1_value, **kw)

    monkeypatch.setattr(enrichment.DataResolver, "resolve", spy)
    cim, _ = _measured_deal()
    enrichment.enrich_cim_data(cim, census_api_key=None)

    assert seen, "enrich_cim_data resolved no fields — the spy missed"
    supplied = set().union(*seen)
    assert supplied <= {"tier2_fn"}, (
        f"a call site now supplies {supplied - {'tier2_fn'}}; decide what "
        f"provenance it should carry in analysis.assumptions._add_cim_rows")


def test_resolve_returns_a_populated_tier_1_value_untouched():
    """The precondition `merge_source_logs` rests on. Its no-demotion rule
    is sound only because a second pass NEVER re-measures a field that is
    already filled — if `resolve` ever consulted tier 2 over a populated
    tier 1, the stored measurement and the value in use could diverge with
    nothing to detect it. `engine`'s own docstring describes re-enrichment
    as "a second chance at geocoding", an intent one step from breaking
    this.

    MUTATION: move the tier-1 early return below the tier-2 block.
    """
    from extract.enrichment import DataResolver

    calls = []
    log = {}
    out = DataResolver(log).resolve(
        "population_3mi", tier1_value=87_450,
        tier2_fn=lambda: calls.append(1) or 99_999)

    assert out == 87_450
    assert calls == [], "tier 2 was consulted over a populated tier 1"
    assert log["population_3mi"]["tier"] == 1


# ── 7. The surfaces that called a measured run all-defaults ─────────


def _register_text(rows):
    """Appendix B alone, rendered from hand-built rows.

    Direct rather than through `_run`, because the defect needs
    `chosen == 0` WITH a measurement present and every pipeline fixture
    fills at least one value from a default. Reaching that state through a
    real run would mean building a CIM that states everything, which tests
    the fixture rather than the renderer.
    """
    from docx import Document

    from output.memo_writer import _add_assumption_register

    doc = Document()
    _add_assumption_register(doc, A.to_dicts(rows))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for r in table.rows:
            parts += [c.text for c in r.cells]
    return "\n".join(parts)


def _measured_only_rows():
    """One measured figure, one shipped default, nothing chosen."""
    return [
        A.Assumption(key="cim.population_3mi", label="3-Mile Population",
                     group=A.G_DEAL, value=87_450, provenance=A.EXTERNAL,
                     detail="Census API — ring centred on the ZCTA centroid"),
        A.Assumption(key="GATES.min_irr_5yr", label="Minimum IRR",
                     group=A.G_RETURNS, value=0.10, provenance=A.CONFIG),
    ]


def test_a_measured_run_is_not_described_as_all_shipped_defaults():
    """The memo half of the residual. With `external` outside `CHOSEN` —
    correct — `chosen` is 0 on a run whose population came off the Census
    API, and B.1's empty branch then told an IC reader that every input
    came from the CIM as stated. It did not.

    MUTATION: restore the single unconditional empty-B.1 sentence.
    """
    text = _register_text(_measured_only_rows())

    assert "every input came from the CIM as stated" not in text
    assert "apart from the measured figures noted above" in text


def test_the_appendix_headline_does_not_claim_a_measurement_is_a_default():
    """The lead sentence counted `chosen` and called it "came from
    something other than the model's shipped defaults", which reads as a
    complete partition of the register. It is not one: a measurement is
    neither a default nor chosen, so at `chosen == 0` the sentence stated
    something false about the number Gate 1 turns on.

    MUTATION: revert the headline wording — "0 came from something other
    than the model's shipped defaults" reappears above a measured row.
    """
    text = _register_text(_measured_only_rows())

    assert "came from something other than the model's shipped defaults" \
        not in text
    assert "came from a human or a fallback" in text


def test_the_measurement_gets_its_own_sentence_like_a_transcribed_page():
    """#100's shape, reused: a fact that is not a chosen assumption but
    must not be reachable only by reading B.2 in full. It names the
    figures, so a reader knows WHICH number was measured.

    MUTATION: drop the note paragraph — the disclosure survives only in
    B.2's hundred-and-forty-row table.
    """
    text = _register_text(_measured_only_rows())

    assert "Note on external data" in text
    assert "3-Mile Population" in text
    assert "measured by this system from public data" in text


def test_the_note_is_absent_when_nothing_was_measured():
    """It is rendered only when it happened — the same contract the
    transcription note keeps. A standing "nothing was measured" sentence
    in every memo is noise that trains a reader to skip the paragraph.
    """
    rows = [r for r in _measured_only_rows() if r.provenance != A.EXTERNAL]
    text = _register_text(rows)

    assert "Note on external data" not in text
    assert "every input came from the CIM as stated" in text


def test_the_results_page_counts_a_measurement_separately(tmp_path):
    """The web surface. `_tab_summary.html` printed "all model defaults"
    off `chosen == 0`; it needs `external` in the same counts dict to say
    anything else, and `summarize` is where both come from.

    MUTATION: drop `external` from `PROVENANCE_LABELS` — `summarize` stops
    emitting the key and the template silently falls back to the
    all-defaults branch, since a missing key is falsy in a Django template.
    """
    counts = A.summarize(_measured_only_rows())
    assert counts["chosen"] == 0
    assert counts["external"] == 1
