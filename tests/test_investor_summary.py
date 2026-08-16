"""Item G — the LP-facing 2-page investor summary.

Three things are worth testing and they are different:

1. **It is a second RENDERING, not a second computation.** Every figure
   traces to the result dicts the IC memo receives. The fixtures below
   feed deliberately distinctive values and assert those exact values
   appear, so a helper that quietly recomputed anything would print a
   different number and fail.

2. **The two-page guarantee is a CONTENT BUDGET, not a page count.**
   python-docx does not paginate. The writer constrains Word's layout —
   pinned geometry, EXACTLY line spacing, EXACTLY row heights — and
   `output.page_budget` refuses a document whose content exceeds the
   page. These tests assert the budget numbers, the pinned geometry that
   makes the budget meaningful, and the calibration table underneath it.
   The opt-in `soffice` test is the only thing that re-validates the
   calibration against a real renderer.

3. **The disclosures cannot be dropped.** The securities legend and the
   levered assumption stamp are asserted directly, including on the
   degraded paths, because those are the renderings most likely to skip
   a block.
"""

import shutil
import subprocess

import pytest
from docx import Document
from docx.shared import Inches

import config as cfg
from output import MAX_FILENAME_STEM, safe_filename
from output.memo_writer import (_GC_PENDING_NOTICE, _IS_BODY_PT, _IS_HEAD_PT,
                                _IS_MAX_MITIGANT_CHARS, _IS_MAX_NAME_CHARS,
                                _IS_MICRO_PT, _SUMMARY_LEGEND, _derive_thesis,
                                _is_build, generate_investor_summary)
from output.page_budget import (InvestorSummaryOverflow,
                                InvestorSummaryUnderflow, MARGIN_X_IN,
                                MARGIN_Y_IN, PAGE_BUDGET_PT, PAGE_HEIGHT_IN,
                                PAGE_MIN_PT, PAGE_WIDTH_IN, estimate_lines,
                                text_width_pt)

_BUILD_ARGS = ("property_name", "cim_data", "market_analysis",
               "physical_analysis", "scenario_results", "risk_analysis",
               "rent_analysis", "value_add", "va_results", "gate_results",
               "gate_summary", "check_summary", "sources_uses", "levered",
               "debt", "assumption_fill_log", "thesis")


# ── Fixtures ─────────────────────────────────────────────────────────

class _CIM:
    def __init__(self, **kw):
        defaults = dict(
            property_name="Sunbelt Storage Portfolio",
            address="4100 Industrial Parkway", city="Abilene", state="TX",
            asking_price=10_000_000, price_per_sf=200.0, nrsf=50_000,
            physical_occupancy=0.92, economic_occupancy=0.78,
        )
        defaults.update(kw)
        for k, v in defaults.items():
            setattr(self, k, v)


def _scenarios(**over):
    def one(irr, moic, exit_cap):
        d = {"irr": irr, "moic": moic, "exit_cap": exit_cap,
             "entry_cap": 0.065, "yield_on_cost": 0.0712,
             "noi_projection": [650_000, 690_000, 720_000, 750_000, 780_000],
             # Deliberately DIFFERENT from sources_uses["total_uses"]
             # (10,350,000): total_uses carries financing costs and
             # total_basis does not, so a test that used one number for
             # both could not tell them apart.
             "total_basis": 10_250_000,
             "hold_years": 5}
        d.update(over)
        return d
    return {"bear": one(0.041, 1.19, 0.0785),
            "base": one(0.1234, 1.87, 0.0685),
            "bull": one(0.211, 2.44, 0.0585)}


STAMP = [
    {"key": "pref_rate", "label": "8.00% preferred return, annually compounded"},
    {"key": "promote_split", "label": "20% promote on the LP residual"},
    {"key": "am_fee_treatment",
     "label": "AM fee above the waterfall - 1.00% of LP equity"},
]
LEVERED = {"base": {"lp_net_irr": 0.1567, "lp_moic": 2.03,
                    "am_fee_total": 208_000, "am_fee_pct": 0.01,
                    "gp_promote": 415_000, "assumption_stamp": STAMP},
           "bear": {"lp_net_irr": 0.012, "lp_moic": 1.05},
           "bull": {"lp_net_irr": 0.2891, "lp_moic": 2.99}}

SOURCES_USES = {"total_uses": 10_350_000, "senior_debt": 6_400_000,
                "total_equity": 3_950_000, "gp_equity": 395_000,
                "lp_equity": 3_555_000, "gp_coinvest_pct": 0.10}

DEBT = {"loan": 6_400_000, "binding_constraint": "max_ltv"}

PHYSICAL = {"property_profile": {
                "property_name": "Sunbelt Storage Portfolio",
                "city_state": "Abilene, TX", "year_built": 2005,
                "nrsf": 50_000, "total_units": 420, "cc_pct": 0.45,
                "physical_occupancy": 0.92, "economic_occupancy": 0.78},
            "price_vs_replacement": {
                "comparable": True, "asking_price": 10_000_000,
                "asking_per_sf": 200.0, "replacement_cost": 13_000_000,
                "discount_to_replacement": 0.2308}}

MARKET = {"demographics": {"population_3mi": 87_450, "median_hhi_3mi": 64_300},
          "msa_info": {"msa_name": "Abilene", "is_top_50": False}}

GATES = [{"gate": 5, "name": "No Oversupply Flag",
          "actual": "6.2 SF/capita", "result": "PASS"}]

RISKS = {"risks": [
    {"risk": "Low severity item", "severity": "Low", "mitigation": "Monitor"},
    {"risk": "Property tax reassessment on sale", "severity": "High",
     "mitigation": "Model reassessed basis at the state formula"},
    {"risk": "Economic occupancy 14 points below physical", "severity": "High",
     "mitigation": "Audit concessions and delinquency in diligence"},
    {"risk": "New supply within 3 miles", "severity": "Medium",
     "mitigation": "Confirm pipeline with the planning department"}],
    "why_deal_could_fail": ["The in-place-to-market rent gap closes from "
                            "above in a falling street-rate market."]}

VALUE_ADD = {"revenue_opportunities": [
                 {"category": "Economic Occupancy Recovery",
                  "description": "Close the 14-point collection gap.",
                  "est_annual_impact": 118_000, "timeline": "6-12 months"},
                 {"category": "Revenue Management / ECRI",
                  "description": "Institute a quarterly ECRI programme.",
                  "est_annual_impact": 62_000, "timeline": "12-18 months"}],
             "expense_opportunities": [
                 {"category": "Payroll", "description": "Move to kiosk staffing.",
                  "est_annual_impact": 24_000, "timeline": "6 months"}],
             "estimated_noi_uplift": 204_000}

VA_RESULTS = {"base": {"current_occupancy": 0.78, "target_occupancy": 0.88,
                       "months_to_stabilize": 24, "in_place_rent_psf": 11.40,
                       "target_rent_psf": 13.10, "market_rent_psf": 13.75}}

RENT = {"rent_gap_pct": 0.14}
GATE_SUMMARY = {"recommendation": "PROCEED TO DILIGENCE"}
CHECK_SUMMARY = {"blocking_failed": 0}


def _kwargs(**over):
    kw = dict(property_name="Sunbelt Storage Portfolio", cim_data=_CIM(),
              market_analysis=MARKET, physical_analysis=PHYSICAL,
              scenario_results=_scenarios(), risk_analysis=RISKS,
              rent_analysis=RENT, value_add=VALUE_ADD, va_results=VA_RESULTS,
              gate_results=GATES, gate_summary=GATE_SUMMARY,
              check_summary=CHECK_SUMMARY, sources_uses=SOURCES_USES,
              levered=LEVERED, debt=DEBT, assumption_fill_log=None,
              thesis=None)
    kw.update(over)
    return kw


def _build(**over):
    kw = _kwargs(**over)
    return _is_build(**{k: kw[k] for k in _BUILD_ARGS})


def _generate(tmp_path, **over):
    return generate_investor_summary(output_dir=str(tmp_path), **_kwargs(**over))


def _text(path) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts += [c.text for c in row.cells]
    return "\n".join(parts)


def longest_realistic_deal():
    """Worst case that could actually occur: a name at the cap, three
    scenarios, three risks each at the mitigant cap, widest currency."""
    long_mit = "M" * _IS_MAX_MITIGANT_CHARS
    return _kwargs(
        property_name="P" * _IS_MAX_NAME_CHARS,
        cim_data=_CIM(property_name="P" * _IS_MAX_NAME_CHARS,
                      asking_price=987_654_321, price_per_sf=1234.56),
        risk_analysis={"risks": [
            {"risk": "R" * 120, "severity": s, "mitigation": long_mit}
            for s in ("High", "High", "Medium")],
            "why_deal_could_fail": ["F" * 200]},
        sources_uses={**SOURCES_USES, "total_uses": 987_654_321,
                      "total_equity": 456_789_012, "senior_debt": 530_865_309},
    )


# ── The two-page content budget ──────────────────────────────────────

def test_budget_fits_both_pages():
    _, p1, p2 = _build()
    assert p1.total_pt <= PAGE_BUDGET_PT, f"page 1 {p1.total_pt:.0f}pt"
    assert p2.total_pt <= PAGE_BUDGET_PT, f"page 2 {p2.total_pt:.0f}pt"
    # An under-filled page 2 means the ladder ate content that belonged.
    assert p2.total_pt >= PAGE_MIN_PT, f"page 2 only {p2.total_pt:.0f}pt"


def test_budget_fits_the_longest_realistic_deal():
    kw = longest_realistic_deal()
    _, p1, p2 = _is_build(**{k: kw[k] for k in _BUILD_ARGS})
    assert p1.total_pt <= PAGE_BUDGET_PT, f"page 1 {p1.total_pt:.0f}pt"
    assert p2.total_pt <= PAGE_BUDGET_PT, f"page 2 {p2.total_pt:.0f}pt"


# ── Portfolio notice (portfolio-suspect CIMs) ────────────────────────

_PORTFOLIO_SIGNAL = {
    "is_portfolio": True,
    "evidence": ["multiple cities on the cover page: Granville, New Albany"]}


def test_portfolio_notice_renders_and_is_budgeted(tmp_path):
    """The LP document is the one surface that leaves the firm, so a
    portfolio-suspect deal must carry the caveat there — charged to the
    page-1 budget like every real block."""
    _, p1, _ = _build(cim_data=_CIM(portfolio_signal=_PORTFOLIO_SIGNAL))
    assert any(lbl == "portfolio/notice" for lbl, _ in p1.blocks)
    assert p1.total_pt <= PAGE_BUDGET_PT, f"page 1 {p1.total_pt:.0f}pt"

    text = _text(_generate(
        tmp_path, cim_data=_CIM(portfolio_signal=_PORTFOLIO_SIGNAL)))
    assert "more than one property" in text
    # Investor-facing: the detection EVIDENCE stays out of the LP document —
    # that is analyst diligence detail, and the IC memo carries it.
    assert "multiple cities" not in text


def test_no_portfolio_notice_for_a_single_asset(tmp_path):
    _, p1, _ = _build()
    assert not any(lbl == "portfolio/notice" for lbl, _ in p1.blocks)
    assert "more than one property" not in _text(_generate(tmp_path))


def test_portfolio_notice_fits_the_longest_realistic_deal():
    kw = longest_realistic_deal()
    kw["cim_data"].portfolio_signal = dict(_PORTFOLIO_SIGNAL)
    _, p1, _ = _is_build(**{k: kw[k] for k in _BUILD_ARGS})
    assert p1.total_pt <= PAGE_BUDGET_PT, f"page 1 {p1.total_pt:.0f}pt"


def test_over_budget_raises_rather_than_shrinking(tmp_path):
    """No silent shrink-to-fit. The document goes to investors; a page
    that quietly lost a block is worse than a loud failure."""
    monster = {"risks": [{"risk": "R" * 400, "severity": "High",
                          "mitigation": "M" * 400} for _ in range(3)],
               "why_deal_could_fail": ["F" * 4000]}
    with pytest.raises(InvestorSummaryOverflow) as exc:
        generate_investor_summary(output_dir=str(tmp_path),
                                  **_kwargs(risk_analysis=monster))
    assert "over by" in str(exc.value)


def test_under_filled_page_two_also_raises():
    """The one-sided assert this replaces would have passed a nearly
    blank page 2. The floor applies only to a COMPLETE deal — see
    `test_thin_deal_still_produces_a_document` for the other side."""
    from output.page_budget import PageBudget
    budget = PageBudget("Page 2")
    budget.add("tiny", 10.0)
    with pytest.raises(InvestorSummaryUnderflow, match="rendered only"):
        budget.check(floor_pt=PAGE_MIN_PT)


# ── The geometry the budget depends on ───────────────────────────────

def test_saved_document_pins_page_and_margins(tmp_path):
    """A later margin edit must fail here, not in an LP's inbox — the
    budget is computed against these exact numbers."""
    doc = Document(_generate(tmp_path))
    section = doc.sections[0]
    assert section.page_width == Inches(PAGE_WIDTH_IN)
    assert section.page_height == Inches(PAGE_HEIGHT_IN)
    assert section.left_margin == Inches(MARGIN_X_IN)
    assert section.right_margin == Inches(MARGIN_X_IN)
    assert section.top_margin == Inches(MARGIN_Y_IN)
    assert section.bottom_margin == Inches(MARGIN_Y_IN)


def test_styles_carry_the_sizes_the_budget_assumes(tmp_path):
    doc = Document(_generate(tmp_path))
    for name, size in (("LPBody", _IS_BODY_PT), ("LPHead", _IS_HEAD_PT),
                       ("LPMicro", _IS_MICRO_PT)):
        assert doc.styles[name].font.size.pt == size
        assert doc.styles[name].font.name == "Calibri"


def test_exactly_one_page_break_and_it_is_page_break_before(tmp_path):
    """`add_page_break()` injects a stray empty paragraph, which is a
    line of unbudgeted height. `page_break_before` does not."""
    xml = Document(_generate(tmp_path)).element.body.xml
    assert xml.count("pageBreakBefore") == 1
    assert 'w:type="page"' not in xml


# ── The calibration underneath the budget ────────────────────────────

@pytest.mark.parametrize("text,size,col_pt,expected", [
    ("", 9, 300, 1),
    ("short", 9, 300, 1),
    ("i" * 100, 9, 300, 1),          # narrow bucket
    ("M" * 100, 9, 300, 3),          # wide bucket
    ("x" * 100, 9, 300, 2),          # default bucket
    ("x" * 100, 9, 150, 4),          # half the column, double the lines
])
def test_golden_calibration_table(text, size, col_pt, expected):
    assert estimate_lines(text, size, col_pt) == expected


def test_bold_is_measured_wider_than_regular():
    assert text_width_pt("Sample", 9, bold=True) > text_width_pt("Sample", 9)


def test_non_latin_is_measured_wide_not_rejected():
    """This guard used to RAISE on non-ASCII. The pipeline immediately
    fed it a "≥" from `analysis.filters`' own gate text and killed the
    document, so it now measures unknown glyphs at full width — the
    direction the budget is allowed to be wrong in."""
    wide = text_width_pt("自自自", 9)
    latin = text_width_pt("xxx", 9)
    assert wide > latin


def test_pipeline_symbols_are_folded_before_they_reach_the_budget(tmp_path):
    """Regression: gate text carries >=, risk text carries arrows. A
    maths symbol in a risk description must not fail the document."""
    body = _text(_generate(tmp_path, risk_analysis={"risks": [
        {"risk": "Population >= 50,000 unverified", "severity": "High",
         "mitigation": "Order a demographic report; occupancy -> target"}]}))
    assert ">= 50,000" in body
    assert "-> target" in body


def test_smart_quotes_from_a_cim_do_not_fail_the_document(tmp_path):
    """Broker PDFs are full of typographic punctuation. It is folded to
    ASCII before it reaches a cell, so it cannot trip the width table."""
    name = "O’Brien’s “Premier” Storage — Abilene"
    body = _text(_generate(tmp_path, cim_data=_CIM(property_name=name),
                           property_name=name))
    assert "O'Brien's" in body


# ── Second rendering, never a second computation ─────────────────────

def test_target_return_box_shows_both_lenses_and_the_delta(tmp_path):
    body = _text(_generate(tmp_path))
    assert "12.3%" in body          # unlevered base IRR
    assert "15.7%" in body          # LP net base IRR
    assert "1.87x" in body and "2.03x" in body
    # Leverage delta: 15.67% - 12.34% = +333 bps, a difference between
    # two published figures.
    assert "+333 bps" in body
    assert "accretive" in body


def test_dilutive_leverage_is_stated_plainly(tmp_path):
    """On config defaults the loan constant frequently sits above yield
    on cost, so LP net can land BELOW unlevered. A sophisticated
    allocator spots that instantly; the document must not bury it."""
    dilutive = {**LEVERED, "base": {**LEVERED["base"], "lp_net_irr": 0.0912}}
    body = _text(_generate(tmp_path, levered=dilutive))
    assert "DILUTIVE" in body
    assert "-322 bps" in body


def test_figures_trace_to_the_result_dicts(tmp_path):
    body = _text(_generate(tmp_path))
    assert "$10,000,000" in body            # asking price
    assert "23.1%" in body                  # discount to replacement
    assert "6.500%" in body                 # entry cap (3dp, as _fmt_cap)
    assert "6.850%" in body                 # base exit cap
    assert "$6,400,000" in body             # senior debt
    assert "$10,350,000" in body            # total uses (capital stack)


def test_fee_and_promote_load_is_volunteered(tmp_path):
    """This audience asks what was deducted to reach an LP net number
    within ten minutes. Printing it converts an interrogation into a
    checked box."""
    body = _text(_generate(tmp_path))
    assert "$208,000" in body               # AM fee, hold total
    assert "1.00%/yr" in body
    assert "$415,000" in body               # GP promote, hold total
    assert "10% of equity" in body          # GP co-invest


def test_levered_figures_carry_their_assumption_stamp(tmp_path):
    """CLAUDE.md key design decision 7. This is the only surface that
    leaves the firm, so the reader has no memo section 6 to cross-check
    which conventions produced the number."""
    body = _text(_generate(tmp_path))
    for row in STAMP:
        assert row["label"] in body
    assert "subject to the final partnership agreement" in body


def test_unlevered_deal_gains_no_orphaned_stamp(tmp_path):
    body = _text(_generate(tmp_path, levered=None))
    assert "LP net returns are computed under" not in body


def _stamped(tmp_path, statuses):
    """Render the summary with the STAMP rows carrying `statuses`."""
    stamp = [{**row, "status": s} for row, s in zip(STAMP, statuses)]
    levered = {**LEVERED,
               "base": {**LEVERED["base"], "assumption_stamp": stamp}}
    return _text(_generate(tmp_path, levered=levered))


def test_a_confirmed_convention_is_not_called_a_proposed_term(tmp_path):
    """This document goes to investors, so overstating and understating
    are both costly. Calling a convention the operator confirmed against
    the executed LPA "proposed, subject to the final partnership
    agreement" understates it — and blanketing all five with that caveat
    is what the single sentence used to do."""
    body = _stamped(tmp_path, ["confirmed", "confirmed", "confirmed"])
    assert "3 of 3 confirmed against the executed partnership agreement" in body
    assert "proposed terms" not in body


def test_a_moot_row_is_never_counted_as_a_confirmed_one(tmp_path):
    """The overstatement this sentence exists to avoid, caught by the
    characterization snapshot when the first version of it read "2 of 5
    are confirmed" for a deal where the operator had confirmed ONE.

    Nobody read the LPA's ordering clause — it stopped being able to move
    a dollar once the pref was confirmed to compound. Reporting that as a
    confirmation claims more of the document has been read than has been.
    """
    body = _stamped(tmp_path, ["confirmed", "open", "moot"])
    assert "1 of 3 confirmed against the executed partnership agreement" in body
    assert "1 made moot by it" in body
    assert "the rest are proposed terms" in body
    assert "2 of 3 confirmed" not in body


def test_a_fully_settled_stamp_drops_the_proposed_caveat_entirely(tmp_path):
    body = _stamped(tmp_path, ["confirmed", "moot", "moot"])
    assert "1 of 3 confirmed" in body and "2 made moot by it" in body
    assert "proposed terms" not in body


def test_an_all_open_stamp_keeps_the_original_caveat_verbatim(tmp_path):
    """The status field is additive: a stamp from before it existed, or
    one where nothing is confirmed, must read exactly as it always did."""
    body = _stamped(tmp_path, ["open", "open", "open"])
    assert ("These are proposed terms, subject to the final partnership "
            "agreement.") in body
    # And a row with no `status` key at all defaults to the same place.
    assert ("These are proposed terms, subject to the final partnership "
            "agreement.") in _text(_generate(tmp_path))


# ── The plan to achieve the return ───────────────────────────────────

def test_plan_section_carries_initiatives_impacts_and_the_bridge(tmp_path):
    body = _text(_generate(tmp_path))
    assert "Plan to Achieve the Return" in body
    assert "$204,000" in body                       # total NOI uplift
    assert "Economic Occupancy Recovery" in body
    assert "$118,000" in body and "6-12 months" in body
    # The quantified bridge from the monthly engine.
    assert "occupancy 78% -> 88%" in body
    assert "24 months" in body
    assert "$11.40" in body and "$13.10" in body and "$13.75" in body


def test_stabilized_deal_falls_back_to_sources_of_return(tmp_path):
    """No value-add narrative exists for a stabilized asset, so the
    section says where the return comes from arithmetically — all
    differences between figures printed on page 1."""
    body = _text(_generate(tmp_path, value_add=None, va_results=None))
    assert "stabilized asset" in body
    assert "+62 bp spread" in body                  # 7.12% YoC - 6.50% cap
    assert "20%" in body                            # NOI growth across hold
    assert "+35 bp wider" in body                   # 6.85% exit - 6.50% entry


# ── Risks always carry mitigants ─────────────────────────────────────

def test_every_rendered_risk_carries_its_mitigant(tmp_path):
    body = _text(_generate(tmp_path))
    assert body.count("Mitigant:") == 3
    assert "Audit concessions and delinquency" in body
    assert "Primary failure mode:" in body


def test_risks_are_ordered_by_severity_and_capped(tmp_path):
    body = _text(_generate(tmp_path))
    assert body.index("Property tax reassessment") < body.index("New supply")
    assert "Low severity item" not in body          # capped at 3


def test_mitigants_are_truncated_but_never_dropped(tmp_path):
    body = _text(_generate(tmp_path, risk_analysis={"risks": [
        {"risk": "A risk", "severity": "High", "mitigation": "M" * 400}]}))
    assert "Mitigant:" in body
    assert "..." in body


def test_a_risk_without_a_mitigation_still_gets_a_line(tmp_path):
    body = _text(_generate(tmp_path, risk_analysis={"risks": [
        {"risk": "Unmitigated", "severity": "High"}]}))
    assert "Mitigant: Under diligence." in body


# ── Thesis: derived, or overridden ───────────────────────────────────

def test_thesis_ladder_is_gated_on_computed_values():
    bullets = _derive_thesis(_CIM(), PHYSICAL, RENT, VALUE_ADD, MARKET)
    assert len(bullets) == 3
    assert "23% below replacement cost" in bullets[0]
    assert "14 points" in bullets[1]                # 92% - 78%
    assert "14% below market" in bullets[2]


def test_thesis_omits_the_spread_bullet_below_the_config_threshold():
    flag = cfg.GATES["econ_phys_spread_flag"]
    cim = _CIM(physical_occupancy=0.92, economic_occupancy=0.92 - flag / 2)
    bullets = _derive_thesis(cim, PHYSICAL, RENT, VALUE_ADD, MARKET)
    assert not any("trails physical" in b for b in bullets)


def test_spread_bullet_survives_a_zero_economic_occupancy():
    """A property collecting nothing is the loudest mismanagement case;
    truthiness on the pair silently dropped it."""
    bullets = _derive_thesis(_CIM(physical_occupancy=0.85,
                                  economic_occupancy=0.0),
                             PHYSICAL, RENT, VALUE_ADD, MARKET)
    assert any("85 points" in b for b in bullets)


def test_operator_override_replaces_the_derived_thesis(tmp_path):
    body = _text(_generate(tmp_path, thesis=["Operator wrote this one."]))
    assert "Operator wrote this one." in body
    assert "below replacement cost" not in body


# ── Market snapshot ──────────────────────────────────────────────────

def test_msa_is_read_from_the_key_market_analysis_actually_emits(tmp_path):
    """`analysis.market._assess_msa` emits `msa_name`, never `msa`."""
    from analysis.market import _assess_msa

    class _M:
        msa = "Dallas"
        city = "Dallas"
    real = _assess_msa(_M())
    assert "msa_name" in real
    body = _text(_generate(tmp_path, market_analysis={
        "demographics": {"population_3mi": 87_450}, "msa_info": real}))
    assert "Dallas - top-50 MSA" in body


def test_no_msa_rank_is_ever_printed(tmp_path):
    """`msa_info` carries a bool and a name. There is no rank number in
    this codebase; inventing one would be a fabricated fact."""
    assert "#" not in _text(_generate(tmp_path))


def test_sf_per_capita_is_conditional_on_not_being_tbd(tmp_path):
    assert "6.2 SF/capita" in _text(_generate(tmp_path))
    hidden = _text(_generate(tmp_path, gate_results=[
        {"gate": 5, "name": "No Oversupply Flag", "actual": "TBD"}]))
    assert "Supply (3-mile)" not in hidden


# ── Degradation ──────────────────────────────────────────────────────

def test_thin_deal_still_produces_a_document(tmp_path):
    path = generate_investor_summary(
        output_dir=str(tmp_path), property_name="Thin Deal", cim_data=_CIM(),
        market_analysis={}, physical_analysis={}, scenario_results={},
        risk_analysis={})
    body = _text(path)
    assert "Thin Deal" in body
    assert "Market Snapshot" not in body
    assert "Scenario Returns" not in body
    assert _SUMMARY_LEGEND in body


def test_every_document_carries_the_securities_legend(tmp_path):
    for over in ({}, {"levered": None}, {"va_results": None},
                 {"risk_analysis": {}}, {"sources_uses": None}):
        assert _SUMMARY_LEGEND in _text(_generate(tmp_path, **over))


def test_screening_result_and_blocking_checks_reach_the_footer(tmp_path):
    body = _text(_generate(tmp_path, check_summary={"blocking_failed": 2}))
    assert "PROCEED TO DILIGENCE" in body
    assert "2 blocking model check(s) unresolved" in body


# ── The shared filename helper ───────────────────────────────────────

def test_long_names_do_not_collide_across_the_writers():
    """`excel_writer`'s private copy was uncapped, and `generate_excel`
    is called UNWRAPPED, so a long name aborted the whole run."""
    a = "Fund IV Self-Storage Portfolio - " + "X" * 60 + " - Abilene, TX"
    b = "Fund IV Self-Storage Portfolio - " + "X" * 60 + " - Waco, TX"
    assert len(safe_filename(a)) <= MAX_FILENAME_STEM
    assert safe_filename(a) != safe_filename(b)
    assert safe_filename("Sunbelt Storage") == "Sunbelt_Storage"


# ── Opt-in: the only test that re-validates the calibration ──────────

def pdf_page_count(raw: bytes) -> int:
    """Pages in a PDF, from its bytes, without a PDF library.

    Extracted from the render test and unit-tested below BECAUSE it is
    the fragile half. `/Type /Page` is a PREFIX of `/Type /Pages`, the
    page-tree root — so the obvious `raw.count(b"/Type /Page")` returns
    N+1, and the calibration test it backed could only ever have passed
    on a one-page document. That defect survived because the test around
    it has never run: it skips wherever LibreOffice is absent, which is
    everywhere.

    Returns 0 when the structure is compressed into object streams,
    which the caller treats as "cannot count" rather than "zero pages" —
    `pdfinfo` is used in preference wherever it exists for exactly that
    reason.
    """
    import re

    return len(re.findall(rb"/Type\s*/Page(?![s/\w])", raw))


def _render_page_count(docx_path, tmp_path, soffice) -> int:
    """Render to PDF and count, preferring `pdfinfo` over byte-scraping.

    Two counters, and the preference order is the point: `pdfinfo` is
    authoritative and `pdf_page_count` is a heuristic over bytes that a
    future LibreOffice could compress out of sight. Falling back rather
    than requiring poppler keeps the LOCAL opt-in run possible with only
    LibreOffice installed; CI installs both and so never takes the
    fallback.
    """
    subprocess.run([soffice, "--headless", "--convert-to", "pdf",
                    "--outdir", str(tmp_path), str(docx_path)],
                   check=True, capture_output=True, timeout=300)
    pdf = next(p for p in tmp_path.iterdir() if p.suffix == ".pdf")

    if shutil.which("pdfinfo"):
        out = subprocess.run(["pdfinfo", str(pdf)], check=True,
                             capture_output=True, timeout=60).stdout.decode()
        line = next(ln for ln in out.splitlines() if ln.startswith("Pages:"))
        return int(line.split(":", 1)[1].strip())

    pages = pdf_page_count(pdf.read_bytes())
    assert pages, ("could not count pages: no pdfinfo, and the PDF's object "
                   "structure is not scannable — install poppler-utils")
    return pages


def test_real_render_is_two_pages(tmp_path):
    """The ONLY thing that re-validates the content budget against a real
    renderer. Everything else in this file asserts the budget's own
    arithmetic, which cannot notice that the arithmetic models the wrong
    thing.

    **`CIM_REQUIRE_SOFFICE=1` turns the skip into a failure**, and the
    calibration CI job sets it. Without that, an apt-get step that
    silently installed nothing would leave this test skipping and the job
    reporting green — a calibration check that validates nothing while
    looking like it does, which is worse than not having one.
    """
    import os

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice is None:
        if os.environ.get("CIM_REQUIRE_SOFFICE"):
            pytest.fail(
                "CIM_REQUIRE_SOFFICE is set but neither soffice nor "
                "libreoffice is on PATH — the calibration job installed "
                "nothing and would otherwise have reported green")
        pytest.skip("LibreOffice not installed; the calibration check is "
                    "opt-in locally. CI runs it in the `page-budget` job; "
                    "set CIM_REQUIRE_SOFFICE=1 to make its absence fail.")

    pages = _render_page_count(_generate(tmp_path), tmp_path, soffice)
    assert pages == 2, (
        f"rendered {pages} pages, not 2 — the content budget in "
        f"output/page_budget.py no longer models what Word does. Read that "
        f"module's docstring before changing a bucket width.")


@pytest.mark.parametrize("raw,expected", [
    # The trap: one page-tree root plus two pages is TWO pages, not three.
    (b"<</Type /Pages /Count 2>> <</Type /Page>> <</Type /Page>>", 2),
    # LibreOffice omits the space; the root must still not be counted.
    (b"<</Type/Pages/Count 2>><</Type/Page>><</Type/Page>>", 2),
    (b"<</Type /Pages /Count 1>> <</Type /Page /MediaBox[0 0 612 792]>>", 1),
    # Compressed object streams: unscannable, and reported as such rather
    # than as a confident zero.
    (b"%PDF-1.7\n<</Type/ObjStm/N 8>>stream\x00\x01binary", 0),
])
def test_the_page_counter_does_not_count_the_page_tree_root(raw, expected):
    """The unit test the render test cannot be: it runs everywhere,
    including on the box that has no LibreOffice, so the half of the
    calibration check most likely to be wrong is covered even when the
    half that needs a renderer cannot run."""
    assert pdf_page_count(raw) == expected


# ── Review findings ──────────────────────────────────────────────────

def test_total_basis_is_the_unlevered_figure_not_total_uses(tmp_path):
    """CLAUDE.md decision 3: total basis EXCLUDES financing costs, and
    `total_uses == total_basis + financing_costs`. Printing total_uses
    under a "Total Basis" label put a financing-inflated number on the
    one document built to keep the unlevered lens financing-free."""
    body = _text(_generate(tmp_path))
    rows = [ln.strip() for ln in body.splitlines()]
    basis_idx = rows.index("Total Basis")
    assert rows[basis_idx + 1] == "$10,250,000"      # scenario total_basis
    assert rows[basis_idx + 1] != "$10,350,000"      # NOT total_uses


def test_no_levered_figure_survives_a_missing_assumption_stamp(tmp_path):
    """CLAUDE.md decision 7, made structural. The figures and the stamp
    used to be gated independently, so a payload carrying a levered
    scenario but no stamp printed a bare LP net IRR — on the only
    document that leaves the firm."""
    stampless = {"base": {"lp_net_irr": 0.1567, "lp_moic": 2.03,
                          "am_fee_total": 208_000, "gp_promote": 415_000},
                 "bear": {"lp_net_irr": 0.012},
                 "bull": {"lp_net_irr": 0.2891}}
    body = _text(_generate(tmp_path, levered=stampless))
    for leaked in ("15.7%", "2.03x", "$208,000", "$415,000", "28.9%"):
        assert leaked not in body
    assert "LP net returns are computed under" not in body
    # The unlevered lens is untouched — it never depended on the stamp.
    assert "12.3%" in body


def test_bear_bull_band_alone_cannot_print_without_a_stamp(tmp_path):
    """The band reads bear/bull directly, so it needed the same gate."""
    body = _text(_generate(tmp_path, levered={"bear": {"lp_net_irr": 0.012},
                                              "bull": {"lp_net_irr": 0.2891}}))
    assert "1.2%" not in body and "28.9%" not in body


# ── Item G: the distribution gate ────────────────────────────────────
#
# The gate used to live in a backlog paragraph and a code comment, which
# are the two places the analyst clicking "Investor Summary (.docx)" will
# never look. These tests hold it as behaviour instead.

def test_an_uncleared_summary_says_so_on_its_own_first_line(tmp_path,
                                                            monkeypatch):
    """MUTATION: delete the `_is_gc_notice` call from `_is_build`.

    On the PAGE, not only on screen: the failure this guards is a file
    already detached from the app — attached to an email, sitting in a
    data room — and a caveat beside the download button is invisible the
    moment the .docx moves.
    """
    monkeypatch.setattr(cfg, "INVESTOR_SUMMARY_GC_CLEARED", False)
    body = _text(_generate(tmp_path))
    assert _GC_PENDING_NOTICE in body
    assert body.strip().startswith("INTERNAL DRAFT")


def test_clearing_the_gate_removes_the_notice_and_keeps_the_legend(
        tmp_path, monkeypatch):
    """The flag has to actually be a flag — a notice that cannot be
    removed is one the operator will strip by editing the constant, and
    the legend would go with it.

    The legend is the half that must survive BOTH states: it says what
    the document is, where the notice only says who has not yet read it.
    """
    monkeypatch.setattr(cfg, "INVESTOR_SUMMARY_GC_CLEARED", True)
    body = _text(_generate(tmp_path))
    assert _GC_PENDING_NOTICE not in body
    assert _SUMMARY_LEGEND in body


def test_the_notice_is_measured_not_just_written(tmp_path, monkeypatch):
    """Every block on this document is charged to the page budget, and a
    notice exempted from that is how a document that fits in the tests
    overflows in Word.

    Asserts the DIFFERENCE, so it fails if the notice renders free.
    """
    def _page1_pt(cleared):
        monkeypatch.setattr(cfg, "INVESTOR_SUMMARY_GC_CLEARED", cleared)
        _doc, page1, _page2 = _build()
        return page1.total_pt

    charged = _page1_pt(False) - _page1_pt(True)
    assert charged > 0, "the notice rendered without being charged to the page"
    # And it is charged as the block it is, not as a rounding wobble.
    assert charged >= _IS_MICRO_PT


def test_the_gc_flag_is_not_settings_page_editable():
    """A legal clearance is not a per-deal underwriting assumption, and an
    analyst must not be able to clear it from the screen that edits cap
    rates. It is absent from the override registry by construction; this
    says so out loud, because the registry is derived LIVE from config
    and a future refactor could sweep it in."""
    from webapp.forms import override_key_registry

    keys = override_key_registry()
    assert not [k for k in keys if "GC_CLEARED" in k.upper()]


def test_the_notice_survives_the_punctuation_fold(tmp_path, monkeypatch):
    """`_is_para` folds typographic punctuation to ASCII so the budget can
    measure it, so a constant containing an em dash would differ from the
    text on the page — and every assertion above would pass while the
    notice was subtly not the constant. Caught in development, pinned
    here."""
    monkeypatch.setattr(cfg, "INVESTOR_SUMMARY_GC_CLEARED", False)
    assert _GC_PENDING_NOTICE.isascii()
    assert _SUMMARY_LEGEND.isascii()
    assert _GC_PENDING_NOTICE in _text(_generate(tmp_path))
